import os
from copy import deepcopy

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from worker.docx_renderer import normalize_text, ensure_dict, safe_list, format_emission_date


FICHA_TEMPLATE_PATH = os.path.join(os.path.dirname(__file__), "assets", "template_ficha_cadastro_erosao.docx")

# --- Cell map: (row, unique_cell_index) for each field in the template table ---
# Row 0: EMPREENDIMENTO (1 merged cell spanning all 12 cols)
# Row 1: CADASTRO DE FOCOS EROSIVOS (title, skip)
# Row 2: Ficha n (cols 0-8), Data (cols 9-11)
# Row 3: Profissional (1 merged cell)
# Row 4: LOCALIZACAO (section header, skip)
# Row 5: UTM E (cols 0-8), Atitude (cols 9-11)
# Row 6: UTM S (cols 0-8), Fotos (cols 9-11)
# Row 7: area checkboxes: Faixa Servidao (col 0), Area Terceiros (col 3), Area Publica (col 9)
# Row 8: Referencia (1 merged cell)
# Row 9: CLASSIFICACAO CRITICIDADE (section header, skip)
# Row 10: criticidade checkboxes: Baixo (col 0), Medio (col 2), Alto (col 6), Muito Alto (col 10)
# Row 11: SITUACAO ATUAL (section header, skip)
# Row 12: Estagio Erosivo label (col 0), Ativo (col 2), Estavel (col 6), Regen. Natural (col 10)
# Row 13: TIPO/CARACTERISTICAS (section header, skip)
# Row 14: Laminar (col 0), Sulco (col 2), Ravina (col 6), Vocoroca (col 10)
# Row 15: Presenca agua label (col 0), Sim (col 4), Nao (col 7), Nao verificado (col 11)
# Row 16: DECLIVIDADE (section header, skip)
# Row 17: 0-6 (col 0), 6-12 (col 2), 12-20 (col 6), >20 (col 10)
# Row 18: DIMENSOES (section header, skip)
# Row 19: Largura label (col 0), Ate 1m (col 2), 1-10m (col 6), >30m (col 10)
# Row 20: Altura label (col 0), Ate 1m (col 2), 1-10m (col 6), >30m (col 10)
# Row 21: CARACTERIZACAO (section header, skip)
# Row 22: Relevo label (col 0), Suave (col 2), Ondulado (col 6), Escarpado (col 10)
# Row 23: Tipo Solo label (col 0), Argiloso (col 2), Arenoso (col 6), Lateritico (col 10)
# Row 24: Usos Solo label (col 0), Pastagem (col 1), Cultivo (col 5), Campo (col 8), Veg Arborea (col 11)
# Row 25: CARACTERIZACAO (section header, skip)
# Row 26: Obstaculos label (col 0), Acesso (col 1), Cerca (col 5), Curso dagua (col 8), Tubulacao (col 11)
# Row 27: Outros (1 merged cell)
# Row 28: MEDIDA PREVENTIVA (1 merged cell)
# Row 29: OBSERVACOES - CROQUIS (1 merged cell)
# Row 30: empty (1 merged cell)


CHECKED = "( X )"
UNCHECKED_PATTERNS = ["(   )", "(    )"]


def _get_unique_cells(row):
    """Return list of (col_index, cell) for unique cells in a row (skip merged duplicates)."""
    seen = set()
    result = []
    for k, cell in enumerate(row.cells):
        cid = id(cell._tc)
        if cid not in seen:
            seen.add(cid)
            result.append((k, cell))
    return result


def _set_cell_text(cell, text):
    """Replace cell text while preserving the first paragraph's formatting."""
    paragraphs = cell.paragraphs
    if not paragraphs:
        return
    p = paragraphs[0]
    runs = list(p.runs)
    if not runs:
        p.text = text
        return
    runs[0].text = text
    for run in runs[1:]:
        run.text = ""


def _set_checkbox(cell, checked):
    """Toggle a checkbox cell between checked/unchecked state."""
    p = cell.paragraphs[0] if cell.paragraphs else None
    if not p:
        return
    current_text = p.text
    if checked:
        for pattern in UNCHECKED_PATTERNS:
            if pattern in current_text:
                new_text = current_text.replace(pattern, CHECKED, 1)
                runs = list(p.runs)
                if runs:
                    runs[0].text = new_text
                    for run in runs[1:]:
                        run.text = ""
                else:
                    p.text = new_text
                return
    # If unchecked, leave as-is (template default is unchecked)


def _find_cell_by_col(unique_cells, col_index):
    """Find a cell by its starting column index in the unique cells list."""
    for idx, cell in unique_cells:
        if idx == col_index:
            return cell
    return None


def _append_text_to_label(cell, label_prefix, value):
    """For cells like 'Ficha n' or 'UTM E:', append value after the label."""
    if not value:
        return
    p = cell.paragraphs[0] if cell.paragraphs else None
    if not p:
        return
    runs = list(p.runs)
    if not runs:
        p.text = f"{label_prefix} {value}"
        return
    # Keep existing runs (with formatting), append value to last run or add new
    last_run = runs[-1]
    current_full = "".join(r.text for r in runs)
    # Replace full text in first run, clear others
    runs[0].text = f"{label_prefix} {value}"
    for run in runs[1:]:
        run.text = ""


def _classify_declividade(graus):
    """Classify slope in degrees to template ranges."""
    if graus is None or not isinstance(graus, (int, float)):
        return None
    if graus <= 6:
        return 0  # 0-6
    if graus <= 12:
        return 1  # 6-12
    if graus <= 20:
        return 2  # 12-20
    return 3  # >20


def _classify_dimension(metros):
    """Classify dimension in meters to template ranges: Ate 1m, 1-10m, >30m."""
    if metros is None or not isinstance(metros, (int, float)):
        return None
    if metros <= 1:
        return 0  # Ate 1 metro
    if metros <= 10:
        return 1  # 1 a 10 metros
    return 2  # Maior que 30 metros


def _resolve_estagio_erosivo(sinais_avanco, vegetacao_interior):
    """Derive erosion stage from activity indicators (same logic as criticality.js resolveActivityScore)."""
    if sinais_avanco and not vegetacao_interior:
        return "ativo"  # A4: active advance without vegetation
    if sinais_avanco and vegetacao_interior:
        return "ativo"  # A3: active advance with vegetation (still active)
    if not sinais_avanco and vegetacao_interior:
        return "regeneracao"  # A1: vegetation, no advance = stabilized/regenerating
    # not sinais_avanco and not vegetacao_interior -> indeterminate, treat as estavel
    return "estavel"


CRITICIDADE_CODE_MAP = {
    "C1": 0,  # Baixo
    "C2": 1,  # Medio
    "C3": 2,  # Alto
    "C4": 3,  # Muito Alto
}

ESTAGIO_COL_MAP = {
    "ativo": 2,      # col 2
    "estavel": 6,    # col 6
    "regeneracao": 10,  # col 10
}

FEICAO_MAP = {
    "laminar": 0,    # col 0
    "sulco": 2,      # col 2
    "ravina": 6,     # col 6
    "vocoroca": 10,  # col 10
    "movimento_massa": 10,  # same as vocoroca
}

PRESENCA_AGUA_MAP = {
    "sim": 4,        # col 4
    "nao": 7,        # col 7
    "nao_verificado": 11,  # col 11
}

TIPO_SOLO_MAP = {
    "argiloso": 2,   # col 2
    "arenoso": 6,    # col 6
    "lateritico": 10,  # col 10
    "solos_rasos": 2,  # map to argiloso position (closest match)
}

USOS_SOLO_MAP = {
    "pastagem": 1,   # col 1
    "cultivo": 5,    # col 5
    "campo": 8,      # col 8
    "veg_arborea": 11,  # col 11
}

OBSTACULOS_MAP = {
    "acesso": 1,     # col 1
    "cerca": 5,      # col 5
    "curso_agua": 8,  # col 8
    "tubulacao": 11,  # col 11
}

AREA_MAP = {
    "faixa_servidao": 0,   # col 0
    "area_terceiros": 3,   # col 3
    "area_publica": 9,     # col 9
}

CRITICIDADE_SOLUCOES = {
    "C1": "Cobertura vegetal (gramineas, ressemeadura); Curvas de nivel, plantio em faixas; Mulching / palhada / biomanta leve",
    "C2": "Barraginhas e pequenos terracos; Sangradouros laterais / lombadas de agua; Canaletas vegetadas / valetas rasas; Hidrossemeadura + biomantas leves",
    "C3": "Reconformacao de taludes; Sarjetas de crista / canaletas revestidas; Escadas hidraulicas / bacias de dissipacao; Check dams (degraus com pedra/gabioes)",
    "C4": "Rede completa de drenagem da bacia; Drenos profundos para piping; Diques de terra / barragens; Estruturas de contencao (muros, gabioes); PRAD especifico",
}


def _fill_ficha_table(table, erosion, project_name):
    """Fill a single ficha table with erosion data."""
    rows = table.rows

    # --- Row 0: EMPREENDIMENTO ---
    cells_r0 = _get_unique_cells(rows[0])
    cell_empreend = _find_cell_by_col(cells_r0, 0)
    if cell_empreend:
        _append_text_to_label(cell_empreend, "EMPREENDIMENTO:", normalize_text(project_name))

    # --- Row 2: Ficha n / Data ---
    cells_r2 = _get_unique_cells(rows[2])
    cell_ficha = _find_cell_by_col(cells_r2, 0)
    cell_data = _find_cell_by_col(cells_r2, 9)
    if cell_ficha:
        _append_text_to_label(cell_ficha, "Ficha n\u00ba", normalize_text(erosion.get("id")))
    if cell_data:
        data_value = format_emission_date(erosion.get("updatedAt") or erosion.get("createdAt"))
        _append_text_to_label(cell_data, "Data:", data_value)

    # --- Row 3: Profissional ---
    cells_r3 = _get_unique_cells(rows[3])
    cell_prof = _find_cell_by_col(cells_r3, 0)
    if cell_prof:
        _append_text_to_label(cell_prof, "Profissional:", normalize_text(erosion.get("updatedBy")))

    # --- Row 5: UTM E / Altitude ---
    cells_r5 = _get_unique_cells(rows[5])
    cell_utme = _find_cell_by_col(cells_r5, 0)
    cell_alt = _find_cell_by_col(cells_r5, 9)
    location = ensure_dict(erosion.get("locationCoordinates"))
    if cell_utme:
        utm_e = normalize_text(location.get("utmEasting"))
        utm_zone = normalize_text(location.get("utmZone"))
        utm_hemi = normalize_text(location.get("utmHemisphere"))
        fuso_suffix = ""
        if utm_zone:
            fuso_suffix = f" (Fuso {utm_zone}{utm_hemi})" if utm_hemi else f" (Fuso {utm_zone})"
        _append_text_to_label(cell_utme, "UTM E:", f"{utm_e}{fuso_suffix}" if utm_e else "")
    if cell_alt:
        altitude = normalize_text(location.get("altitude"))
        _append_text_to_label(cell_alt, "Atitude:", altitude)

    # --- Row 6: UTM S / Fotos ---
    cells_r6 = _get_unique_cells(rows[6])
    cell_utms = _find_cell_by_col(cells_r6, 0)
    cell_fotos = _find_cell_by_col(cells_r6, 9)
    if cell_utms:
        utm_s = normalize_text(location.get("utmNorthing"))
        _append_text_to_label(cell_utms, "UTM S:", utm_s)
    if cell_fotos:
        fotos = normalize_text(erosion.get("fotos"))
        _append_text_to_label(cell_fotos, "Fotos:", fotos)

    # --- Row 7: Area checkboxes ---
    cells_r7 = _get_unique_cells(rows[7])
    local_contexto = ensure_dict(erosion.get("localContexto"))
    local_tipo = normalize_text(local_contexto.get("localTipo")).lower()
    exposicao = normalize_text(local_contexto.get("exposicao")).lower()

    # Determine which area checkbox to check
    area_to_check = set()
    if local_tipo in ("faixa_servidao", "via_acesso_exclusiva", "base_torre"):
        area_to_check.add("faixa_servidao")
    elif local_tipo == "fora_faixa_servidao" or exposicao == "area_terceiros":
        area_to_check.add("area_terceiros")

    for area_key, col in AREA_MAP.items():
        cell = _find_cell_by_col(cells_r7, col)
        if cell and area_key in area_to_check:
            _set_checkbox(cell, True)

    # --- Row 8: Referencia ---
    cells_r8 = _get_unique_cells(rows[8])
    cell_ref = _find_cell_by_col(cells_r8, 0)
    if cell_ref:
        torre_ref = normalize_text(location.get("reference")) or normalize_text(erosion.get("torreRef"))
        _append_text_to_label(cell_ref, "Refer\u00eancia:", torre_ref)

    # --- Row 10: Criticidade checkboxes ---
    cells_r10 = _get_unique_cells(rows[10])
    crit_code = normalize_text(erosion.get("criticalityCode")).upper()
    crit_col_map = {
        "C1": 0,   # Baixo
        "C2": 2,   # Medio
        "C3": 6,   # Alto
        "C4": 10,  # Muito Alto
    }
    if crit_code in crit_col_map:
        cell = _find_cell_by_col(cells_r10, crit_col_map[crit_code])
        if cell:
            _set_checkbox(cell, True)

    # --- Row 12: Estagio Erosivo ---
    cells_r12 = _get_unique_cells(rows[12])
    sinais = erosion.get("sinaisAvanco")
    vegetacao = erosion.get("vegetacaoInterior")
    if isinstance(sinais, bool) or isinstance(vegetacao, bool):
        estagio = _resolve_estagio_erosivo(bool(sinais), bool(vegetacao))
        col = ESTAGIO_COL_MAP.get(estagio)
        if col is not None:
            cell = _find_cell_by_col(cells_r12, col)
            if cell:
                _set_checkbox(cell, True)

    # --- Row 14: Tipo feicao checkboxes ---
    cells_r14 = _get_unique_cells(rows[14])
    tipos_feicao = safe_list(erosion.get("tiposFeicao"))
    for tipo in tipos_feicao:
        col = FEICAO_MAP.get(normalize_text(tipo).lower())
        if col is not None:
            cell = _find_cell_by_col(cells_r14, col)
            if cell:
                _set_checkbox(cell, True)

    # --- Row 15: Presenca de agua ---
    cells_r15 = _get_unique_cells(rows[15])
    presenca = normalize_text(erosion.get("presencaAguaFundo")).lower()
    if presenca in PRESENCA_AGUA_MAP:
        cell = _find_cell_by_col(cells_r15, PRESENCA_AGUA_MAP[presenca])
        if cell:
            _set_checkbox(cell, True)

    # --- Row 17: Declividade ---
    cells_r17 = _get_unique_cells(rows[17])
    decliv_graus = erosion.get("declividadeGraus")
    if isinstance(decliv_graus, (int, float)):
        idx = _classify_declividade(decliv_graus)
        if idx is not None:
            decliv_cols = [0, 2, 6, 10]
            cell = _find_cell_by_col(cells_r17, decliv_cols[idx])
            if cell:
                _set_checkbox(cell, True)

    # --- Row 19: Largura Maxima ---
    cells_r19 = _get_unique_cells(rows[19])
    profund = erosion.get("profundidadeMetros")
    # Note: the DB stores profundidadeMetros which maps to depth, not width
    # We only fill if dimensionamento data is available
    # Leave blank as per rules - no inference

    # --- Row 20: Altura Maxima ---
    cells_r20 = _get_unique_cells(rows[20])
    # Same: only fill from explicit data, profundidadeMetros maps to depth
    if isinstance(profund, (int, float)):
        idx = _classify_dimension(profund)
        if idx is not None:
            dim_cols = [2, 6, 10]
            cell = _find_cell_by_col(cells_r20, dim_cols[idx])
            if cell:
                _set_checkbox(cell, True)

    # --- Row 22: Relevo (no direct mapping in DB, leave blank) ---

    # --- Row 23: Tipo de Solo ---
    cells_r23 = _get_unique_cells(rows[23])
    tipo_solo = normalize_text(erosion.get("tipoSolo")).lower()
    if tipo_solo in TIPO_SOLO_MAP:
        cell = _find_cell_by_col(cells_r23, TIPO_SOLO_MAP[tipo_solo])
        if cell:
            _set_checkbox(cell, True)

    # --- Row 24: Usos do Solo ---
    cells_r24 = _get_unique_cells(rows[24])
    usos = safe_list(erosion.get("usosSolo"))
    for uso in usos:
        uso_key = normalize_text(uso).lower()
        if uso_key in USOS_SOLO_MAP:
            cell = _find_cell_by_col(cells_r24, USOS_SOLO_MAP[uso_key])
            if cell:
                _set_checkbox(cell, True)

    # --- Row 26: Obstaculos ---
    cells_r26 = _get_unique_cells(rows[26])
    for uso in usos:
        uso_key = normalize_text(uso).lower()
        if uso_key in OBSTACULOS_MAP:
            cell = _find_cell_by_col(cells_r26, OBSTACULOS_MAP[uso_key])
            if cell:
                _set_checkbox(cell, True)

    # --- Row 27: Outros ---
    cells_r27 = _get_unique_cells(rows[27])
    cell_outros = _find_cell_by_col(cells_r27, 0)
    if cell_outros:
        outros = normalize_text(erosion.get("usoSoloOutro"))
        if outros:
            _append_text_to_label(cell_outros, "Outros:", outros)

    # --- Row 28: Medida Preventiva ---
    cells_r28 = _get_unique_cells(rows[28])
    cell_medida = _find_cell_by_col(cells_r28, 0)
    if cell_medida:
        medida = normalize_text(erosion.get("medidaPreventiva"))
        if not medida and crit_code in CRITICIDADE_SOLUCOES:
            medida = CRITICIDADE_SOLUCOES[crit_code]
        if medida:
            _append_text_to_label(cell_medida, "MEDIDA PREVENTIVA:", medida)

    # --- Row 29: Observacoes/Croquis ---
    cells_r29 = _get_unique_cells(rows[29])
    cell_obs = _find_cell_by_col(cells_r29, 0)
    if cell_obs:
        obs = normalize_text(erosion.get("dimensionamento"))
        if obs:
            _append_text_to_label(cell_obs, "OBSERVA\u00c7\u00d5ES - CROQUIS:", obs)


def _clone_table(table):
    """Deep-copy a table XML element to produce an independent clone."""
    return deepcopy(table._tbl)


def _add_page_break(document):
    """Insert a page break paragraph into the document body."""
    p = OxmlElement("w:p")
    pPr = OxmlElement("w:pPr")
    p.append(pPr)
    r = OxmlElement("w:r")
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    r.append(br)
    p.append(r)
    document._element.body.append(p)


def render_ficha_cadastro_docx(context, output_path, image_loader):
    """Render ficha de cadastro de erosao DOCX from database context."""
    render_model = ensure_dict(context.get("renderModel"))
    ficha_data = ensure_dict(render_model.get("fichaCadastro"))
    erosions = safe_list(ficha_data.get("erosions"))
    project = ensure_dict(context.get("project"))
    project_name = normalize_text(project.get("nome")) or normalize_text(project.get("id"))

    if not erosions:
        # Create empty document with a message
        doc = Document()
        doc.add_paragraph("Nenhuma erosao encontrada para gerar fichas de cadastro.")
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        doc.save(output_path)
        return output_path

    # Load template
    template_doc = Document(FICHA_TEMPLATE_PATH)
    template_table = template_doc.tables[0]

    # Create output document preserving template page setup
    output_doc = Document(FICHA_TEMPLATE_PATH)
    body = output_doc._element.body

    # Remove original table from output
    original_tbl = output_doc.tables[0]._tbl
    body.remove(original_tbl)

    # Also remove any existing paragraphs (template may have a blank one)
    for p in list(body.findall(qn("w:p"))):
        # Keep sectPr paragraphs
        pPr = p.find(qn("w:pPr"))
        if pPr is not None and pPr.find(qn("w:sectPr")) is not None:
            continue
        body.remove(p)

    # For each erosion, clone the template table, fill it, and append
    for idx, erosion in enumerate(erosions):
        if idx > 0:
            _add_page_break(output_doc)

        # Clone template table
        cloned_tbl = _clone_table(template_table)

        # Insert cloned table before the sectPr (if exists)
        sect_pr = body.find(qn("w:sectPr"))
        if sect_pr is not None:
            sect_pr.addprevious(cloned_tbl)
        else:
            body.append(cloned_tbl)

        # Access the cloned table via python-docx Table wrapper
        from docx.table import Table
        cloned_table = Table(cloned_tbl, output_doc)

        # Fill the cloned table
        _fill_ficha_table(cloned_table, erosion, project_name)

    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    output_doc.save(output_path)
    return output_path
