"""Renderer do PAEC: preenche o template DOCX tokenizado com a ficha da usina.

O template (gerado por worker/tools/paec_tokenizer.py e registrado como media
asset) tem cada campo por-usina num run unico ``{{chave}}`` ou
``{{chave|transform}}``. Aqui a mutacao e minima — so o texto dos runs muda,
preservando a formatacao institucional (Verdana, cabecalhos, tabelas, TOC).

Politica de dados ausentes: o renderer NUNCA falha por falta de valor — o
placeholder vira ``[[PENDENTE: <label>]]`` com realce amarelo (auditavel no
proprio documento) e entra na lista de pendencias devolvida no resultMeta do
job. Falha dura so por template inacessivel/corrompido.
"""

import copy
import io
import re
from collections import OrderedDict

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph as DocxParagraph

from worker.docx_renderer import _picture_size_kwargs
from worker.docx_runs import (
    collect_all_spans,
    iter_parts,
    paragraph_text,
    replace_span_text,
    run_text,
    set_run_highlight,
    set_run_text,
)

PLACEHOLDER_RE = re.compile(r"\{\{([a-z0-9_.]+)(?:\|(upper|title))?\}\}")

PENDING_PREFIX = "[[PENDENTE: "
PENDING_SUFFIX = "]]"


def _apply_transform(value, transform):
    if transform == "upper":
        return value.upper()
    if transform == "title":
        return value.title()
    return value


def _normalize_values(values):
    normalized = {}
    if isinstance(values, dict):
        for key, value in values.items():
            text = "" if value is None else str(value)
            if text.strip():
                normalized[str(key)] = text
    return normalized


def _field_labels(fields):
    labels = {}
    if isinstance(fields, list):
        for field in fields:
            if isinstance(field, dict) and field.get("key"):
                labels[str(field["key"])] = str(field.get("label") or field["key"])
    return labels


_TABLE_SEARCH_WINDOW = 8


def _enclosing_table(paragraph):
    """``w:tbl`` ancestral mais proximo de ``paragraph``, ou None se o
    paragrafo nao esta dentro de tabela nenhuma."""
    node = paragraph.getparent()
    while node is not None:
        if node.tag == qn("w:tbl"):
            return node
        node = node.getparent()
    return None


def _find_following_table(paragraph):
    """Primeira ``w:tbl`` depois de ``paragraph`` no mesmo pai — pula
    paragrafos de legenda intermediarios (ex. a frase de introducao antes da
    tabela de extintores), mas nao cruza mais que ``_TABLE_SEARCH_WINDOW``
    elementos irmaos para nao vazar pra uma tabela de outra secao."""
    parent = paragraph.getparent()
    if parent is None:
        return None
    siblings = list(parent)
    try:
        idx = siblings.index(paragraph)
    except ValueError:
        return None
    for sibling in siblings[idx + 1 : idx + 1 + _TABLE_SEARCH_WINDOW]:
        if sibling.tag == qn("w:tbl"):
            return sibling
    return None


def _find_block_table(paragraph):
    """Tabela do bloco a partir do paragrafo-ancora: se a propria ancora ja
    esta dentro de uma tabela (caso do cabecalho de coluna vir realcado, ex.
    'Nº PONTO DE ENCONTRO'), essa e a tabela; senao procura a proxima tabela
    depois do paragrafo (caso do heading solto antes da tabela, ex. Anexo IV
    ou 12.1.9)."""
    enclosing = _enclosing_table(paragraph)
    return enclosing if enclosing is not None else _find_following_table(paragraph)


def _find_block_anchor_span(document, block):
    """Localiza o span kind=list/manual ainda realcado (nao tokenizado nesta
    fase) que corresponde ao bloco, casando pelo texto salvo em
    ``anchorContext`` na curadoria do mapping.yaml."""
    anchor = (block.get("anchorContext") or "").strip()
    if not anchor:
        return None
    for span in collect_all_spans(document):
        if span.text.strip() == anchor:
            return span
    return None


def _normalize_header_cell(text):
    return re.sub(r"\s+", " ", text).strip()


def _table_header_cells(table):
    rows = table.findall(qn("w:tr"))
    if not rows:
        return []
    return [
        _normalize_header_cell(" ".join(paragraph_text(p) for p in tc.findall(qn("w:p"))))
        for tc in rows[0].findall(qn("w:tc"))
    ]


def _find_table_by_header(document, header_match):
    """Localiza a tabela pelo texto do cabecalho (linha 0), pra blocos sem
    span nenhum marcado no documento (ex. contatos internos/externos no REV
    10 — a tabela tem dado real mas nunca foi realcada na marcacao
    original). Casa por igualdade exata apos normalizar espacos internos
    (celulas de cabecalho as vezes tem varios paragrafos/quebras de linha)."""
    expected = [_normalize_header_cell(h) for h in header_match or []]
    if not expected:
        return None
    for _part, root in iter_parts(document):
        for tbl in root.iter(qn("w:tbl")):
            if _table_header_cells(tbl) == expected:
                return tbl
    return None


def _set_cell_text(tc, text, highlight=None):
    """Substitui o texto de uma celula (``w:tc``), mantendo o primeiro
    paragrafo/run (herda formatacao) e removendo os demais. ``highlight``
    marca a celula (ex. amarelo em ``[[PENDENTE]]``), igual ao tratamento
    de campo escalar sem valor."""
    paragraphs = tc.findall(qn("w:p"))
    if not paragraphs:
        return
    first_p = paragraphs[0]
    for extra_p in paragraphs[1:]:
        tc.remove(extra_p)
    runs = first_p.findall(qn("w:r"))
    if runs:
        run = runs[0]
        set_run_text(run, text)
        for extra_r in runs[1:]:
            first_p.remove(extra_r)
    else:
        run = first_p.makeelement(qn("w:r"), {})
        first_p.append(run)
        set_run_text(run, text)
    if highlight:
        set_run_highlight(run, highlight)


def _render_list_block(document, block, items):
    """Clona a linha-modelo da tabela do bloco uma vez por item da ficha
    (``paec_plant_list_items``), preenchendo as celulas por ``columns[].key``
    na ordem. Sem item nenhum, uma unica linha ``[[PENDENTE]]`` substitui a
    linha-exemplo do modelo (nunca deixa dado de outra usina no documento
    gerado). Retorna True se conseguiu renderizar a tabela.

    Duas estrategias de localizacao, conforme o bloco: ``anchorContext``
    (span ainda realcado no template — extintores, brigadistas etc.) ou
    ``headerMatch`` (tabela sem span nenhum marcado — contatos internos e
    afins; so o texto do cabecalho identifica a tabela).
    """
    columns = block.get("columns") or []
    if not columns:
        return False

    anchor_span = _find_block_anchor_span(document, block)
    if anchor_span is not None:
        table = _find_block_table(anchor_span.paragraph)
    else:
        table = _find_table_by_header(document, block.get("headerMatch"))
    if table is None:
        return False

    rows = table.findall(qn("w:tr"))
    if len(rows) < 2:
        return False
    header_row, template_row = rows[0], rows[1]
    data_rows = rows[1:]

    new_rows = []
    if items:
        for item in items:
            row = copy.deepcopy(template_row)
            cells = row.findall(qn("w:tc"))
            for cell, column in zip(cells, columns):
                value = item.get(column["key"]) if isinstance(item, dict) else None
                _set_cell_text(cell, "" if value is None else str(value))
            new_rows.append(row)
    else:
        row = copy.deepcopy(template_row)
        cells = row.findall(qn("w:tc"))
        if cells:
            label = block.get("label") or block.get("key")
            _set_cell_text(cells[0], f"{PENDING_PREFIX}{label}{PENDING_SUFFIX}", highlight="yellow")
            for cell in cells[1:]:
                _set_cell_text(cell, "")
        new_rows.append(row)

    for row in data_rows:
        table.remove(row)
    anchor = header_row
    for row in new_rows:
        anchor.addnext(row)
        anchor = row

    if anchor_span is not None:
        replace_span_text(anchor_span, anchor_span.text, drop_highlight=True)
    return True


def _render_list_blocks(document, blocks, list_items_map):
    """Renderiza todos os blocos ``kind=list`` do manifest; retorna as chaves
    que NAO deu pra renderizar (sem tabela localizavel ou sem colunas
    curadas ainda) — ficam com a linha-exemplo original do modelo, visivel
    pra auditoria manual."""
    unrendered = []
    for block in blocks or []:
        if not isinstance(block, dict) or block.get("kind") != "list":
            continue
        items = list_items_map.get(block.get("key")) or []
        if not _render_list_block(document, block, items):
            unrendered.append(block.get("key"))
    return unrendered


def _find_section_heading_span(document, section):
    """Localiza o span kind=section_title ainda realcado (apply_mapping nunca
    tokeniza esse kind) que corresponde a secao, casando pelo texto salvo em
    ``defaultTitle`` — igual ``_find_block_anchor_span`` faz pros blocos."""
    title = (section.get("defaultTitle") or "").strip()
    if not title:
        return None
    for span in collect_all_spans(document):
        if span.text.strip() == title:
            return span
    return None


def _remove_section_range(paragraph, next_paragraph):
    """Remove ``paragraph`` (o heading) e todo elemento irmao ate
    ``next_paragraph`` (exclusive) — o conteudo da secao termina onde comeca
    o proximo heading do documento, seja ele do mesmo grupo ou nao. Sem
    ``next_paragraph`` (ultima secao do documento), remove ate o fim do pai."""
    parent = paragraph.getparent()
    if parent is None:
        return
    siblings = list(parent)
    start_idx = siblings.index(paragraph)
    if next_paragraph is not None and next_paragraph.getparent() is parent:
        end_idx = siblings.index(next_paragraph)
    else:
        end_idx = len(siblings)
    for element in siblings[start_idx:end_idx]:
        parent.remove(element)


_HEADING_NUMBER_RE = re.compile(r"^\s*\d+(?:\.\d+)*\.?\s*")


def _renumber_title(title, new_number):
    """Troca o prefixo numerico do titulo (``12.1.3. Rede de Hidrantes`` com
    ``new_number='12.1.2'`` vira ``12.1.2. Rede de Hidrantes``); titulo sem
    prefixo numerico (ex. title_override sem numero) so ganha o prefixo."""
    rest = _HEADING_NUMBER_RE.sub("", title, count=1).strip()
    return f"{new_number}. {rest}" if rest else f"{new_number}."


def _mark_toc_dirty(document):
    """Injeta ``<w:updateFields w:val=\"true\"/>`` em settings.xml — o Word
    oferece atualizar o sumario ao abrir o documento (nao recomputamos o TOC
    na mao). Idempotente: nao duplica se ja presente."""
    settings = document.settings.element
    if settings.find(qn("w:updateFields")) is not None:
        return
    element = settings.makeelement(qn("w:updateFields"), {qn("w:val"): "true"})
    settings.insert(0, element)


def _render_section_flags(document, sections, section_flags):
    """Aplica sectionFlags (Fase 3): remove o conteudo das secoes desligadas
    (do heading ate o proximo heading do documento) e renumera
    sequencialmente as remanescentes de cada ``renumberGroup``, aplicando
    ``titleOverride`` quando presente. Marca o TOC como desatualizado se
    alguma secao mudou. Retorna o conjunto de sectionKeys processadas.
    """
    section_flags = section_flags or {}
    resolved = []
    for section in sections or []:
        if not isinstance(section, dict):
            continue
        span = _find_section_heading_span(document, section)
        if span is not None:
            resolved.append((section, span))

    changed = False
    disabled_keys = set()
    for index, (section, span) in enumerate(resolved):
        flag = section_flags.get(section.get("sectionKey")) or {}
        if flag.get("enabled") is False:
            disabled_keys.add(section["sectionKey"])
            next_paragraph = resolved[index + 1][1].paragraph if index + 1 < len(resolved) else None
            _remove_section_range(span.paragraph, next_paragraph)
            changed = True

    groups = OrderedDict()
    for section, span in resolved:
        if section["sectionKey"] in disabled_keys:
            continue
        group = section.get("renumberGroup")
        if group:
            groups.setdefault(group, []).append((section, span))

    renumbered_keys = set()
    for group, items in groups.items():
        for position, (section, span) in enumerate(items, start=1):
            flag = section_flags.get(section["sectionKey"]) or {}
            base_title = flag.get("titleOverride") or section.get("defaultTitle") or ""
            new_text = _renumber_title(base_title, f"{group}.{position}")
            if new_text != span.text:
                changed = True
            replace_span_text(span, new_text, drop_highlight=True)
            renumbered_keys.add(section["sectionKey"])

    # Demais headings resolvidos (sem grupo, ou grupo mas sem entrada em
    # sections/renumberGroup) so perdem o realce -- nunca fica marca de
    # curadoria visivel num documento final gerado.
    for section, span in resolved:
        key = section["sectionKey"]
        if key in disabled_keys or key in renumbered_keys:
            continue
        replace_span_text(span, span.text, drop_highlight=True)

    if changed:
        _mark_toc_dirty(document)

    return {section["sectionKey"] for section, _span in resolved}


_ANNEX_HEADING_RE = re.compile(r"^anexo\b", re.IGNORECASE)


def _find_image_slot_region(anchor_paragraph):
    """Paragrafos irmaos apos a ancora ate o heading do proximo anexo
    (``anexo ...``, case-insensitive) ou um elemento nao-paragrafo — a
    regiao de conteudo do slot. No REV 10 a regiao pode intercalar legendas
    de texto com as imagens (ex. o unifilar tem "Diagrama de Operação;"
    antes de cada painel), entao ela NAO para no primeiro texto."""
    parent = anchor_paragraph.getparent()
    if parent is None:
        return []
    siblings = list(parent)
    idx = siblings.index(anchor_paragraph)
    region = []
    for sibling in siblings[idx + 1 :]:
        if sibling.tag != qn("w:p"):
            break
        if _ANNEX_HEADING_RE.match(paragraph_text(sibling).strip()):
            break
        region.append(sibling)
    return region


def _insert_image_paragraph(document, after_element, buffer):
    """Insere um paragrafo novo com a imagem centralizada logo apos
    ``after_element``. Mistura lxml (posicao) com a API alta do python-docx
    (add_picture cuida do relationship/part da imagem); dimensionamento
    reusa a regra ja calibrada dos outros renderers (largura 15cm, teto de
    altura pra retrato)."""
    new_p = after_element.makeelement(qn("w:p"), {})
    after_element.addnext(new_p)
    paragraph = DocxParagraph(new_p, document)
    run = paragraph.add_run()
    run.add_picture(io.BytesIO(buffer), **_picture_size_kwargs(buffer))
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return new_p


def _insert_pending_paragraph(after_element, label):
    """Paragrafo de texto ``[[PENDENTE: <label>]]`` realcado de amarelo no
    lugar das imagens de um slot vazio — nao da pra 'realcar' uma imagem que
    nao existe, entao a pendencia vira texto visivel."""
    new_p = after_element.makeelement(qn("w:p"), {})
    after_element.addnext(new_p)
    run = new_p.makeelement(qn("w:r"), {})
    new_p.append(run)
    set_run_text(run, f"{PENDING_PREFIX}{label}{PENDING_SUFFIX}")
    set_run_highlight(run, "yellow")
    return new_p


def _render_image_slots(document, image_slots, assets, image_loader):
    """Renderiza os slots de imagem (Fase 4 — rota de fuga, unifilar):
    localiza a ancora do slot (span kind=image, ainda realcado no template),
    REMOVE as imagens-exemplo do modelo (nunca deixa diagrama de outra usina
    no documento gerado — mesmo principio dos blocos tabulares) e insere uma
    imagem da ficha por mediaAssetId, na ordem. Slot sem imagem nenhuma vira
    um paragrafo ``[[PENDENTE]]``. Retorna as assetKeys que NAO deu pra
    renderizar (ancora nao encontrada) — ficam com as imagens-exemplo
    originais, visiveis pra auditoria manual."""
    unrendered = []
    for slot in image_slots or []:
        if not isinstance(slot, dict):
            continue
        anchor_span = _find_block_anchor_span(document, slot)
        if anchor_span is None:
            unrendered.append(slot.get("assetKey"))
            continue

        anchor_paragraph = anchor_span.paragraph
        for element in _find_image_slot_region(anchor_paragraph):
            if list(element.iter(qn("w:drawing"))):
                element.getparent().remove(element)

        media_ids = assets.get(slot.get("assetKey")) if isinstance(assets, dict) else None
        media_ids = [m for m in (media_ids or []) if m]

        insert_after = anchor_paragraph
        if media_ids and image_loader is not None:
            for media_id in media_ids:
                try:
                    media = image_loader(media_id)
                    buffer = media.get("buffer") if isinstance(media, dict) else None
                    if not buffer:
                        raise ValueError("conteudo vazio")
                    insert_after = _insert_image_paragraph(document, insert_after, buffer)
                except Exception:
                    insert_after = _insert_pending_paragraph(
                        insert_after, f"Imagem indisponivel — {slot.get('label') or media_id}"
                    )
        else:
            _insert_pending_paragraph(anchor_paragraph, slot.get("label") or slot.get("assetKey"))

        replace_span_text(anchor_span, anchor_span.text, drop_highlight=True)
    return unrendered


def render_paec_to_docx(context, template_bytes, output_path, image_loader=None):
    """Renderiza o PAEC e retorna ``{"pendencies": [...], "stats": {...}}``.

    ``context`` e o retorno de buildPaecContext (backend): renderModel.paecReport
    com fields (catalogo do manifest), values (ficha) e pendencies pre-computadas
    (blocos manuais etc. — repassadas e complementadas aqui). ``image_loader``
    (mediaAssetId -> {"buffer": bytes}) vem do cache de prefetch do
    job_processor; sem ele os imageSlots viram só o placeholder de pendencia.
    """
    render_model = context.get("renderModel") if isinstance(context, dict) else {}
    paec = render_model.get("paecReport") if isinstance(render_model, dict) else {}

    values = _normalize_values(paec.get("values"))
    labels = _field_labels(paec.get("fields"))

    document = Document(io.BytesIO(template_bytes))

    missing_keys = []
    unresolved_tokens = []

    for _part, root in iter_parts(document):
        for run in root.iter(qn("w:r")):
            text = run_text(run)
            if "{{" not in text:
                continue

            replaced_missing = []

            def _sub(match):
                key, transform = match.group(1), match.group(2)
                if key in values:
                    return _apply_transform(values[key], transform or "none")
                replaced_missing.append(key)
                label = labels.get(key, key)
                return f"{PENDING_PREFIX}{label}{PENDING_SUFFIX}"

            new_text = PLACEHOLDER_RE.sub(_sub, text)
            if new_text != text:
                set_run_text(run, new_text)
                if replaced_missing:
                    set_run_highlight(run, "yellow")
                    missing_keys.extend(replaced_missing)
            if "{{" in new_text:
                # Token que o regex nao reconhece (manifest dessincronizado do
                # template): mantem visivel e reporta.
                unresolved_tokens.append(new_text.strip()[:120])
                set_run_highlight(run, "yellow")

    list_items_map = paec.get("listItems") if isinstance(paec.get("listItems"), dict) else {}
    _render_list_blocks(document, paec.get("blocks"), list_items_map)

    _render_section_flags(document, paec.get("sections"), paec.get("sectionFlags"))

    assets_map = paec.get("assets") if isinstance(paec.get("assets"), dict) else {}
    _render_image_slots(document, paec.get("imageSlots"), assets_map, image_loader)

    pendencies = []
    seen = set()
    for key in missing_keys:
        if key in seen:
            continue
        seen.add(key)
        pendencies.append({
            "kind": "field",
            "key": key,
            "label": labels.get(key, key),
            "section": None,
        })
    for token in unresolved_tokens:
        pendencies.append({
            "kind": "unresolved_token",
            "key": token,
            "label": f"Marcador nao resolvido: {token}",
            "section": None,
        })
    # Pendencias que so o backend conhece (blocos manuais da fase 1); campos ja
    # sao recalculados acima com base no documento real.
    for pendency in paec.get("pendencies") or []:
        if isinstance(pendency, dict) and pendency.get("kind") in {"manual_block", "list", "image"}:
            pendencies.append(pendency)

    document.save(output_path)

    stats = paec.get("stats") if isinstance(paec.get("stats"), dict) else {}
    return {"pendencies": pendencies, "stats": stats}
