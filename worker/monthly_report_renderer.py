"""Renderer DOCX do Relatorio Mensal de Acompanhamento dos Servicos.

Reproduz o relatorio real entregue (modelo contratual APENSO D): capa com a
arte institucional CONCREMAT de pagina inteira, sumario (TOC 1-3), corpo
Verdana com headings pretos, cabecalho interno com logos AXIA + CONCREMAT|CCCC,
rodape "Classificacao: Interna" e, por engenheiro, o quadro semanal de
atividades (3 variantes) + legenda + resumo por projeto.

A estrutura de blocos espelha apps/relat/src/features/monthly-report/utils/
docModel.js — a fixture worker/tests/fixtures/monthly_report_render_model.json
trava a paridade entre a preview React e este renderer.
"""

import datetime
import os

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor, Twips

ASSETS_DIR = os.path.join(os.path.dirname(__file__), "assets")
COVER_IMAGE = os.path.join(ASSETS_DIR, "monthly_report_cover.png")
LOGO_AXIA = os.path.join(ASSETS_DIR, "monthly_report_logo_axia.png")
LOGO_CONCREMAT = os.path.join(ASSETS_DIR, "monthly_report_logo_concremat_cccc.png")

BODY_FONT = "Verdana"  # APENSO D: estilo Normal do modelo contratual

MONTHS = [
    "Janeiro", "Fevereiro", "Março", "Abril", "Maio", "Junho",
    "Julho", "Agosto", "Setembro", "Outubro", "Novembro", "Dezembro",
]

# Cores canonicas das categorias (mesmos hex de utils/constants.js).
CATEGORIES = {
    "vistoria": {"label": "Vistoria de campo", "color": "166534", "light": "ECFDF3"},
    "doc": {"label": "Documentação técnica", "color": "2563EB", "light": "EFF6FF"},
    "relatorio": {"label": "Elaboração de relatório", "color": "0369A1", "light": "E4F1F9"},
    "geo": {"label": "Geoprocessamento", "color": "B45309", "light": "FFFBEB"},
    "reuniao": {"label": "Reuniões / articulação", "color": "475569", "light": "F1F5F9"},
    "outro": {"label": "Outros", "color": "7F1D1D", "light": "FEF2F2"},
}

QUADRO_FILL_OUT = "F1F3F6"
QUADRO_FILL_HOLIDAY = "FDF0F4"
QUADRO_FILL_WEEKEND = "F8FAFC"
QUADRO_BORDER_COLOR = "C9CFD8"
HOLIDAY_COLOR = "993556"
FOOTER_COLOR = "2E74B5"

SECTION2_LEAD = "As atividades são realizadas conforme solicitação da contratante a cada um dos engenheiros, conforme demonstrado a seguir."
ENG_LEAD = "As atividades realizadas no período deste relatório foram distribuídas na seguinte disposição:"
RESUMO_DESC = "Descrição das atividades desenvolvidas em cada empreendimento ao longo do período, com destaque para entregas e marcos relevantes."
EMPTY_INTRO = "Introdução ainda não escrita — use o texto-modelo na etapa 1."
EMPTY_CONCLUSAO = "Conclusão ainda não escrita — use o texto-modelo na etapa 3."


def _norm(value):
    return str(value or "").strip()


# ----------------------------------------------------------------------------
# Datas / feriados (porta de utils/holidays.js + calendar.js)
# ----------------------------------------------------------------------------
def date_key(d):
    return f"{d.year:04d}-{d.month:02d}-{d.day:02d}"


def parse_date_key(k):
    y, m, d = (int(x) for x in str(k).split("-"))
    return datetime.date(y, m, d)


def get_date_range(ref_year, ref_month):
    # ref_month 0-11. Periodo: dia 16 do mes anterior -> dia 15 do mes de ref.
    end = datetime.date(ref_year, ref_month + 1, 15)
    start_month = ref_month
    start_year = ref_year
    if start_month == 0:
        start_month = 12
        start_year -= 1
    start = datetime.date(start_year, start_month, 16)
    return start, end


def build_holiday_map(holidays):
    """Feriados sao lista explicita do relatorio (sem computo de oficiais)."""
    result = {}
    for h in holidays or []:
        if isinstance(h, dict) and _norm(h.get("date")):
            result[_norm(h["date"])] = _norm(h.get("name")) or "Feriado"
    return result


def is_working_day(d, holiday_keys):
    if d.weekday() >= 5:  # 5=sab, 6=dom
        return False
    return date_key(d) not in holiday_keys


def activities_visible_on_date(activities, date_str, holiday_keys):
    """Atividade de 1 dia aparece sempre; multi-dia so em dias uteis."""
    d = parse_date_key(date_str)
    visible = []
    for a in activities:
        if date_str < a["startDate"] or date_str > a["endDate"]:
            continue
        if a["startDate"] == a["endDate"]:
            visible.append(a)
        elif is_working_day(d, holiday_keys):
            visible.append(a)
    return visible


def build_quadro_weeks(report, engineer_activities):
    """Matriz semanal do quadro — mesma estrutura de buildQuadroWeeks (docModel.js)."""
    start, end = get_date_range(int(report["refYear"]), int(report["refMonth"]))
    hmap = build_holiday_map(report.get("holidays"))
    holiday_keys = set(hmap.keys())

    def js_get_day(d):
        return (d.weekday() + 1) % 7  # domingo=0

    grid_start = start - datetime.timedelta(days=js_get_day(start))
    grid_end = end + datetime.timedelta(days=6 - js_get_day(end))

    weeks = []
    week = []
    cursor = grid_start
    while cursor <= grid_end:
        key = date_key(cursor)
        in_range = start <= cursor <= end
        visible = activities_visible_on_date(engineer_activities, key, holiday_keys) if in_range else []
        week.append({
            "dateKey": key,
            "dayNum": cursor.day,
            "inRange": in_range,
            "isWeekend": cursor.weekday() >= 5,
            "holidayName": hmap.get(key) if in_range else None,
            "activities": [
                {
                    "category": a.get("category"),
                    "description": _norm(a.get("description")),
                    "isFirstDay": key == a["startDate"],
                }
                for a in visible
            ],
        })
        if len(week) == 7:
            weeks.append(week)
            week = []
        cursor += datetime.timedelta(days=1)
    return weeks


# ----------------------------------------------------------------------------
# Helpers XML python-docx
# ----------------------------------------------------------------------------
def _run(paragraph, text, size=None, bold=False, italic=False, color=None, font=None):
    r = paragraph.add_run(text)
    r.font.name = font or BODY_FONT
    if size is not None:
        r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    if color:
        r.font.color.rgb = RGBColor.from_string(color)
    return r


def _set_cell_background(cell, hex_color):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    cell._tc.get_or_add_tcPr().append(shd)


def _set_paragraph_shading(paragraph, fill):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    paragraph._p.get_or_add_pPr().append(shd)


def _set_paragraph_left_border(paragraph, color, size=18):
    p_bdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), str(size))
    left.set(qn("w:space"), "2")
    left.set(qn("w:color"), color)
    p_bdr.append(left)
    paragraph._p.get_or_add_pPr().append(p_bdr)


def _set_table_borders(table, color, size=4):
    tbl_pr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(size))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        borders.append(el)
    tbl_pr.append(borders)


def _set_row_cant_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant = OxmlElement("w:cantSplit")
    tr_pr.append(cant)


def _float_picture_full_page(run, page_width, page_height):
    """Converte a imagem inline do run em ancora atras do texto, cobrindo a pagina."""
    drawing = run._r.find(qn("w:drawing"))
    inline = drawing.find(qn("wp:inline"))

    anchor = OxmlElement("wp:anchor")
    for attr, val in (
        ("distT", "0"), ("distB", "0"), ("distL", "0"), ("distR", "0"),
        ("simplePos", "0"), ("relativeHeight", "0"), ("behindDoc", "1"),
        ("locked", "0"), ("layoutInCell", "1"), ("allowOverlap", "1"),
    ):
        anchor.set(attr, val)

    simple = OxmlElement("wp:simplePos")
    simple.set("x", "0")
    simple.set("y", "0")
    anchor.append(simple)

    for tag in ("wp:positionH", "wp:positionV"):
        pos = OxmlElement(tag)
        pos.set("relativeFrom", "page")
        offset = OxmlElement("wp:posOffset")
        offset.text = "0"
        pos.append(offset)
        anchor.append(pos)

    extent = inline.find(qn("wp:extent"))
    extent.set("cx", str(int(page_width)))
    extent.set("cy", str(int(page_height)))
    anchor.append(extent)

    effect = inline.find(qn("wp:effectExtent"))
    if effect is not None:
        anchor.append(effect)
    anchor.append(OxmlElement("wp:wrapNone"))

    anchor.append(inline.find(qn("wp:docPr")))
    cnv = inline.find(qn("wp:cNvGraphicFramePr"))
    if cnv is not None:
        anchor.append(cnv)

    graphic = inline.find(qn("a:graphic"))
    for ext in graphic.iter(qn("a:ext")):
        ext.set("cx", str(int(page_width)))
        ext.set("cy", str(int(page_height)))
    anchor.append(graphic)

    drawing.remove(inline)
    drawing.append(anchor)


def _enable_update_fields(document):
    """Word atualiza o TOC (numeros de pagina) ao abrir o documento."""
    settings = document.settings.element
    el = settings.find(qn("w:updateFields"))
    if el is None:
        el = OxmlElement("w:updateFields")
        settings.append(el)
    el.set(qn("w:val"), "true")


def _add_toc_field(paragraph):
    r = paragraph.add_run()._r
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "Sumário — atualize os campos ao abrir no Word."
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for el in (begin, instr, separate, placeholder, end):
        r.append(el)


def _setup_styles(document):
    """Normal Verdana 10pt + Headings 1-3 pretos (TOC usa os outline levels)."""
    normal = document.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(10)
    normal.font.color.rgb = RGBColor.from_string("1A1A1A")

    for name, size in (("Heading 1", 12), ("Heading 2", 11), ("Heading 3", 10.5)):
        style = document.styles[name]
        style.font.name = BODY_FONT
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string("000000")
        style.paragraph_format.space_before = Pt(14)
        style.paragraph_format.space_after = Pt(8)


def _body_paragraph(document, text, indent=True):
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.15
    if indent:
        p.paragraph_format.first_line_indent = Cm(1.25)
    _run(p, text, size=10)
    return p


def _free_text(document, text, empty_placeholder):
    paragraphs = [t.strip() for t in str(text or "").split("\n") if t.strip()]
    if not paragraphs:
        p = document.add_paragraph()
        p.paragraph_format.space_after = Pt(8)
        _run(p, empty_placeholder, size=10, italic=True, color="9AA4B2")
        return
    for t in paragraphs:
        _body_paragraph(document, t)


# ----------------------------------------------------------------------------
# Capa, cabecalho e rodape
# ----------------------------------------------------------------------------
def _build_cover(document, month_label, period_label):
    section = document.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    for attr in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(section, attr, Cm(0))

    art = document.paragraphs[0] if document.paragraphs else document.add_paragraph()
    run = art.add_run()
    run.add_picture(COVER_IMAGE, width=section.page_width, height=section.page_height)
    _float_picture_full_page(run, section.page_width, section.page_height)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Cm(3.6)
    title.paragraph_format.space_after = Pt(0)
    _run(title, "Relatório Mensal de Acompanhamento dos Serviços", size=16, bold=True, color="FFFFFF")

    month = document.add_paragraph()
    month.alignment = WD_ALIGN_PARAGRAPH.CENTER
    month.paragraph_format.space_before = Cm(11.8)
    month.paragraph_format.space_after = Pt(4)
    _run(month, month_label, size=13, bold=True, color="000000")

    period = document.add_paragraph()
    period.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(period, period_label, size=9, color="595959")


def _build_content_header_footer(section):
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False

    header = section.header
    table = header.add_table(rows=1, cols=2, width=Cm(17.0))
    left, right = table.rows[0].cells
    lp = left.paragraphs[0]
    lp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    lp.add_run().add_picture(LOGO_AXIA, width=Cm(3.4))
    rp = right.paragraphs[0]
    rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rp.add_run().add_picture(LOGO_CONCREMAT, width=Cm(5.2))
    # Limpa o paragrafo vazio default do header (fica antes da tabela).
    if header.paragraphs and not header.paragraphs[0].text and len(header.paragraphs) > 1:
        header.paragraphs[0]._p.getparent().remove(header.paragraphs[0]._p)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(fp, "Classificação: Interna", size=9, color=FOOTER_COLOR)


# ----------------------------------------------------------------------------
# Quadro semanal (3 variantes) + legenda
# ----------------------------------------------------------------------------
def _activity_paragraph(cell, activity, variant):
    cat = CATEGORIES.get(activity["category"], {"color": "64748B", "light": "ECEFF3"})
    first = activity["isFirstDay"]
    text = activity["description"]
    p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.line_spacing = 1.0

    if variant == "preenchido":
        if first:
            _set_paragraph_shading(p, cat["color"])
            _run(p, f" {text} ", size=6.5, bold=True, color="FFFFFF")
        else:
            _set_paragraph_shading(p, cat["light"])
            _run(p, f" ↳ {text} ", size=6.5, color=cat["color"])
        return p

    if variant == "barra":
        _set_paragraph_shading(p, cat["light"])
        _set_paragraph_left_border(p, cat["color"])
        _run(p, f" {'' if first else '↳ '}{text}", size=6.5, color="1E293B")
        return p

    # marcador — estilo do relatorio real entregue.
    _run(p, "● " if first else "↳ ", size=6.5, bold=True, color=cat["color"])
    _run(p, text, size=6.5, color="333333")
    return p


def _build_quadro_table(document, weeks, variant):
    weekdays = ["DOM", "SEG", "TER", "QUA", "QUI", "SEX", "SÁB"]
    table = document.add_table(rows=1 + len(weeks), cols=7)
    table.autofit = False
    _set_table_borders(table, QUADRO_BORDER_COLOR)
    col_width = Twips(1377)

    for i, wd in enumerate(weekdays):
        cell = table.rows[0].cells[i]
        cell.width = col_width
        _set_cell_background(cell, "EEF3FB")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        _run(p, wd, size=7, bold=True, color="64748B")

    for wi, week in enumerate(weeks):
        row = table.rows[1 + wi]
        row.height = Pt(50)
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        _set_row_cant_split(row)
        for ci, cell_data in enumerate(week):
            cell = row.cells[ci]
            cell.width = col_width
            if not cell_data["inRange"]:
                _set_cell_background(cell, QUADRO_FILL_OUT)
            elif cell_data["holidayName"]:
                _set_cell_background(cell, QUADRO_FILL_HOLIDAY)
            elif cell_data["isWeekend"]:
                _set_cell_background(cell, QUADRO_FILL_WEEKEND)

            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(1)
            _run(p, str(cell_data["dayNum"]), size=7.5, bold=True,
                 color="334155" if cell_data["inRange"] else "A8B0BC")

            if cell_data["holidayName"]:
                hp = cell.add_paragraph()
                hp.paragraph_format.space_after = Pt(1)
                hp.paragraph_format.line_spacing = 1.0
                _run(hp, f"★ {cell_data['holidayName']}", size=6, color=HOLIDAY_COLOR)

            for activity in cell_data["activities"]:
                _activity_paragraph(cell, activity, variant)

    return table


def _build_legend(document):
    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(12)
    _run(p, "LEGENDA   ", size=7, bold=True, color="64748B")
    for cat in CATEGORIES.values():
        _run(p, "● ", size=8, bold=True, color=cat["color"])
        _run(p, f"{cat['label']}   ", size=8, color="334155")
    _run(p, "★ ", size=8, bold=True, color=HOLIDAY_COLOR)
    _run(p, "Feriado", size=8, color="334155")
    return p


# ----------------------------------------------------------------------------
# Render principal
# ----------------------------------------------------------------------------
def render_context_to_docx(context, output_path):
    render_model = context.get("renderModel", {}) if isinstance(context, dict) else {}
    report = render_model.get("monthlyReport", {}) if isinstance(render_model, dict) else {}

    ref_year = int(report.get("refYear"))
    ref_month = int(report.get("refMonth"))
    quadro_style = _norm(report.get("quadroStyle")) or "marcador"
    engineers = [e for e in (report.get("engineers") or []) if isinstance(e, dict)]

    start, end = get_date_range(ref_year, ref_month)
    month_label = f"{MONTHS[ref_month].upper()} - {ref_year}"
    period_label = f"Período: {start.strftime('%d/%m/%Y')} a {end.strftime('%d/%m/%Y')}"

    document = Document()
    _setup_styles(document)
    _enable_update_fields(document)

    # --- Capa (secao 1, margens zero, arte de pagina inteira) ---
    _build_cover(document, month_label, period_label)

    # --- Miolo (secao 2: margens 2cm, logos no cabecalho, rodape) ---
    section = document.add_section(WD_SECTION.NEW_PAGE)
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    for attr in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(section, attr, Cm(2.0))
    _build_content_header_footer(section)

    # Sumario
    toc_title = document.add_paragraph()
    toc_title.paragraph_format.space_after = Pt(12)
    _run(toc_title, "SUMÁRIO:", size=11, bold=True)
    _add_toc_field(document.add_paragraph())
    document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    # 1 INTRODUCAO
    document.add_paragraph("1 INTRODUÇÃO", style="Heading 1")
    _free_text(document, report.get("intro"), EMPTY_INTRO)

    # 2 ATIVIDADES
    document.add_paragraph("2 ATIVIDADES REALIZADAS NO PERÍODO", style="Heading 1")
    _body_paragraph(document, SECTION2_LEAD)

    for i, engineer in enumerate(engineers):
        n = i + 1
        name = _norm(engineer.get("name")) or f"Engenheiro {n}"
        document.add_paragraph(f"2.{n} Atividades que o eng. {name} realizou", style="Heading 2")
        _body_paragraph(document, ENG_LEAD)

        weeks = build_quadro_weeks(report, [a for a in (engineer.get("activities") or []) if isinstance(a, dict)])
        _build_quadro_table(document, weeks, quadro_style)
        _build_legend(document)

        document.add_paragraph(f"2.{n}.1 Resumo por projeto:", style="Heading 3")
        _body_paragraph(document, RESUMO_DESC)
        for project in engineer.get("projects") or []:
            if not isinstance(project, dict):
                continue
            p_name = _norm(project.get("name"))
            p_text = " ".join(str(project.get("description") or "").split())
            if not p_name and not p_text:
                continue
            bullet = document.add_paragraph(style="List Bullet")
            bullet.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            bullet.paragraph_format.space_after = Pt(8)
            bullet.paragraph_format.line_spacing = 1.15
            _run(bullet, f"{p_name or 'Empreendimento'}: ", size=10, bold=True)
            _run(bullet, p_text, size=10)

    # 3 CONCLUSAO
    document.add_paragraph("3 CONCLUSÃO", style="Heading 1")
    _free_text(document, report.get("conclusao"), EMPTY_CONCLUSAO)

    document.save(output_path)
    return output_path
