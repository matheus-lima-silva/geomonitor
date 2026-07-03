"""Tokenizer do modelo canonico marcado do PAEC.

Transforma o DOCX marcado (realce amarelo = campo por usina, realce vermelho =
bloco tabular por usina) num template reutilizavel com placeholders
``{{chave}}`` + um manifest JSON, mediado por um arquivo de curadoria humana
(``mapping.yaml``). Roda uma vez por revisao do modelo — a geracao por usina
(worker) so consome o template tokenizado.

Subcomandos:

  extract   MARCADO.docx -o mapping.yaml
            Localiza os spans realcados (mesclando runs fragmentados pelo
            Word), sugere chave/tipo/transform por heuristica e emite o
            mapping.yaml para curadoria. Inclui auditoria de paginas.

  apply     MARCADO.docx -m mapping.yaml -o tokenized.docx --manifest manifest.json --revision "REV 10"
            Reaplica o mapping sobre o DOCX marcado: cada span kind=field vira
            ``{{chave}}`` (ou ``{{chave|upper}}``) em run unico, preservando a
            formatacao e removendo o realce. Emite o manifest do template.

  render    TOKENIZED.docx -v values.json -o out.docx
            Preenche os placeholders com um mapa chave->valor (validacao
            local; o renderer de producao vive em worker/paec_renderer.py).

  roundtrip MARCADO.docx -m mapping.yaml
            extract -> apply -> render com os valores originais e compara o
            texto com o modelo marcado. Deve ser identico.

Uso: python -m worker.tools.paec_tokenizer <subcomando> ...
"""

import argparse
import io
import json
import re
import sys
import unicodedata
from collections import OrderedDict

from docx import Document
from docx.oxml.ns import qn

from worker.docx_runs import (
    collect_all_spans,
    collect_page_geometry,
    iter_parts,
    iter_paragraphs,
    paragraph_text,
    replace_span_text,
    run_text,
    set_run_text,
)

PLACEHOLDER_RE = re.compile(r"\{\{([a-z0-9_.]+)(?:\|(upper|title))?\}\}")

SPAN_KINDS = ("field", "list", "image", "manual", "section_title", "whitespace", "ignore")

_HEADING_RE = re.compile(r"^\s*\d+(\.\d+)+\.?\s+\S")
_CNPJ_RE = re.compile(r"\d{2}\.\d{3}\.\d{3}/\d{4}-\d{2}")
_PHONE_RE = re.compile(r"\(?\d{2}\)?\s?9?\s?\d{4}[-\s]?\d{4}")
_EMAIL_RE = re.compile(r"[\w.+-]+@[\w-]+\.[\w.]+")
_URL_RE = re.compile(r"(www\.|https?://|\w+\.com(\.br)?$)")

_NBSP = "\u00a0"


# ---------------------------------------------------------------------------
# Heuristicas de sugestao
# ---------------------------------------------------------------------------
def _strip_accents(text):
    return "".join(
        c for c in unicodedata.normalize("NFKD", text) if not unicodedata.combining(c)
    )


def normalize_for_grouping(text):
    """Chave de agrupamento: ocorrencias do mesmo texto (a menos de caixa,
    espacos e acentos) apontam para o MESMO campo."""
    text = _strip_accents(text.replace(_NBSP, " ")).casefold()
    return re.sub(r"\s+", " ", text).strip()


def slugify(text, max_tokens=5):
    text = _strip_accents(text.replace(_NBSP, " ")).lower()
    tokens = re.findall(r"[a-z0-9]+", text)[:max_tokens]
    return "_".join(tokens)[:48] or "campo"


def _classify_text(text):
    """Prefixo semantico para a chave sugerida, quando reconhecivel."""
    if _CNPJ_RE.search(text):
        return "cnpj"
    if _EMAIL_RE.search(text):
        return "email"
    if _PHONE_RE.search(text):
        return "telefone"
    if _URL_RE.search(text.strip()):
        return "web"
    return None


def _looks_like_heading_label(text):
    """``_HEADING_RE`` alone also matches numeric values shaped like a
    section number (ex. ``95.814 L``, a volume). Require enough letters in
    the remainder to tell a real heading title from a quantity+unit."""
    if not _HEADING_RE.match(text):
        return False
    letters = re.findall(r"[^\W\d_]", text, flags=re.UNICODE)
    return len(letters) >= 4


def suggest_kind(span):
    text = span.text.strip()
    if not text:
        return "whitespace"
    if span.color == "red":
        return "list"
    if _looks_like_heading_label(text):
        return "section_title"
    if text.casefold().startswith("anexo"):
        return "manual"
    return "field"


def _variant_transform(variant, canonical):
    """Transform que leva ``canonical`` a ``variant``; None se nao derivavel."""
    if variant == canonical:
        return "none"
    if variant == canonical.upper():
        return "upper"
    if variant == canonical.title():
        return "title"
    return None


def build_mapping(document, source_name):
    """Extrai os spans e monta o dicionario do mapping.yaml (pre-curadoria).

    Agrupamento em duas passadas: primeiro reune as variantes de texto de cada
    campo (a menos de caixa/acentos/espacos) e escolhe como canonica uma
    variante que NAO seja toda maiuscula (se existir); depois atribui chave e
    transform por ocorrencia. Variante nao derivavel por upper/title da
    canonica ganha chave propria (_v2, _v3...) para manter o round-trip exato.
    """
    spans = collect_all_spans(document)
    prepared = [
        (span, suggest_kind(span), span.text.strip(), normalize_for_grouping(span.text))
        for span in spans
    ]

    # passa 1: variantes por grupo (apenas campos)
    group_variants = OrderedDict()
    for _span, kind, stripped, norm in prepared:
        if kind == "field":
            group_variants.setdefault(norm, []).append(stripped)

    canonicals = {
        norm: next((v for v in variants if v != v.upper()), variants[0])
        for norm, variants in group_variants.items()
    }

    # passa 2: chave por grupo + transform por ocorrencia
    class_counters = {}
    group_keys = {}
    variant_keys = {}  # (norm, variant) -> (key, transform)

    # ``slugify`` trunca em 5 tokens e descarta pontuacao: textos distintos
    # (ex. "...e Colombia" vs "...e P. Colombia", "Telefone/Fax" vs
    # "Telefone / Fax") podem colidir no mesmo slug apesar de pertencerem a
    # grupos NORMALIZADOS diferentes. ``_claim_key`` garante que cada
    # identidade (norm de campo, ou texto literal de bloco/secao/anexo)
    # recebe uma chave unica, desambiguando com sufixo numerico quando a
    # chave-base ja pertence a outra identidade.
    used_keys = {}  # chave -> identidade dona

    def _claim_key(candidate, identity):
        owner = used_keys.get(candidate)
        if owner is None or owner == identity:
            used_keys[candidate] = identity
            return candidate
        suffix = 2
        while True:
            alt = f"{candidate}_{suffix}"
            owner = used_keys.get(alt)
            if owner is None or owner == identity:
                used_keys[alt] = identity
                return alt
            suffix += 1

    def _base_key(canonical):
        prefix = _classify_text(canonical)
        if prefix:
            class_counters[prefix] = class_counters.get(prefix, 0) + 1
            return f"{prefix}_{class_counters[prefix]}"
        return slugify(canonical)

    entries = []
    for index, (span, kind, stripped, norm) in enumerate(prepared):
        key = None
        transform = "none"
        if kind == "field":
            if norm not in group_keys:
                group_keys[norm] = _claim_key(_base_key(canonicals[norm]), norm)
            base = group_keys[norm]
            cache = variant_keys.get((norm, stripped))
            if cache is None:
                derived = _variant_transform(stripped, canonicals[norm])
                if derived is not None:
                    cache = (base, derived)
                else:
                    suffix = sum(1 for (n, _v) in variant_keys if n == norm) + 1
                    cache = (f"{base}_v{suffix}", "none")
                variant_keys[(norm, stripped)] = cache
            key, transform = cache
        elif kind in ("list", "image", "manual", "section_title"):
            key = _claim_key(slugify(stripped), stripped)

        entries.append(
            OrderedDict(
                [
                    ("id", index),
                    ("part", span.part),
                    ("color", span.color),
                    ("text", span.text),
                    ("context", span.context.strip()[:200]),
                    ("kind", kind),
                    ("key", key),
                    ("transform", transform),
                    ("section", None),
                    ("label", None),
                    ("reviewed", False),
                ]
            )
        )

    yellow = sum(1 for s in spans if s.color == "yellow")
    red = sum(1 for s in spans if s.color == "red")
    return OrderedDict(
        [
            ("source", source_name),
            ("stats", {"spans": len(spans), "yellow": yellow, "red": red}),
            ("pageAudit", collect_page_geometry(document)),
            ("spans", entries),
        ]
    )


# ---------------------------------------------------------------------------
# apply
# ---------------------------------------------------------------------------
def _humanize(key):
    last = key.split(".")[-1]
    return last.replace("_", " ").strip().capitalize() or key


def _split_ws(text):
    """(indentacao, corpo, sufixo) — o placeholder nunca engole os espacos
    usados para centralizar texto na capa."""
    body = text.strip()
    if not body:
        return text, "", ""
    lead = len(text) - len(text.lstrip())
    trail = len(text) - len(text.rstrip())
    return text[:lead], text[lead : len(text) - trail], text[len(text) - trail :]


def _renumber_group(title):
    m = re.match(r"^\s*(\d+(?:\.\d+)*)\.\d+\.?\s", title)
    return m.group(1) if m else None


def apply_mapping(document, mapping, template_name, revision_label):
    """Tokeniza o documento em memoria conforme o mapping; retorna o manifest.

    Spans kind=field viram ``{{chave}}``; whitespace e ignore so perdem o
    realce (texto original preservado, sem virar campo no manifest); os
    demais kinds (list/image/manual/section_title) sao apenas catalogados no
    manifest nesta fase — as fases 2-4 evoluem a tokenizacao deles.
    ``mapping["generatedListBlocks"]`` (opcional) adiciona blocos kind=list
    sem span marcado no documento (tabela existe mas nao foi realcada na
    marcacao original) -- entram direto no manifest, localizados no render
    por texto do cabecalho da tabela em vez de por realce.
    """
    spans = collect_all_spans(document)
    entries = mapping["spans"]
    if len(spans) != len(entries):
        raise ValueError(
            f"mapping tem {len(entries)} spans mas o documento tem {len(spans)}; "
            "confira se e o MESMO arquivo marcado usado no extract"
        )
    for span, entry in zip(spans, entries):
        if span.text != entry["text"]:
            raise ValueError(
                f"span {entry['id']} divergiu do mapping: "
                f"esperado {entry['text']!r}, encontrado {span.text!r}"
            )

    # pre-scan: valor canonico (primeira ocorrencia transform=none) e tipo
    samples = {}
    multiline = set()
    for entry in entries:
        if entry["kind"] != "field":
            continue
        key = entry["key"]
        body = _split_ws(entry["text"])[1]
        if "\n" in body:
            multiline.add(key)
        if key not in samples and (entry.get("transform") or "none") == "none":
            samples[key] = body

    fields = OrderedDict()
    blocks = OrderedDict()
    image_slots = OrderedDict()
    sections = OrderedDict()
    unreviewed = 0

    for span, entry in zip(spans, entries):
        kind = entry["kind"]
        key = entry.get("key")
        if not entry.get("reviewed"):
            unreviewed += 1

        if kind == "whitespace":
            replace_span_text(span, span.text, drop_highlight=True)
        elif kind == "ignore":
            replace_span_text(span, span.text, drop_highlight=True)
        elif kind == "field":
            transform = entry.get("transform") or "none"
            placeholder = (
                f"{{{{{key}}}}}" if transform == "none" else f"{{{{{key}|{transform}}}}}"
            )
            leading, _body, trailing = _split_ws(span.text)
            replace_span_text(span, f"{leading}{placeholder}{trailing}", drop_highlight=True)
            field = fields.setdefault(
                key,
                OrderedDict(
                    [
                        ("key", key),
                        ("label", entry.get("label") or _humanize(key)),
                        ("section", entry.get("section")),
                        ("type", "multiline" if key in multiline else "text"),
                        ("required", True),
                        ("sampleValue", samples.get(key, "")),
                        ("occurrences", []),
                    ]
                ),
            )
            if entry.get("section") and not field["section"]:
                field["section"] = entry["section"]
            field["occurrences"].append(
                {"spanId": entry["id"], "part": span.part, "transform": transform}
            )
        elif kind == "section_title":
            sections.setdefault(
                key,
                {
                    "sectionKey": key,
                    "defaultTitle": span.text.strip(),
                    "renumberGroup": _renumber_group(span.text),
                },
            )
        elif kind == "image":
            image_slots.setdefault(
                key,
                {
                    "assetKey": key,
                    "label": entry.get("label") or _humanize(key),
                    "page": {"preset": "a4-portrait"},
                    "anchorContext": entry["context"],
                },
            )
        else:  # list | manual
            block = blocks.setdefault(
                key,
                {
                    "key": key,
                    "kind": kind,
                    "label": entry.get("label") or _humanize(key),
                    "anchorContext": entry["context"],
                    "columns": entry.get("columns") or [],
                    "spanIds": [],
                },
            )
            block["spanIds"].append(entry["id"])

    # Blocos tabulares sem span marcado no modelo (ex. contatos internos e
    # externos, plantoes das gerencias no REV 10 — a tabela existe e tem
    # dados reais, mas nem o cabecalho foi realcado na marcacao original).
    # Curados manualmente em `generatedListBlocks` (fora de `spans`, pois nao
    # correspondem a nenhum span real do documento); localizados no render
    # por texto do cabecalho (`headerMatch`), nao por realce.
    for generated in mapping.get("generatedListBlocks") or []:
        key = generated["key"]
        blocks[key] = {
            "key": key,
            "kind": "list",
            "label": generated.get("label") or _humanize(key),
            "anchorContext": None,
            "headerMatch": generated.get("headerMatch") or [],
            "columns": generated.get("columns") or [],
            "spanIds": [],
        }

    manifest = OrderedDict(
        [
            ("templateName", template_name),
            ("revisionLabel", revision_label),
            ("fields", list(fields.values())),
            ("blocks", list(blocks.values())),
            ("imageSlots", list(image_slots.values())),
            ("sections", list(sections.values())),
            ("pageAudit", mapping.get("pageAudit", [])),
            ("stats", {"unreviewedSpans": unreviewed, **mapping.get("stats", {})}),
        ]
    )
    return manifest


# ---------------------------------------------------------------------------
# render (validacao local / round-trip)
# ---------------------------------------------------------------------------
def _apply_transform(value, transform):
    if transform == "upper":
        return value.upper()
    if transform == "title":
        return value.title()
    return value


def render_values(document, values):
    """Substitui ``{{chave}}``/``{{chave|transform}}`` em todos os runs de
    todas as partes. Retorna as chaves nao resolvidas (sem valor)."""
    unresolved = []
    for _part, root in iter_parts(document):
        for run in root.iter(qn("w:r")):
            text = run_text(run)
            if "{{" not in text:
                continue

            def _sub(match):
                key, transform = match.group(1), match.group(2)
                if key not in values:
                    unresolved.append(key)
                    return match.group(0)
                return _apply_transform(str(values[key]), transform or "none")

            new_text = PLACEHOLDER_RE.sub(_sub, text)
            if new_text != text:
                set_run_text(run, new_text)
    return unresolved


def document_text(document):
    """Texto visivel de todas as partes, uma linha por paragrafo (inclui
    tabelas e caixas de texto). Base da comparacao round-trip."""
    lines = []
    for _part, root in iter_parts(document):
        for paragraph in iter_paragraphs(root):
            lines.append(paragraph_text(paragraph))
    return "\n".join(lines)


def sample_values(manifest):
    return {f["key"]: f["sampleValue"] for f in manifest["fields"]}


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------
def _load_yaml(path):
    import yaml

    with open(path, encoding="utf-8") as handle:
        return yaml.safe_load(handle)


def _dump_yaml(data, path):
    import yaml

    class _Dumper(yaml.SafeDumper):
        pass

    _Dumper.add_representer(OrderedDict, lambda d, v: d.represent_dict(v.items()))
    with open(path, "w", encoding="utf-8") as handle:
        yaml.dump(data, handle, Dumper=_Dumper, allow_unicode=True, sort_keys=False, width=120)


def _cmd_extract(args):
    doc = Document(args.docx)
    mapping = build_mapping(doc, source_name=args.docx)
    _dump_yaml(mapping, args.output)
    stats = mapping["stats"]
    print(
        f"{stats['spans']} spans ({stats['yellow']} amarelos, {stats['red']} vermelhos) "
        f"-> {args.output}"
    )


def _cmd_apply(args):
    doc = Document(args.docx)
    mapping = _load_yaml(args.mapping)
    manifest = apply_mapping(doc, mapping, args.name, args.revision)
    doc.save(args.output)
    with open(args.manifest, "w", encoding="utf-8") as handle:
        json.dump(manifest, handle, ensure_ascii=False, indent=2)
    print(
        f"{len(manifest['fields'])} campos, {len(manifest['blocks'])} blocos, "
        f"{len(manifest['sections'])} secoes, {len(manifest['imageSlots'])} slots de imagem"
    )
    if manifest["stats"]["unreviewedSpans"]:
        print(
            f"AVISO: {manifest['stats']['unreviewedSpans']} spans ainda nao revisados na curadoria"
        )


def _cmd_render(args):
    doc = Document(args.docx)
    with open(args.values, encoding="utf-8") as handle:
        values = json.load(handle)
    unresolved = render_values(doc, values)
    doc.save(args.output)
    if unresolved:
        print(f"chaves sem valor: {sorted(set(unresolved))}")


def _cmd_roundtrip(args):
    original_text = document_text(Document(args.docx))

    doc = Document(args.docx)
    mapping = _load_yaml(args.mapping)
    manifest = apply_mapping(doc, mapping, "roundtrip", "roundtrip")

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    rendered = Document(buffer)
    unresolved = render_values(rendered, sample_values(manifest))
    rendered_text = document_text(rendered)

    if unresolved:
        print(f"FALHA: chaves sem valor no round-trip: {sorted(set(unresolved))}")
        return 1
    if rendered_text == original_text:
        print("ROUND-TRIP OK: texto identico ao modelo marcado")
        return 0
    _print_first_diffs(original_text, rendered_text)
    return 1


def _print_first_diffs(expected, actual, limit=10):
    shown = 0
    for i, (e, a) in enumerate(zip(expected.split("\n"), actual.split("\n"))):
        if e != a:
            print(f"linha {i}:\n  esperado: {e!r}\n  obtido:   {a!r}")
            shown += 1
            if shown >= limit:
                break
    print(f"FALHA: {shown}+ linhas divergentes")


def main(argv=None):
    parser = argparse.ArgumentParser(prog="paec_tokenizer", description=__doc__)
    sub = parser.add_subparsers(dest="command", required=True)

    p = sub.add_parser("extract", help="extrai spans e gera mapping.yaml de curadoria")
    p.add_argument("docx")
    p.add_argument("-o", "--output", required=True)
    p.set_defaults(func=_cmd_extract)

    p = sub.add_parser("apply", help="tokeniza o docx marcado conforme o mapping")
    p.add_argument("docx")
    p.add_argument("-m", "--mapping", required=True)
    p.add_argument("-o", "--output", required=True)
    p.add_argument("--manifest", required=True)
    p.add_argument("--name", default="PAEC AXIA")
    p.add_argument("--revision", required=True)
    p.set_defaults(func=_cmd_apply)

    p = sub.add_parser("render", help="preenche placeholders com valores (validacao local)")
    p.add_argument("docx")
    p.add_argument("-v", "--values", required=True)
    p.add_argument("-o", "--output", required=True)
    p.set_defaults(func=_cmd_render)

    p = sub.add_parser("roundtrip", help="valida extract->apply->render == original")
    p.add_argument("docx")
    p.add_argument("-m", "--mapping", required=True)
    p.set_defaults(func=_cmd_roundtrip)

    args = parser.parse_args(argv)
    return args.func(args) or 0


if __name__ == "__main__":
    sys.exit(main())
