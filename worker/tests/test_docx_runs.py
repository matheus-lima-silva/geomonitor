import io

import pytest
from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, qn

from worker.docx_runs import (
    collect_all_spans,
    collect_highlight_spans,
    collect_page_geometry,
    iter_parts,
    paragraph_text,
    replace_span_text,
    run_text,
    set_run_text,
)


def _add_runs(paragraph, *specs):
    """specs: (texto, cor|None). Cores: 'yellow' | 'red'."""
    colors = {"yellow": WD_COLOR_INDEX.YELLOW, "red": WD_COLOR_INDEX.RED}
    for text, color in specs:
        run = paragraph.add_run(text)
        if color:
            run.font.highlight_color = colors[color]
    return paragraph


def _reload(document):
    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return Document(buffer)


# ---------------------------------------------------------------------------
# collect_highlight_spans
# ---------------------------------------------------------------------------
def test_merge_de_runs_fragmentados_com_mesmo_highlight():
    doc = Document()
    p = doc.add_paragraph()
    _add_runs(
        p,
        ("Usina de ", None),
        ("Ger", "yellow"),
        ("encia ", "yellow"),
        ("Regional", "yellow"),
        (" gera energia.", None),
    )
    spans = collect_highlight_spans(doc.element)
    assert len(spans) == 1
    assert spans[0].text == "Gerencia Regional"
    assert spans[0].color == "yellow"
    assert len(spans[0].runs) == 3
    assert spans[0].context == "Usina de Gerencia Regional gera energia."


def test_run_sem_highlight_e_troca_de_cor_fecham_o_span():
    doc = Document()
    p = doc.add_paragraph()
    _add_runs(
        p,
        ("A", "yellow"),
        ("meio", None),
        ("B", "yellow"),
        ("C", "red"),
    )
    spans = collect_highlight_spans(doc.element)
    assert [(s.text, s.color) for s in spans] == [
        ("A", "yellow"),
        ("B", "yellow"),
        ("C", "red"),
    ]


def test_proof_err_entre_runs_nao_quebra_o_span():
    doc = Document()
    p = doc.add_paragraph()
    _add_runs(p, ("Mari", "yellow"), ("mbondo", "yellow"))
    runs = p._p.findall(qn("w:r"))
    proof = p._p.makeelement(qn("w:proofErr"), {qn("w:type"): "spellStart"})
    runs[0].addnext(proof)
    spans = collect_highlight_spans(doc.element)
    assert len(spans) == 1
    assert spans[0].text == "Marimbondo"


def test_spans_em_tabela_e_header():
    doc = Document()
    table = doc.add_table(rows=1, cols=1)
    _add_runs(table.rows[0].cells[0].paragraphs[0], ("(34) 99730-4626", "yellow"))

    header = doc.sections[0].header
    header.is_linked_to_previous = False
    _add_runs(header.paragraphs[0], ("OO|SEGSRMMB", "yellow"))

    spans = collect_all_spans(doc)
    texts = {(s.part.split("/")[-1], s.text) for s in spans}
    assert ("document", "(34) 99730-4626") in texts
    assert any(part.startswith("header") and text == "OO|SEGSRMMB" for part, text in texts)


def test_spans_dentro_de_caixa_de_texto_vml():
    doc = Document()
    xml = (
        '<w:p %s xmlns:v="urn:schemas-microsoft-com:vml"><w:r><w:pict><v:shape><v:textbox><w:txbxContent>'
        '<w:p><w:r><w:rPr><w:highlight w:val="yellow"/></w:rPr><w:t>GERENTE </w:t></w:r>'
        '<w:r><w:rPr><w:highlight w:val="yellow"/></w:rPr><w:t>PCH X</w:t></w:r></w:p>'
        "</w:txbxContent></v:textbox></v:shape></w:pict></w:r></w:p>"
    ) % nsdecls("w")
    doc.element.body.insert(0, parse_xml(xml))
    spans = collect_highlight_spans(doc.element)
    assert len(spans) == 1
    assert spans[0].text == "GERENTE PCH X"


def test_span_com_quebra_de_linha_no_run():
    doc = Document()
    p = doc.add_paragraph()
    run = p.add_run("Tel.: ")
    run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    run.add_break()
    run2 = p.add_run("17 98820-6833")
    run2.font.highlight_color = WD_COLOR_INDEX.YELLOW
    spans = collect_highlight_spans(doc.element)
    assert len(spans) == 1
    assert spans[0].text == "Tel.: \n17 98820-6833"


# ---------------------------------------------------------------------------
# replace_span_text / set_run_text
# ---------------------------------------------------------------------------
def test_replace_span_preserva_formatacao_e_remove_highlight():
    doc = Document()
    p = doc.add_paragraph()
    first = p.add_run("Mari")
    first.bold = True
    first.font.highlight_color = WD_COLOR_INDEX.YELLOW
    second = p.add_run("mbondo")
    second.font.highlight_color = WD_COLOR_INDEX.YELLOW

    span = collect_highlight_spans(doc.element)[0]
    replace_span_text(span, "{{usina}}")

    reloaded = _reload(doc)
    runs = reloaded.paragraphs[0].runs
    assert [r.text for r in runs] == ["{{usina}}"]
    assert runs[0].bold is True
    assert runs[0].font.highlight_color is None
    assert collect_highlight_spans(reloaded.element) == []


def test_set_run_text_multilinha_gera_br_e_tab():
    doc = Document()
    run = doc.add_paragraph().add_run("x")
    set_run_text(run._r, "linha1\nlinha2\tcol")
    assert run_text(run._r) == "linha1\nlinha2\tcol"
    assert len(run._r.findall(qn("w:br"))) == 1
    assert len(run._r.findall(qn("w:tab"))) == 1


def test_set_run_text_preserva_espacos_de_borda():
    doc = Document()
    run = doc.add_paragraph().add_run("x")
    set_run_text(run._r, "  {{usina}}  ")
    reloaded = _reload(doc)
    assert reloaded.paragraphs[0].runs[0].text == "  {{usina}}  "


# ---------------------------------------------------------------------------
# auditoria de paginas
# ---------------------------------------------------------------------------
def test_collect_page_geometry_reporta_secoes():
    doc = Document()
    sections = collect_page_geometry(doc)
    assert len(sections) == 1
    assert sections[0]["widthTwips"] > 0
    assert sections[0]["heightTwips"] > 0
    assert sections[0]["orientation"] == "portrait"
