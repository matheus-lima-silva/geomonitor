import json
import os

import pytest
from docx import Document

from worker.monthly_report_renderer import (
    activities_visible_on_date,
    build_holiday_map,
    build_quadro_weeks,
    get_date_range,
    render_context_to_docx,
)

FIXTURE_PATH = os.path.join(os.path.dirname(__file__), "fixtures", "monthly_report_render_model.json")


def _load_fixture():
    with open(FIXTURE_PATH, encoding="utf-8") as handle:
        return json.load(handle)


def _build_context(report_overrides=None):
    fixture = _load_fixture()
    report = dict(fixture["monthlyReport"])
    report.update(report_overrides or {})
    return {
        "job": {"id": "JOB-1", "kind": "monthly_report"},
        "renderModel": {"monthlyReport": report},
    }


def _render(tmp_path, report_overrides=None):
    output = os.path.join(str(tmp_path), "out.docx")
    render_context_to_docx(_build_context(report_overrides), output)
    return output


def _all_text(path):
    doc = Document(path)
    chunks = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    chunks.append(p.text)
    return "\n".join(chunks)


# ----------------------------------------------------------------------------
# Paridade de dominio com utils JS (mesma fixture do snapshot Vitest)
# ----------------------------------------------------------------------------
def test_get_date_range_periodo_16_a_15():
    import datetime
    start, end = get_date_range(2026, 4)  # Maio/2026
    assert start == datetime.date(2026, 4, 16)
    assert end == datetime.date(2026, 5, 15)


def test_quadro_weeks_cobre_o_periodo_em_semanas_completas():
    fixture = _load_fixture()["monthlyReport"]
    weeks = build_quadro_weeks(fixture, fixture["engineers"][0]["activities"])
    assert len(weeks) == 5  # 12/04 a 16/05/2026
    assert all(len(week) == 7 for week in weeks)
    assert weeks[0][0]["dateKey"] == "2026-04-12"
    assert weeks[4][6]["dateKey"] == "2026-05-16"


def test_feriado_explicito_esconde_multi_dia_e_mostra_nome():
    fixture = _load_fixture()["monthlyReport"]
    weeks = build_quadro_weeks(fixture, fixture["engineers"][0]["activities"])
    cells = {c["dateKey"]: c for week in weeks for c in week}

    tiradentes = cells["2026-04-21"]
    assert tiradentes["holidayName"] == "Tiradentes"
    assert tiradentes["activities"] == []

    # Continuacao reaparece no dia util seguinte, sem isFirstDay.
    day22 = cells["2026-04-22"]
    assert any("LT 500kv" in a["description"] and not a["isFirstDay"] for a in day22["activities"])

    day16 = cells["2026-04-16"]
    assert any("LT 500kv" in a["description"] and a["isFirstDay"] for a in day16["activities"])


def test_feriados_sao_lista_explicita_sem_computo_de_oficiais():
    # Corpus Christi 2026 (04/06) nao esta na lista => dia util normal.
    hmap = build_holiday_map([{"date": "2026-04-21", "name": "Tiradentes"}])
    assert set(hmap.keys()) == {"2026-04-21"}

    visible = activities_visible_on_date(
        [{"startDate": "2026-04-30", "endDate": "2026-05-04", "category": "doc", "description": "x"}],
        "2026-05-01",  # sexta; Dia do Trabalho NAO marcado
        set(hmap.keys()),
    )
    assert len(visible) == 1


# ----------------------------------------------------------------------------
# Estrutura do DOCX gerado
# ----------------------------------------------------------------------------
def test_docx_estrutura_completa(tmp_path):
    output = _render(tmp_path)
    doc = Document(output)
    text = _all_text(output)

    # Capa + corpo = 2 secoes; capa com margens zero, corpo 2cm.
    assert len(doc.sections) == 2
    assert doc.sections[0].left_margin.cm == pytest.approx(0, abs=0.01)
    assert doc.sections[1].left_margin.cm == pytest.approx(2.0, abs=0.01)

    assert "Relatório Mensal de Acompanhamento dos Serviços" in text
    assert "MAIO - 2026" in text
    assert "Período: 16/04/2026 a 15/05/2026" in text
    assert "SUMÁRIO:" in text
    assert "1 INTRODUÇÃO" in text
    assert "2 ATIVIDADES REALIZADAS NO PERÍODO" in text
    assert "2.1 Atividades que o eng. Matheus Lima realizou" in text
    assert "2.1.1 Resumo por projeto:" in text
    assert "2.2 Atividades que o eng. Victor Britto realizou" in text
    assert "3 CONCLUSÃO" in text
    assert "LT 500 kV Bom Despacho–Ouro Preto:" in text
    assert "UHE Serra da Mesa:" in text


def test_docx_headings_usam_estilos_de_toc(tmp_path):
    doc = Document(_render(tmp_path))
    by_style = {}
    for p in doc.paragraphs:
        by_style.setdefault(p.style.name, []).append(p.text)
    assert "1 INTRODUÇÃO" in by_style.get("Heading 1", [])
    assert "2.1 Atividades que o eng. Matheus Lima realizou" in by_style.get("Heading 2", [])
    assert "2.1.1 Resumo por projeto:" in by_style.get("Heading 3", [])

    # Normal = Verdana (APENSO D).
    normal = doc.styles["Normal"]
    assert normal.font.name == "Verdana"


def test_docx_quadro_por_engenheiro_com_marcadores(tmp_path):
    output = _render(tmp_path)
    doc = Document(output)
    # 1 tabela de quadro por engenheiro (6 linhas = header + 5 semanas, 7 colunas).
    quadros = [t for t in doc.tables if len(t.columns) == 7]
    assert len(quadros) == 2
    for quadro in quadros:
        assert len(quadro.rows) == 6
        header = [cell.text for cell in quadro.rows[0].cells]
        assert header == ["DOM", "SEG", "TER", "QUA", "QUI", "SEX", "SÁB"]

    full = _all_text(output)
    assert "●" in full
    assert "↳" in full
    assert "★ Tiradentes" in full
    assert "LEGENDA" in full


def test_docx_variante_preenchido_aplica_shading_no_paragrafo(tmp_path):
    output = _render(tmp_path, {"quadroStyle": "preenchido"})
    doc = Document(output)
    quadro = next(t for t in doc.tables if len(t.columns) == 7)
    shaded = []
    for row in quadro.rows[1:]:
        for cell in row.cells:
            for p in cell.paragraphs:
                ppr = p._p.find("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pPr")
                if ppr is None:
                    continue
                shd = ppr.find("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}shd")
                if shd is not None:
                    shaded.append(shd.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill"))
    assert "166534" in shaded  # vistoria solida no primeiro dia


def test_docx_update_fields_e_toc(tmp_path):
    output = _render(tmp_path)
    doc = Document(output)
    settings_xml = doc.settings.element.xml
    assert "updateFields" in settings_xml

    body_xml = doc.element.body.xml
    assert "TOC \\o" in body_xml


def test_docx_header_com_logos_e_rodape_classificacao(tmp_path):
    doc = Document(_render(tmp_path))
    content = doc.sections[1]
    header_xml = content.header._element.xml
    assert "blip" in header_xml  # imagens dos logos embutidas
    footer_text = "\n".join(p.text for p in content.footer.paragraphs)
    assert "Classificação: Interna" in footer_text


def test_docx_capa_com_arte_ancorada_atras_do_texto(tmp_path):
    doc = Document(_render(tmp_path))
    body_xml = doc.element.body.xml
    assert "behindDoc=\"1\"" in body_xml


def test_docx_placeholders_quando_textos_vazios(tmp_path):
    full = _all_text(_render(tmp_path, {"intro": "", "conclusao": ""}))
    assert "Introdução ainda não escrita" in full
    assert "Conclusão ainda não escrita" in full
