"""Cobertura do Quadro 1 (Graduacao de Criticidade) renderizado via marcador.

O texto da "Descricao das Atividades" pode conter um marcador que o worker
substitui pela tabela de graduacao de criticidade ao gerar o DOCX.
"""

from docx import Document

from worker.docx_renderer import (
    QUADRO_CRITICIDADE_HEADER,
    QUADRO_CRITICIDADE_MARKER,
    QUADRO_CRITICIDADE_ROWS,
    add_criticality_grading_table,
    add_numbered_text_section,
    is_quadro_criticidade_marker,
)


def test_marker_predicate_matches_only_the_marker_line():
    assert is_quadro_criticidade_marker(QUADRO_CRITICIDADE_MARKER)
    # Tolerante a acentos/caixa/espacos no boundary JS->Python.
    assert is_quadro_criticidade_marker("  [INSERIR QUADRO 1 de graduacao de criticidade]  ")
    # Nao casa o paragrafo que descreve o quadro (comeca com "O Quadro 1...").
    assert not is_quadro_criticidade_marker(
        "O Quadro 1 define quatro níveis de criticidade a partir de critérios técnicos."
    )
    assert not is_quadro_criticidade_marker("Tabela 1 – Graduação de criticidade.")
    assert not is_quadro_criticidade_marker("")


def test_section_with_marker_renders_the_grading_table():
    document = Document()
    text = "\n\n".join(
        [
            "Parágrafo introdutório das atividades.",
            "Tabela 1 – Graduação de criticidade.",
            QUADRO_CRITICIDADE_MARKER,
            "Parágrafo após o quadro.",
        ]
    )

    add_numbered_text_section(
        document,
        "DESCRIÇÃO DAS ATIVIDADES",
        text,
        marker_predicate=is_quadro_criticidade_marker,
        marker_renderer=add_criticality_grading_table,
    )

    assert len(document.tables) == 1
    table = document.tables[0]
    assert len(table.rows) == 1 + len(QUADRO_CRITICIDADE_ROWS)
    assert len(table.columns) == len(QUADRO_CRITICIDADE_HEADER)

    header_cells = [cell.text for cell in table.rows[0].cells]
    assert header_cells == QUADRO_CRITICIDADE_HEADER
    # Primeira coluna de criterio preservada.
    assert table.rows[1].cells[0].text == "Tipo de Erosão"
    assert table.rows[1].cells[4].text.startswith("Voçorocas Ativas")

    # A linha-marcador nao vira paragrafo de corpo; os demais textos sim.
    paragraph_texts = [p.text for p in document.paragraphs]
    assert QUADRO_CRITICIDADE_MARKER not in paragraph_texts
    assert "Parágrafo introdutório das atividades." in paragraph_texts
    assert "Parágrafo após o quadro." in paragraph_texts


def test_section_without_marker_renders_no_table():
    document = Document()
    text = "Apenas um parágrafo de atividades, sem quadro de criticidade."

    add_numbered_text_section(
        document,
        "DESCRIÇÃO DAS ATIVIDADES",
        text,
        marker_predicate=is_quadro_criticidade_marker,
        marker_renderer=add_criticality_grading_table,
    )

    assert document.tables == []
