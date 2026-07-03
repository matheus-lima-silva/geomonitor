import io
import json
import os
from unittest.mock import Mock

import pytest
from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml.ns import qn

from worker.docx_runs import collect_all_spans
from worker.paec_renderer import render_paec_to_docx
from worker.job_processor import build_output_file_name, process_paec_report_job
from worker.runtime import WorkerClient

FIXTURE_PATH = os.path.join(os.path.dirname(__file__), "fixtures", "paec_render_model.json")


def _load_fixture():
    with open(FIXTURE_PATH, encoding="utf-8") as handle:
        return json.load(handle)


def _build_context(overrides=None):
    fixture = _load_fixture()
    paec = dict(fixture["paecReport"])
    paec.update(overrides or {})
    return {
        "job": {"id": "JOB-1", "kind": "paec_report", "paecPlantId": paec["plant"]["id"]},
        "renderModel": {"paecReport": paec},
    }


def _template_bytes():
    """Template tokenizado sintetico com os padroes reais: placeholder puro,
    com transform, com indentacao de capa e em celula de tabela."""
    doc = Document()
    doc.add_paragraph("Plano da usina {{usina}}.")
    doc.add_paragraph("PROTEGER AS INSTALACOES DA {{usina|upper}};")
    doc.add_paragraph("          {{usina}} ")
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].paragraphs[0].add_run("CNPJ")
    table.rows[0].cells[1].paragraphs[0].add_run("{{cnpj_1}}")
    table.rows[1].cells[0].paragraphs[0].add_run("Endereco")
    table.rows[1].cells[1].paragraphs[0].add_run("{{endereco_rep}}")
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def _template_bytes_with_list_block():
    """Template sintetico reproduzindo o padrao real: heading realcado de
    vermelho (kind=list, ainda nao tokenizado nesta fase) seguido da tabela
    com header + 1 linha-exemplo (a chave em COLUNAS do fixture e nome/telefone)."""
    doc = Document()
    doc.add_paragraph("Plano da usina {{usina}}.")
    heading = doc.add_paragraph()
    heading.add_run("anexo IV - RELACAO DE BRIGADISTAS").font.highlight_color = WD_COLOR_INDEX.RED
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].paragraphs[0].add_run("NOME")
    table.rows[0].cells[1].paragraphs[0].add_run("TELEFONE")
    table.rows[1].cells[0].paragraphs[0].add_run("Exemplo Marimbondo")
    table.rows[1].cells[1].paragraphs[0].add_run("(00) 0000-0000")
    doc.add_paragraph("depois da tabela")
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def _template_bytes_with_generated_list_block():
    """Template sintetico reproduzindo o padrao dos blocos SEM span marcado
    (contatos internos/externos no REV 10): tabela com header + 1 linha-
    exemplo, sem realce nenhum (nem no cabecalho, nem em nenhuma outra
    parte perto dela) -- so o texto do header identifica a tabela."""
    doc = Document()
    doc.add_paragraph("Plano da usina {{usina}}.")
    doc.add_paragraph("anexo II - plano de comunicacao - CONTATOS INTERNOS")
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].paragraphs[0].add_run("APOIO")
    table.rows[0].cells[1].paragraphs[0].add_run("TELEFONE")
    table.rows[1].cells[0].paragraphs[0].add_run("Exemplo Marimbondo")
    table.rows[1].cells[1].paragraphs[0].add_run("(00) 0000-0000")
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def _template_bytes_with_sections():
    """Template sintetico reproduzindo o padrao real das secoes 12.1.x:
    heading realcado (kind=section_title, nunca tokenizado pelo apply)
    seguido de paragrafos de conteudo ate o proximo heading."""
    doc = Document()
    doc.add_paragraph("Plano da usina {{usina}}.")
    for number, label in [("12.1.1", "Recurso A"), ("12.1.2", "Recurso B"), ("12.1.3", "Recurso C")]:
        heading = doc.add_paragraph()
        heading.add_run(f"{number}. {label}").font.highlight_color = WD_COLOR_INDEX.YELLOW
        doc.add_paragraph(f"Conteudo do {label}.")
    doc.add_paragraph("Fim do documento.")
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


SAMPLE_SECTIONS = [
    {"sectionKey": "recurso_a", "defaultTitle": "12.1.1. Recurso A", "renumberGroup": "12.1"},
    {"sectionKey": "recurso_b", "defaultTitle": "12.1.2. Recurso B", "renumberGroup": "12.1"},
    {"sectionKey": "recurso_c", "defaultTitle": "12.1.3. Recurso C", "renumberGroup": "12.1"},
]


def _all_text(path):
    doc = Document(path)
    chunks = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    chunks.append(p.text)
    return "\n".join(chunks)


def _render(tmp_path, overrides=None, template=None):
    output = os.path.join(str(tmp_path), "out.docx")
    result = render_paec_to_docx(_build_context(overrides), template or _template_bytes(), output)
    return output, result


# ---------------------------------------------------------------------------
# render_paec_to_docx
# ---------------------------------------------------------------------------
def test_substitui_placeholders_com_transform(tmp_path):
    output, _result = _render(tmp_path)
    text = _all_text(output)
    assert "Plano da usina PCH Anta." in text
    assert "PROTEGER AS INSTALACOES DA PCH ANTA;" in text
    assert "{{usina" not in text


def test_preserva_indentacao_da_capa(tmp_path):
    output, _result = _render(tmp_path)
    doc = Document(output)
    cover = [p.text for p in doc.paragraphs if p.text.strip() == "PCH Anta" and p.text != "PCH Anta"]
    assert cover, "paragrafo de capa indentado deveria manter os espacos"
    assert cover[0].startswith("          ")


def test_valor_multilinha_vira_quebras_reais(tmp_path):
    output, _result = _render(tmp_path)
    doc = Document(output)
    endereco_cell = doc.tables[0].rows[1].cells[1]
    assert endereco_cell.paragraphs[0].text == "Rua A, 123\nCentro - Sapucaia/RJ".replace("\n", "\n")
    # o \n do valor vira w:br (nao texto literal "\\n")
    xml = endereco_cell.paragraphs[0]._p.xml
    assert "<w:br/>" in xml


def test_campo_sem_valor_vira_pendente_realcado(tmp_path):
    output, result = _render(tmp_path)
    text = _all_text(output)
    assert "[[PENDENTE: CNPJ]]" in text

    doc = Document(output)
    spans = collect_all_spans(doc)
    assert any(s.color == "yellow" and "PENDENTE: CNPJ" in s.text for s in spans)

    fields = [p for p in result["pendencies"] if p["kind"] == "field"]
    assert [p["key"] for p in fields] == ["cnpj_1"]
    assert fields[0]["label"] == "CNPJ"


def test_token_desconhecido_vira_unresolved_token(tmp_path):
    doc = Document()
    doc.add_paragraph("{{image:rota_de_fuga}}")
    buffer = io.BytesIO()
    doc.save(buffer)

    output, result = _render(tmp_path, template=buffer.getvalue())
    unresolved = [p for p in result["pendencies"] if p["kind"] == "unresolved_token"]
    assert len(unresolved) == 1
    assert "image:rota_de_fuga" in unresolved[0]["key"]
    # o token permanece visivel e realcado no documento
    assert "{{image:rota_de_fuga}}" in _all_text(output)


def test_repassa_pendencias_de_blocos_do_backend(tmp_path):
    _output, result = _render(tmp_path)
    kinds = [(p["kind"], p["key"]) for p in result["pendencies"]]
    assert ("list", "brigadistas") in kinds
    assert ("manual_block", "anexo_vii_rota_de_fuga") in kinds
    assert result["stats"] == {"fieldsFilled": 2, "fieldsTotal": 3}


def test_bloco_list_clona_linha_por_item_e_remove_realce_do_anexo(tmp_path):
    output, _result = _render(tmp_path, template=_template_bytes_with_list_block())

    doc = Document(output)
    table = doc.tables[0]
    assert [c.text for c in table.rows[0].cells] == ["NOME", "TELEFONE"]
    assert [c.text for c in table.rows[1].cells] == ["Fulano de Tal", "(11) 1111-1111"]
    assert [c.text for c in table.rows[2].cells] == ["Ciclana da Silva", "(11) 2222-2222"]
    assert len(table.rows) == 3
    # a linha-exemplo do modelo (Marimbondo) nao pode sobrar no documento
    text = _all_text(output)
    assert "Exemplo Marimbondo" not in text

    spans = collect_all_spans(doc)
    assert not any("RELACAO DE BRIGADISTAS" in s.text for s in spans), (
        "realce do anexo devia ter sido removido apos a tabela ser renderizada"
    )


def test_bloco_list_sem_item_vira_uma_linha_pendente(tmp_path):
    output, _result = _render(
        tmp_path,
        overrides={"listItems": {}},
        template=_template_bytes_with_list_block(),
    )
    doc = Document(output)
    table = doc.tables[0]
    assert len(table.rows) == 2
    assert table.rows[1].cells[0].text == "[[PENDENTE: Relacao de brigadistas]]"
    assert table.rows[1].cells[1].text == ""

    spans = collect_all_spans(doc)
    assert any(s.color == "yellow" and "PENDENTE: Relacao de brigadistas" in s.text for s in spans)


def test_bloco_list_sem_span_localiza_pela_tabela_via_headerMatch(tmp_path):
    """Bloco gerado (mapping.yaml['generatedListBlocks']) nao tem
    anchorContext -- nenhum span foi marcado no documento original pra essa
    tabela. Precisa ser localizado so pelo texto do cabecalho."""
    output, _result = _render(
        tmp_path,
        overrides={
            "blocks": [{
                "key": "contatos_internos",
                "kind": "list",
                "label": "Contatos internos",
                "anchorContext": None,
                "headerMatch": ["APOIO", "TELEFONE"],
                "columns": [
                    {"key": "apoio", "label": "Apoio"},
                    {"key": "telefone", "label": "Telefone"},
                ],
            }],
            "listItems": {
                "contatos_internos": [{"apoio": "Gerente", "telefone": "(21) 1111-1111"}],
            },
        },
        template=_template_bytes_with_generated_list_block(),
    )

    doc = Document(output)
    table = doc.tables[0]
    assert [c.text for c in table.rows[0].cells] == ["APOIO", "TELEFONE"]
    assert [c.text for c in table.rows[1].cells] == ["Gerente", "(21) 1111-1111"]
    assert len(table.rows) == 2
    assert "Exemplo Marimbondo" not in _all_text(output)


def test_secao_desligada_some_do_documento_e_remanescentes_renumeram(tmp_path):
    output, _result = _render(
        tmp_path,
        overrides={
            "sections": SAMPLE_SECTIONS,
            "sectionFlags": {"recurso_b": {"enabled": False}},
        },
        template=_template_bytes_with_sections(),
    )
    text = _all_text(output)
    assert "Recurso B" not in text
    assert "Conteudo do Recurso B" not in text
    assert "12.1.1. Recurso A" in text
    assert "Conteudo do Recurso A." in text
    # Recurso C era 12.1.3, com B fora vira 12.1.2 (renumeracao sequencial)
    assert "12.1.2. Recurso C" in text
    assert "12.1.3. Recurso C" not in text
    assert "Conteudo do Recurso C." in text
    assert "Fim do documento." in text


def test_secao_com_title_override_troca_o_titulo_mantendo_a_posicao(tmp_path):
    output, _result = _render(
        tmp_path,
        overrides={
            "sections": SAMPLE_SECTIONS,
            "sectionFlags": {"recurso_b": {"titleOverride": "Recurso B Customizado"}},
        },
        template=_template_bytes_with_sections(),
    )
    doc = Document(output)
    headings = [p.text for p in doc.paragraphs if p.text.startswith("12.1.")]
    assert headings == ["12.1.1. Recurso A", "12.1.2. Recurso B Customizado", "12.1.3. Recurso C"]
    # conteudo da secao (nao e o heading) continua intacto, so o titulo mudou
    assert "Conteudo do Recurso B." in _all_text(output)


def test_secoes_processadas_perdem_o_realce_mesmo_sem_flag(tmp_path):
    output, _result = _render(
        tmp_path,
        overrides={"sections": SAMPLE_SECTIONS, "sectionFlags": {}},
        template=_template_bytes_with_sections(),
    )
    doc = Document(output)
    spans = collect_all_spans(doc)
    assert not any("Recurso" in s.text for s in spans)


def test_secao_desligada_marca_toc_para_atualizar(tmp_path):
    output, _result = _render(
        tmp_path,
        overrides={
            "sections": SAMPLE_SECTIONS,
            "sectionFlags": {"recurso_b": {"enabled": False}},
        },
        template=_template_bytes_with_sections(),
    )
    doc = Document(output)
    assert doc.settings.element.find(qn("w:updateFields")) is not None


def test_sem_secoes_flags_nao_marca_toc(tmp_path):
    output, _result = _render(
        tmp_path,
        overrides={"sections": SAMPLE_SECTIONS, "sectionFlags": {}},
        template=_template_bytes_with_sections(),
    )
    doc = Document(output)
    assert doc.settings.element.find(qn("w:updateFields")) is None


def test_ficha_completa_nao_gera_pendencia_de_campo(tmp_path):
    values = {"usina": "PCH Anta", "cnpj_1": "00.001.180/0038-18", "endereco_rep": "Rua A"}
    _output, result = _render(tmp_path, overrides={"values": values})
    assert [p for p in result["pendencies"] if p["kind"] in ("field", "unresolved_token")] == []


# ---------------------------------------------------------------------------
# job_processor: handler + nome de arquivo
# ---------------------------------------------------------------------------
def test_build_output_file_name_paec():
    context = _build_context()
    assert build_output_file_name(context) == "PAEC - PCH Anta.docx"


def test_process_paec_report_job_baixa_template_e_completa(tmp_path):
    client = Mock(spec=WorkerClient)
    client.download_media_content.return_value = {
        "statusCode": 200,
        "buffer": _template_bytes(),
        "contentType": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    }
    client.create_output_media.return_value = {
        "id": "MEDIA-out",
        "upload": {"href": "/api/media/MEDIA-out/upload", "method": "PUT"},
    }

    result = process_paec_report_job(client, "JOB-1", _build_context(), str(tmp_path))

    client.download_media_content.assert_called_once_with("MEDIA-tok")
    assert result["status"] == "completed"
    assert result["outputDocxMediaId"] == "MEDIA-out"
    assert any(p["key"] == "cnpj_1" for p in result["resultMeta"]["pendencies"])

    created = client.create_output_media.call_args
    assert created.kwargs["purpose"] == "paec_report_docx"
    assert created.kwargs["file_name"] == "PAEC - PCH Anta.docx"
    client.upload_media_binary.assert_called_once()
    client.complete_media_upload.assert_called_once()


def test_process_paec_report_job_falha_sem_template_media(tmp_path):
    client = Mock(spec=WorkerClient)
    context = _build_context({"template": {"id": "PAECT-1", "revisionLabel": "REV", "tokenizedDocxMediaId": ""}})
    result = process_paec_report_job(client, "JOB-1", context, str(tmp_path))
    assert result["status"] == "failed"
    assert "tokenizedDocxMediaId" in result["errorLog"]
    client.download_media_content.assert_not_called()
