import base64
import io
import json
import re
import struct
import unittest
import zipfile

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor

from worker.docx_renderer import (
    apply_axia_formatting_compound,
    apply_body_font,
    create_document_from_template,
)
from worker.exif_gps import extract_gps_latlon
from worker.runtime import WorkerClient, WorkerRuntime


SAMPLE_PNG_BYTES = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAwMCAO7Z2ioAAAAASUVORK5CYII="
)


def build_png(width, height):
    """Monta um PNG RGB minimo (IHDR + IDAT + IEND) via stdlib (sem Pillow),
    usado para exercitar o teto de altura de fotos verticais."""
    import zlib

    def _chunk(typ, data):
        body = typ + data
        return struct.pack(">I", len(data)) + body + struct.pack(">I", zlib.crc32(body) & 0xFFFFFFFF)

    sig = b"\x89PNG\r\n\x1a\n"
    ihdr = struct.pack(">IIBBBBB", width, height, 8, 2, 0, 0, 0)
    raw = b"".join(b"\x00" + b"\xff\x00\x00" * width for _ in range(height))
    return sig + _chunk(b"IHDR", ihdr) + _chunk(b"IDAT", zlib.compress(raw)) + _chunk(b"IEND", b"")


def build_minimal_jpeg_with_gps(
    lat_dms=(22, 30, 0),
    lat_ref="S",
    lon_dms=(43, 15, 0),
    lon_ref="W",
    break_den=False,
    bad_offset=False,
):
    """Monta um JPEG minimo (SOI + APP1/Exif + EOI) com um GPS IFD, via struct.

    Layout do bloco TIFF (little-endian), offsets relativos ao inicio do TIFF:
      8   IFD0 (1 entry -> GPS IFD pointer @26)
      26  GPS IFD (4 entries: LatRef, Lat@80, LonRef, Lon@104)
      80  3 rationals da latitude
      104 3 rationals da longitude
    """
    def rationals(dms):
        out = b""
        for value in dms:
            den = 0 if break_den else 1
            out += struct.pack("<II", int(value), den)
        return out

    lat_off = 9999 if bad_offset else 80
    tiff = b"II" + struct.pack("<HI", 42, 8)
    # IFD0 @8
    tiff += struct.pack("<H", 1)
    tiff += struct.pack("<HHI", 0x8825, 4, 1) + struct.pack("<I", 26)
    tiff += struct.pack("<I", 0)
    # GPS IFD @26
    tiff += struct.pack("<H", 4)
    tiff += struct.pack("<HHI", 0x0001, 2, 2) + (lat_ref.encode() + b"\x00\x00\x00")[:4]
    tiff += struct.pack("<HHI", 0x0002, 5, 3) + struct.pack("<I", lat_off)
    tiff += struct.pack("<HHI", 0x0003, 2, 2) + (lon_ref.encode() + b"\x00\x00\x00")[:4]
    tiff += struct.pack("<HHI", 0x0004, 5, 3) + struct.pack("<I", 104)
    tiff += struct.pack("<I", 0)
    # data @80 (lat), @104 (lon)
    tiff += rationals(lat_dms)
    tiff += rationals(lon_dms)

    exif = b"Exif\x00\x00" + tiff
    app1 = b"\xff\xe1" + struct.pack(">H", len(exif) + 2) + exif
    return b"\xff\xd8" + app1 + b"\xff\xd9"


def read_docx_entry(docx_bytes, entry_name):
    with zipfile.ZipFile(io.BytesIO(docx_bytes)) as docx_zip:
        return docx_zip.read(entry_name).decode("utf-8")


def read_docx_plain_text(docx_bytes):
    """Return concatenated text from every <w:t> element in word/document.xml.

    Useful for assertions on content that spans multiple runs separated by
    field codes or formatting boundaries (e.g. SEQ Foto fields).
    """
    xml = read_docx_entry(docx_bytes, "word/document.xml")
    texts = re.findall(r"<w:t(?:\s[^>]*)?>([^<]*)</w:t>", xml)
    return "".join(texts)


def read_kmz_entry(kmz_bytes, entry_name):
    with zipfile.ZipFile(io.BytesIO(kmz_bytes)) as kmz_zip:
        return kmz_zip.read(entry_name)


def read_kmz_namelist(kmz_bytes):
    with zipfile.ZipFile(io.BytesIO(kmz_bytes)) as kmz_zip:
        return kmz_zip.namelist()


def read_metadata_header(docx_bytes):
    with zipfile.ZipFile(io.BytesIO(docx_bytes)) as docx_zip:
        for name in sorted(docx_zip.namelist()):
            if not name.startswith("word/header") or not name.endswith(".xml"):
                continue
            content = docx_zip.read(name).decode("utf-8")
            if "N° do Documento:" in content:
                return content
    return ""


def build_dossier_context():
    return {
        "job": {"id": "JOB-1", "kind": "project_dossier"},
        "project": {"id": "PRJ-01", "nome": "Projeto 1"},
        "defaults": {"projectId": "PRJ-01"},
        "renderModel": {
            "dossier": {
                "id": "DOS-1",
                "nome": "Dossie 1",
                "observacoes": "Observacoes",
                "scopeJson": {"includeFotos": True},
            },
            "sections": {
                "licencas": [{"id": "LO-1", "orgao": "IBAMA"}],
                "inspecoes": [],
                "erosoes": [],
                "entregas": [],
                "workspaces": [{"id": "RW-1", "nome": "Workspace 1", "status": "draft"}],
                "photos": [{
                    "id": "RPH-1",
                    "caption": "Foto 1",
                    "workspaceId": "RW-1",
                    "towerId": "T-01",
                    "includeInReport": True,
                    "mediaAssetId": "MED-PHOTO-1",
                }],
            },
        },
    }


def build_compound_context():
    return {
        "job": {"id": "JOB-2", "kind": "report_compound"},
        "project": None,
        "defaults": None,
        "renderModel": {
            "compound": {
                "id": "RC-1",
                "nome": "Composto 1",
                "sharedTextsJson": {"introducao": "Introducao global"},
                "orderJson": ["RW-2", "RW-1"],
                "workspaceIds": ["RW-1", "RW-2"],
            },
            "workspaces": [
                {
                    "workspace": {"id": "RW-2", "nome": "Workspace 2", "status": "draft"},
                    "project": {"id": "PRJ-02", "nome": "Projeto 2"},
                    "photos": [],
                },
                {
                    "workspace": {"id": "RW-1", "nome": "Workspace 1", "status": "draft"},
                    "project": {"id": "PRJ-01", "nome": "Projeto 1"},
                    "photos": [{
                        "id": "RPH-2",
                        "caption": "Foto 2",
                        "workspaceId": "RW-1",
                        "includeInReport": True,
                        "mediaAssetId": "MED-PHOTO-2",
                    }],
                },
            ],
        },
    }


def build_workspace_kmz_context():
    return {
        "job": {"id": "JOB-KMZ-1", "kind": "workspace_kmz", "workspaceId": "RW-1"},
        "project": {
            "id": "PRJ-01",
            "nome": "Projeto 1",
            "linhaFonteKml": "Linha Norte C1",
            "linhaCoordenadas": [
                {"latitude": -22.1, "longitude": -43.1},
                {"latitude": -22.2, "longitude": -43.2},
            ],
            "torresCoordenadas": [
                {"numero": "1", "latitude": -22.3, "longitude": -43.3},
            ],
        },
        "defaults": {"projectId": "PRJ-01"},
        "renderModel": {
            "workspaceKmz": {
                "token": "kmz-1",
                "workspaceId": "RW-1",
            },
            "workspace": {"id": "RW-1", "nome": "Workspace 1", "projectId": "PRJ-01"},
            "photos": [{
                "id": "RPH-KMZ-1",
                "caption": "Foto KMZ 1",
                "workspaceId": "RW-1",
                "towerId": "T-01",
                "includeInReport": True,
                "mediaAssetId": "MED-KMZ-PHOTO-1",
            }],
        },
    }


class StubClient:
    def __init__(self, configured=True, job=None, contexts=None, download_fail_media_ids=None, fail_upload=False, fail_create_media=False):
        self._configured = configured
        self.job = job
        self.contexts = contexts or {}
        self.download_fail_media_ids = set(download_fail_media_ids or [])
        self.fail_upload = fail_upload
        self.fail_create_media = fail_create_media
        self.completed = []
        self.failed = []
        self.created_media = []
        self.uploaded_media = []
        self.completed_media = []
        self.progress_reports = []

    def is_configured(self):
        return self._configured

    def claim_next_job(self):
        return self.job

    def get_job_context(self, job_id):
        return self.contexts.get(job_id)

    def download_media_content(self, media_id):
        if media_id in self.download_fail_media_ids:
            raise RuntimeError(f"media {media_id} indisponivel")
        return {
            "buffer": SAMPLE_PNG_BYTES,
            "contentType": "image/png",
        }

    def create_output_media(self, job_id, file_name, content_type, size_bytes=0, purpose="report_output_docx"):
        if self.fail_create_media:
            from worker.runtime import WorkerClientError
            raise WorkerClientError("API respondeu 415: Content-Type nao suportado")
        media = {
            "id": f"MED-{job_id}",
            "upload": {
                "href": "https://geomonitor-api.example.com/api/media/upload",
                "method": "PUT",
                "headers": {"Content-Type": content_type},
            },
        }
        self.created_media.append((job_id, file_name, content_type, size_bytes, purpose, media))
        return media

    def upload_media_binary(self, upload_descriptor, content, content_type):
        if self.fail_upload:
            raise RuntimeError("falha no upload final")
        self.uploaded_media.append((upload_descriptor, content, content_type))

    def complete_media_upload(self, media_id, stored_size_bytes=None, sha256=""):
        self.completed_media.append((media_id, stored_size_bytes, sha256))

    def report_job_progress(self, job_id, processed, total, phase="rendering"):
        self.progress_reports.append((job_id, processed, total, phase))

    def mark_complete(self, job_id, output_docx_media_id=None, output_kmz_media_id=None):
        self.completed.append((job_id, output_docx_media_id, output_kmz_media_id))

    def mark_failed(self, job_id, error_log):
        self.failed.append((job_id, error_log))


class FakeResponse:
    def __init__(self, status, payload=None):
        self.status = status
        self.payload = payload

    def read(self):
        if self.payload is None:
            return b""
        return json.dumps(self.payload).encode("utf-8")

    def __enter__(self):
        return self

    def __exit__(self, exc_type, exc, tb):
        return False


class WorkerRuntimeTests(unittest.TestCase):
    def test_run_once_returns_disabled_when_worker_is_not_configured(self):
        runtime = WorkerRuntime(client=StubClient(configured=False))

        result = runtime.run_once()

        self.assertEqual(result["status"], "disabled")
        self.assertEqual(runtime.describe()["lastResultStatus"], "disabled")

    def test_run_once_returns_idle_when_queue_is_empty(self):
        runtime = WorkerRuntime(client=StubClient(configured=True, job=None))

        result = runtime.run_once()

        self.assertEqual(result["status"], "idle")
        self.assertEqual(runtime.describe()["claimedCount"], 0)

    def test_run_once_completes_project_dossier_job(self):
        job = {"id": "JOB-1", "kind": "project_dossier"}
        client = StubClient(
            configured=True,
            job=job,
            contexts={"JOB-1": build_dossier_context()},
        )
        runtime = WorkerRuntime(client=client)

        result = runtime.run_once()

        self.assertEqual(result["status"], "completed")
        self.assertEqual(client.completed[0][0], "JOB-1")
        self.assertEqual(client.completed[0][1], "MED-JOB-1")
        self.assertEqual(runtime.describe()["completedCount"], 1)

        uploaded_docx = client.uploaded_media[0][1]
        document_text = read_docx_plain_text(uploaded_docx)
        header_xml = read_metadata_header(uploaded_docx)

        self.assertIn("Foto 1", document_text)
        self.assertIn("Torre T-01", document_text)
        self.assertIn("LT Projeto 1", header_xml)

    def test_run_once_completes_report_compound_job(self):
        job = {"id": "JOB-2", "kind": "report_compound"}
        client = StubClient(
            configured=True,
            job=job,
            contexts={"JOB-2": build_compound_context()},
        )
        runtime = WorkerRuntime(client=client)

        result = runtime.run_once()

        self.assertEqual(result["status"], "completed")
        self.assertEqual(client.completed[0][0], "JOB-2")
        self.assertEqual(client.completed[0][1], "MED-JOB-2")

        uploaded_docx = client.uploaded_media[0][1]
        document_text = read_docx_plain_text(uploaded_docx)
        self.assertIn("Introducao global", document_text)
        self.assertIn("ILUSTRA\u00c7\u00c3O FOTOGR\u00c1FICA", document_text)
        self.assertIn("Foto 1", document_text)
        self.assertIn("Foto 2", document_text)

    def test_apply_axia_formatting_compound_aligns_styles_and_margins(self):
        document, used_template, _ = create_document_from_template(
            "LT Test", "COD-001", "2026-04-15", "00"
        )
        self.assertTrue(used_template, "Template file must exist for this test to be meaningful")

        apply_axia_formatting_compound(document)

        sections = list(document.sections)
        self.assertGreater(len(sections), 1, "Template expected to have a back-cover section")

        # All content sections (everything except the back cover) get AXIA margins.
        # Margins round-trip through w:pgMar as twips, so Cm(4) lands ~180 EMU off
        # (2268 twips * 635 EMU = 1440180 vs Cm(4) = 1440000). Tolerate <0.01 cm.
        margin_tolerance = Cm(0.01)
        for section in sections[:-1]:
            self.assertAlmostEqual(section.top_margin, Cm(4), delta=margin_tolerance)
            self.assertAlmostEqual(section.right_margin, Cm(2), delta=margin_tolerance)
            self.assertAlmostEqual(section.left_margin, Cm(2), delta=margin_tolerance)
            self.assertAlmostEqual(section.bottom_margin, Cm(2), delta=margin_tolerance)

        # Back-cover section stays at its original zero-margin layout (quarta capa).
        back_cover = sections[-1]
        self.assertAlmostEqual(back_cover.top_margin, Cm(0), delta=margin_tolerance)
        self.assertAlmostEqual(back_cover.right_margin, Cm(0), delta=margin_tolerance)
        self.assertAlmostEqual(back_cover.left_margin, Cm(0), delta=margin_tolerance)
        self.assertAlmostEqual(back_cover.bottom_margin, Cm(0), delta=margin_tolerance)

        # Normal: DM Sans 11 (corpo reduzido + entrelinha apertada, manual secao 3).
        normal = document.styles["Normal"]
        self.assertEqual(normal.font.name, "DM Sans")
        self.assertEqual(normal.font.size, Pt(11))

        # NormalWeb (body paragraph style) — justified DM Sans 11, entrelinha
        # apertada: single line spacing, before 0 / after 6pt.
        body = document.styles["Normal (Web)"]
        self.assertEqual(body.font.name, "DM Sans")
        self.assertEqual(body.font.size, Pt(11))
        self.assertFalse(body.font.italic)
        self.assertEqual(body.paragraph_format.alignment, WD_ALIGN_PARAGRAPH.JUSTIFY)
        self.assertEqual(body.paragraph_format.space_before, Pt(0))
        self.assertEqual(body.paragraph_format.space_after, Pt(6))

        # Autospacing flags on NormalWeb must be cleared so space_before/after take effect.
        pPr = body.element.find("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pPr")
        self.assertIsNotNone(pPr)
        spacing_el = pPr.find("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}spacing")
        self.assertIsNotNone(spacing_el)
        self.assertNotIn(
            "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}beforeAutospacing",
            spacing_el.attrib,
        )
        self.assertNotIn(
            "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}afterAutospacing",
            spacing_el.attrib,
        )

        # Heading do template: troca o typeface e FIXA o tamanho de secao (14pt),
        # senao o Ttulo1 (sem sz proprio) herdaria o corpo de 11pt e encolheria.
        heading = document.styles["Ttulo1"]
        self.assertEqual(heading.font.name, "DM Sans")
        self.assertEqual(heading.font.size, Pt(14))
        # Cor institucional Azul-marinho AXIA fixada em runtime (manual secao 4).
        self.assertEqual(heading.font.color.rgb, RGBColor.from_string("0A003C"))

        # Caption (photo legend): preserva estilo original do template (italic, cor tema).
        caption = document.styles["caption"]
        self.assertTrue(caption.font.italic)
        self.assertIsNone(caption.font.bold)

    def test_apply_body_font_on_template_styles(self):
        # Caminho do dossie: so a fonte de corpo e aplicada, sem margens/spacing.
        document, used_template, _ = create_document_from_template(
            "LT Test", "COD-001", "2026-04-15", "00"
        )
        self.assertTrue(used_template, "Template file must exist for this test to be meaningful")

        apply_body_font(document)

        self.assertEqual(document.styles["Normal"].font.name, "DM Sans")
        self.assertEqual(document.styles["Normal"].font.size, Pt(11))
        self.assertEqual(document.styles["Normal (Web)"].font.name, "DM Sans")
        self.assertEqual(document.styles["Normal (Web)"].font.size, Pt(11))
        self.assertEqual(document.styles["Ttulo1"].font.name, "DM Sans")
        # Tamanho de secao fixado para nao flutuar com o corpo (basedOn Normal).
        self.assertEqual(document.styles["Ttulo1"].font.size, Pt(14))

    def test_run_once_completes_workspace_kmz_job(self):
        job = {"id": "JOB-KMZ-1", "kind": "workspace_kmz"}
        client = StubClient(
            configured=True,
            job=job,
            contexts={"JOB-KMZ-1": build_workspace_kmz_context()},
        )
        runtime = WorkerRuntime(client=client)

        result = runtime.run_once()

        self.assertEqual(result["status"], "completed")
        self.assertEqual(client.completed[0][0], "JOB-KMZ-1")
        self.assertEqual(client.completed[0][2], "MED-JOB-KMZ-1")

        uploaded_kmz = client.uploaded_media[0][1]
        kml_text = read_kmz_entry(uploaded_kmz, "doc.kml").decode("utf-8")

        self.assertIn("Projeto 1 - Workspace 1", kml_text)
        self.assertIn("Linha Norte C1", kml_text)
        self.assertIn("Foto KMZ 1", kml_text)
        # Nome de arquivo estavel por photoId + ExtendedData garantem o round-trip
        # de organizacao (re-import casa a foto e atribui torre sem novo upload).
        self.assertIn("files/RPH-KMZ-1.jpg", kml_text)
        self.assertIn('<Data name="photoId"><value>RPH-KMZ-1</value></Data>', kml_text)
        self.assertIn('<Data name="mediaAssetId"><value>MED-KMZ-PHOTO-1</value></Data>', kml_text)
        self.assertEqual(read_kmz_entry(uploaded_kmz, "files/RPH-KMZ-1.jpg"), SAMPLE_PNG_BYTES)

        # Progresso reportado durante a renderizacao (estilo barra de upload):
        # ping inicial (0/total) e ping final (total/total).
        self.assertTrue(client.progress_reports)
        self.assertEqual(client.progress_reports[0], ("JOB-KMZ-1", 0, 1, "rendering"))
        self.assertEqual(client.progress_reports[-1], ("JOB-KMZ-1", 1, 1, "rendering"))

    def test_download_media_content_retries_then_succeeds(self):
        import os
        from urllib import error as urllib_error

        calls = {"n": 0}

        class BinaryResponse:
            status = 200
            headers = {"Content-Type": "image/png"}

            def read(self):
                return SAMPLE_PNG_BYTES

            def __enter__(self):
                return self

            def __exit__(self, *args):
                return False

        def fake_urlopen(req, timeout=None):
            calls["n"] += 1
            if calls["n"] < 3:
                raise urllib_error.URLError("instabilidade transitoria")
            return BinaryResponse()

        os.environ["WORKER_DOWNLOAD_MAX_ATTEMPTS"] = "3"
        os.environ["WORKER_DOWNLOAD_BACKOFF_BASE"] = "0"
        try:
            client = WorkerClient(base_url="https://api.example.com", token="t", urlopen=fake_urlopen)
            result = client.download_media_content("MED-1")
        finally:
            os.environ.pop("WORKER_DOWNLOAD_MAX_ATTEMPTS", None)
            os.environ.pop("WORKER_DOWNLOAD_BACKOFF_BASE", None)

        self.assertEqual(calls["n"], 3)
        self.assertEqual(result["buffer"], SAMPLE_PNG_BYTES)

    def test_download_media_content_raises_after_exhausting_retries(self):
        import os
        from urllib import error as urllib_error

        calls = {"n": 0}

        def fake_urlopen(req, timeout=None):
            calls["n"] += 1
            raise urllib_error.URLError("sempre indisponivel")

        os.environ["WORKER_DOWNLOAD_MAX_ATTEMPTS"] = "2"
        os.environ["WORKER_DOWNLOAD_BACKOFF_BASE"] = "0"
        try:
            client = WorkerClient(base_url="https://api.example.com", token="t", urlopen=fake_urlopen)
            with self.assertRaises(Exception):
                client.download_media_content("MED-1")
        finally:
            os.environ.pop("WORKER_DOWNLOAD_MAX_ATTEMPTS", None)
            os.environ.pop("WORKER_DOWNLOAD_BACKOFF_BASE", None)

        self.assertEqual(calls["n"], 2)

    def test_run_once_fails_workspace_kmz_when_all_downloads_fail(self):
        job = {"id": "JOB-KMZ-FAIL", "kind": "workspace_kmz"}
        client = StubClient(
            configured=True,
            job=job,
            contexts={"JOB-KMZ-FAIL": build_workspace_kmz_context()},
            download_fail_media_ids={"MED-KMZ-PHOTO-1"},
        )
        runtime = WorkerRuntime(client=client)

        result = runtime.run_once()

        self.assertEqual(result["status"], "failed")
        self.assertEqual(client.failed[0][0], "JOB-KMZ-FAIL")
        # Mensagem distingue "todas falharam no download" de "sem fotos com media".
        self.assertIn("Todas as fotos falharam", client.failed[0][1])

    def test_run_once_completes_workspace_kmz_with_partial_download_failure(self):
        context = build_workspace_kmz_context()
        context["renderModel"]["photos"].append({
            "id": "RPH-KMZ-2",
            "caption": "Foto KMZ 2",
            "workspaceId": "RW-1",
            "towerId": "T-02",
            "includeInReport": True,
            "mediaAssetId": "MED-KMZ-PHOTO-2",
        })
        job = {"id": "JOB-KMZ-PARTIAL", "kind": "workspace_kmz"}
        client = StubClient(
            configured=True,
            job=job,
            contexts={"JOB-KMZ-PARTIAL": context},
            download_fail_media_ids={"MED-KMZ-PHOTO-2"},
        )
        runtime = WorkerRuntime(client=client)

        result = runtime.run_once()

        self.assertEqual(result["status"], "completed")
        uploaded_kmz = client.uploaded_media[0][1]
        names = read_kmz_namelist(uploaded_kmz)
        self.assertIn("files/RPH-KMZ-1.jpg", names)
        self.assertNotIn("files/RPH-KMZ-2.jpg", names)
        readme = read_kmz_entry(uploaded_kmz, "README.txt").decode("utf-8")
        self.assertIn("RPH-KMZ-2", readme)

    def test_run_once_workspace_kmz_uses_exif_when_no_tower(self):
        jpeg = build_minimal_jpeg_with_gps()  # -22.5, -43.25
        context = build_workspace_kmz_context()
        context["renderModel"]["photos"] = [{
            "id": "RPH-EXIF-1",
            "caption": "Foto sem torre",
            "workspaceId": "RW-1",
            "towerId": "",
            "gpsLat": 0,
            "gpsLon": 0,
            "includeInReport": True,
            "mediaAssetId": "MED-EXIF-1",
        }]

        class GpsStub(StubClient):
            def download_media_content(self, media_id):
                return {"buffer": jpeg, "contentType": "image/jpeg"}

        job = {"id": "JOB-EXIF", "kind": "workspace_kmz"}
        client = GpsStub(configured=True, job=job, contexts={"JOB-EXIF": context})

        result = WorkerRuntime(client=client).run_once()

        self.assertEqual(result["status"], "completed")
        kml = read_kmz_entry(client.uploaded_media[0][1], "doc.kml").decode("utf-8")
        self.assertIn('<Style id="photo-marker">', kml)
        self.assertIn("<styleUrl>#photo-marker</styleUrl>", kml)
        # Marcador na posicao do EXIF (lon,lat), nao em 0,0.
        self.assertIn("-43.25,-22.5", kml)

    def test_run_once_workspace_kmz_no_marker_for_zero_coords_without_exif(self):
        # gps placeholder (0,0), sem torre e sem EXIF (StubClient devolve PNG) ->
        # nenhum marcador (nao plota null-island), e aviso no README.
        context = build_workspace_kmz_context()
        context["renderModel"]["photos"] = [{
            "id": "RPH-ZERO-1",
            "caption": "Foto zero",
            "workspaceId": "RW-1",
            "towerId": "",
            "gpsLat": 0,
            "gpsLon": 0,
            "includeInReport": True,
            "mediaAssetId": "MED-ZERO-1",
        }]
        job = {"id": "JOB-ZERO", "kind": "workspace_kmz"}
        client = StubClient(configured=True, job=job, contexts={"JOB-ZERO": context})

        result = WorkerRuntime(client=client).run_once()

        self.assertEqual(result["status"], "completed")
        kmz = client.uploaded_media[0][1]
        kml = read_kmz_entry(kmz, "doc.kml").decode("utf-8")
        self.assertIn('<Style id="photo-marker">', kml)
        self.assertNotIn("<styleUrl>#photo-marker</styleUrl>", kml)
        readme = read_kmz_entry(kmz, "README.txt").decode("utf-8")
        self.assertIn("ficou sem coordenadas de mapa", readme)

    def test_run_once_marks_job_as_failed_when_upload_breaks(self):
        job = {"id": "JOB-UPLOAD-FAIL", "kind": "project_dossier"}
        client = StubClient(
            configured=True,
            job=job,
            contexts={"JOB-UPLOAD-FAIL": build_dossier_context()},
            fail_upload=True,
        )
        runtime = WorkerRuntime(client=client)

        result = runtime.run_once()

        self.assertEqual(result["status"], "failed")
        self.assertEqual(client.failed[0][0], "JOB-UPLOAD-FAIL")
        self.assertIn("falha no upload final", client.failed[0][1])

    def test_run_once_marks_failed_on_worker_client_error_during_processing(self):
        """Regression: WorkerClientError from API (e.g. 415) must mark job failed."""
        job = {"id": "JOB-415", "kind": "project_dossier"}
        client = StubClient(
            configured=True,
            job=job,
            contexts={"JOB-415": build_dossier_context()},
            fail_create_media=True,
        )
        runtime = WorkerRuntime(client=client)
        result = runtime.run_once()

        self.assertEqual(result["status"], "failed")
        self.assertEqual(len(client.failed), 1)
        self.assertEqual(client.failed[0][0], "JOB-415")
        self.assertIn("415", client.failed[0][1])

    def test_run_once_ignores_individual_photo_download_failure(self):
        context = build_dossier_context()
        context["renderModel"]["sections"]["photos"].append({
            "id": "RPH-2",
            "caption": "Foto com falha",
            "workspaceId": "RW-1",
            "includeInReport": True,
            "mediaAssetId": "MED-PHOTO-FAIL",
        })
        job = {"id": "JOB-PHOTO-FAIL", "kind": "project_dossier"}
        client = StubClient(
            configured=True,
            job=job,
            contexts={"JOB-PHOTO-FAIL": context},
            download_fail_media_ids={"MED-PHOTO-FAIL"},
        )
        runtime = WorkerRuntime(client=client)

        result = runtime.run_once()

        self.assertEqual(result["status"], "completed")
        self.assertEqual(client.completed[0][0], "JOB-PHOTO-FAIL")

    # ------------------------------------------------------------------
    # Regression tests for docx_renderer findings (plan: precious-churning-lecun)
    # ------------------------------------------------------------------

    def _run_dossier(self, context, job_id="JOB-RENDER"):
        job = {"id": job_id, "kind": "project_dossier"}
        client = StubClient(configured=True, job=job, contexts={job_id: context})
        runtime = WorkerRuntime(client=client)
        result = runtime.run_once()
        self.assertEqual(result["status"], "completed")
        return client.uploaded_media[0][1]

    def _run_compound(self, context, job_id="JOB-COMP"):
        job = {"id": job_id, "kind": "report_compound"}
        client = StubClient(configured=True, job=job, contexts={job_id: context})
        runtime = WorkerRuntime(client=client)
        result = runtime.run_once()
        self.assertEqual(result["status"], "completed")
        return client.uploaded_media[0][1]

    def test_project_dossier_preserves_content_section_break(self):
        """T1 — clear_template_body must hand back a content sectPr.

        When the inline section break is preserved, the body ends up with
        at least two <w:sectPr> elements (main content + quarta capa).
        """
        docx = self._run_dossier(build_dossier_context(), job_id="JOB-T1")
        body_xml = read_docx_entry(docx, "word/document.xml")
        sectpr_count = body_xml.count("<w:sectPr")
        self.assertGreaterEqual(
            sectpr_count,
            2,
            msg=f"expected >=2 sectPr (content + quarta capa), got {sectpr_count}",
        )

    def test_report_compound_cover_custom_program_and_standard_report_line(self):
        """titulo_programa custom vai na linha 2; linha 3 vira a canonica.

        A capa do template tem tres linhas: LT, subtitulo do Programa e a linha
        do relatorio. Com titulo_programa custom, a linha 2 recebe o custom; a
        linha 3 e SEMPRE padronizada para COVER_REPORT_LINE (modelo OOSEMB.RT.075),
        substituindo o texto longo do template.
        """
        from worker.docx_renderer import COVER_REPORT_LINE
        context = build_compound_context()
        context["renderModel"]["compound"]["sharedTextsJson"] = {
            "introducao": "Introducao global",
            "nome_lt": "Projeto Alfa",
            "titulo_programa": "Programa Personalizado",
        }
        docx = self._run_compound(context, job_id="JOB-T2")
        plain = read_docx_plain_text(docx)
        self.assertIn("Programa Personalizado", plain)
        self.assertIn(COVER_REPORT_LINE, plain)
        # O texto longo antigo da linha 3 foi padronizado (nao aparece mais).
        self.assertNotIn("Inspeção Técnica do Programa de Monitoramento", plain)

    def test_report_compound_cover_standardized(self):
        """Capa padronizada: linhas 2/3 canonicas, LT em CAIXA ALTA, centralizada.

        Sem titulo_programa, a linha 2 usa o padrao institucional COVER_PROGRAM_LINE.
        A LT vai em CAIXA ALTA SO na capa (o cabecalho mantem caixa mista). As
        linhas da capa ficam centralizadas e com fonte explicita (DM Sans no
        preset padrao).
        """
        from worker.docx_renderer import COVER_PROGRAM_LINE, COVER_REPORT_LINE
        context = build_compound_context()
        context["renderModel"]["compound"]["sharedTextsJson"] = {
            "introducao": "Introducao global",
            "nome_lt": "Projeto Alfa",
        }
        docx = self._run_compound(context, job_id="JOB-COVER-STD")
        plain = read_docx_plain_text(docx)
        self.assertIn(COVER_PROGRAM_LINE, plain)
        self.assertIn(COVER_REPORT_LINE, plain)
        self.assertNotIn("Inspeção Técnica do Programa de Monitoramento", plain)

        # LT em CAIXA ALTA na capa, mas caixa mista no cabecalho.
        self.assertIn("LT PROJETO ALFA", plain)
        header_xml = read_metadata_header(docx)
        self.assertIn("LT Projeto Alfa", header_xml)
        self.assertNotIn("LT PROJETO ALFA", header_xml)

        # Linhas da capa (textboxes) centralizadas e em DM Sans explicito.
        doc_xml = read_docx_entry(docx, "word/document.xml")
        txbx = "".join(re.findall(r"<w:txbxContent>.*?</w:txbxContent>", doc_xml, re.S))
        self.assertTrue(txbx, "textbox da capa nao encontrado")
        self.assertIn('<w:jc w:val="center"', txbx)
        self.assertIn('w:ascii="DM Sans"', txbx)

    def test_report_compound_output_keeps_axia_brand(self):
        """O re-skin AXIA sobrevive ao render: tema, fontes e heading institucional.

        Garante que o relatorio gerado (Monitoramento de Processo Erosivo)
        herda a paleta/fontes do template e o realce de heading aplicado em
        runtime — sem cores do tema Office vazando para a entrega.
        """
        docx = self._run_compound(build_compound_context(), job_id="JOB-BRAND")

        theme = read_docx_entry(docx, "word/theme/theme1.xml")
        self.assertIn('<a:accent1><a:srgbClr val="0000FF"', theme)
        self.assertIn('<a:dk2><a:srgbClr val="0A003C"', theme)
        for office_color in ("4472C4", "ED7D31", "FFC000"):
            self.assertNotIn(office_color, theme)

        font_table = read_docx_entry(docx, "word/fontTable.xml")
        self.assertIn('<w:font w:name="DM Sans">', font_table)
        self.assertIn('w:altName w:val="Arial"', font_table)

        # Capa migrou de Verdana (fora da marca) para DM Sans.
        self.assertNotIn('"Verdana"', read_docx_entry(docx, "word/document.xml"))

        # Heading institucional Ttulo1 fixado em Azul-marinho AXIA.
        styles = read_docx_entry(docx, "word/styles.xml")
        ttulo1 = re.search(
            r'<w:style [^>]*w:styleId="Ttulo1".*?</w:style>', styles, re.S
        )
        self.assertIsNotNone(ttulo1)
        self.assertIn('<w:color w:val="0A003C"', ttulo1.group(0))

    def test_add_photo_entry_keeps_photo_with_caption(self):
        """Foto colada a legenda: keepNext+keepLines na imagem, keepLines na legenda."""
        from docx import Document
        from worker.docx_renderer import add_photo_entry

        document = Document()

        def loader(_media_id):
            return {"buffer": SAMPLE_PNG_BYTES, "contentType": "image/png"}

        add_photo_entry(
            document,
            {"mediaAssetId": "MED-X", "caption": "Erosao na base da torre"},
            loader,
            1,
        )
        # paragrafos finais: [..., imagem, legenda]
        caption_p = document.paragraphs[-1]
        image_p = document.paragraphs[-2]
        self.assertTrue(image_p.paragraph_format.keep_with_next)
        self.assertTrue(image_p.paragraph_format.keep_together)
        self.assertTrue(caption_p.paragraph_format.keep_together)

    def test_picture_size_kwargs_caps_portrait_height(self):
        """Retrato alto trava pela altura (16cm); paisagem/quadrado/lixo -> largura 15cm."""
        from worker.docx_renderer import (
            MAX_IMAGE_HEIGHT_CM,
            MAX_IMAGE_WIDTH_CM,
            _picture_size_kwargs,
        )

        self.assertEqual(
            _picture_size_kwargs(build_png(1500, 2000)),  # retrato 3:4
            {"height": Cm(MAX_IMAGE_HEIGHT_CM)},
        )
        self.assertEqual(
            _picture_size_kwargs(build_png(1080, 1920)),  # retrato 9:16
            {"height": Cm(MAX_IMAGE_HEIGHT_CM)},
        )
        self.assertEqual(
            _picture_size_kwargs(build_png(2000, 1500)),  # paisagem 4:3
            {"width": Cm(MAX_IMAGE_WIDTH_CM)},
        )
        # Quadrado 1x1 (15x15 <= 16) e bytes invalidos caem no default de largura.
        self.assertEqual(_picture_size_kwargs(SAMPLE_PNG_BYTES), {"width": Cm(MAX_IMAGE_WIDTH_CM)})
        self.assertEqual(_picture_size_kwargs(b"nao e imagem"), {"width": Cm(MAX_IMAGE_WIDTH_CM)})

    def test_add_photo_entry_caps_vertical_photo_height(self):
        """Foto retrato inserida fica com altura ~16cm e largura < 15cm; paisagem mantem 15cm."""
        from docx import Document
        from worker.docx_renderer import (
            MAX_IMAGE_HEIGHT_CM,
            MAX_IMAGE_WIDTH_CM,
            add_photo_entry,
        )

        document = Document()
        add_photo_entry(
            document,
            {"mediaAssetId": "MED-V", "caption": "Vertical"},
            lambda _m: {"buffer": build_png(1500, 2000), "contentType": "image/png"},
            1,
        )
        shape = document.inline_shapes[-1]
        self.assertAlmostEqual(shape.height, Cm(MAX_IMAGE_HEIGHT_CM), delta=Cm(0.05))
        self.assertLess(shape.width, Cm(MAX_IMAGE_WIDTH_CM))

        document2 = Document()
        add_photo_entry(
            document2,
            {"mediaAssetId": "MED-H", "caption": "Horizontal"},
            lambda _m: {"buffer": build_png(2000, 1500), "contentType": "image/png"},
            1,
        )
        shape2 = document2.inline_shapes[-1]
        self.assertAlmostEqual(shape2.width, Cm(MAX_IMAGE_WIDTH_CM), delta=Cm(0.05))

    def test_report_compound_header_size_unaffected_by_body_shrink(self):
        """Reducao do corpo (11pt) nao vaza para o cabecalho.

        Os estilos Cabealho/Rodap herdam o tamanho do Normal, mas todos os runs
        de texto do header2 tem <w:sz> explicito (7/8/9pt = 14/16/18 half-pt),
        entao a mudanca do corpo nao altera o cabecalho.
        """
        docx = self._run_compound(build_compound_context(), job_id="JOB-HDR-SZ")
        header_xml = read_metadata_header(docx)
        self.assertTrue(header_xml, "header com 'N° do Documento:' nao encontrado")
        sizes = set(re.findall(r'<w:sz w:val="(\d+)"', header_xml))
        # Tamanhos do cabecalho permanecem 7/8/9pt...
        self.assertTrue(
            {"14", "16", "18"} & sizes,
            f"esperava sz 14/16/18 (7/8/9pt) no cabecalho, achei {sizes}",
        )
        # ...e o corpo de 11pt (sz 22) nao vazou via heranca do Normal.
        self.assertNotIn("22", sizes)

    def test_report_compound_heading_hierarchy_sizes_and_navy(self):
        """Secao 14pt (estilo Ttulo1) e subtitulo de torre 12pt navy (override direto)."""
        ctx = build_compound_context()
        ctx["renderModel"]["compound"]["sharedTextsJson"] = {"introducao": "Intro"}
        ctx["renderModel"]["workspaces"] = [
            {
                "workspace": {"id": "RW-1", "nome": "Workspace 1", "status": "draft"},
                "project": {"id": "PRJ-01", "nome": "Projeto 1"},
                "photos": [
                    {"id": "RPH-1", "caption": "X1", "towerId": "T-01",
                     "includeInReport": True, "mediaAssetId": "MED-A"},
                ],
            },
        ]
        ctx["renderModel"]["compound"]["orderJson"] = ["RW-1"]
        ctx["renderModel"]["compound"]["workspaceIds"] = ["RW-1"]
        docx = self._run_compound(ctx, job_id="JOB-HEAD-HIER")

        # Secao: estilo Ttulo1 fixado em 14pt (sz 28 half-pt).
        styles_xml = read_docx_entry(docx, "word/styles.xml")
        ttulo1 = re.search(
            r'<w:style [^>]*w:styleId="Ttulo1".*?</w:style>', styles_xml, re.S
        ).group(0)
        self.assertIn('<w:sz w:val="28"', ttulo1)

        # Subtitulo de torre (ilvl1): run com 12pt (sz 24) e navy 0A003C.
        doc_xml = read_docx_entry(docx, "word/document.xml")
        paras = re.findall(r"<w:p[ >].*?</w:p>", doc_xml, re.S)
        torre = [p for p in paras if "Torre T-01" in p]
        self.assertTrue(torre, "paragrafo 'Torre T-01' nao encontrado")
        self.assertIn('<w:sz w:val="24"', torre[0])
        self.assertIn('<w:color w:val="0A003C"', torre[0])

    def _normal_rfonts(self, docx):
        styles = read_docx_entry(docx, "word/styles.xml")
        normal = re.search(r'<w:style [^>]*w:styleId="Normal".*?</w:style>', styles, re.S).group(0)
        return re.search(r"<w:rFonts[^>]*>", normal).group(0)

    def test_report_compound_default_preset_embeds_dm_sans(self):
        """Preset padrao (axia): corpo em DM Sans e a fonte vai embutida no DOCX."""
        docx = self._run_compound(build_compound_context(), job_id="JOB-DEF")
        self.assertIn('w:ascii="DM Sans"', self._normal_rfonts(docx))
        dm = re.search(
            r'<w:font w:name="DM Sans">.*?</w:font>',
            read_docx_entry(docx, "word/fontTable.xml"),
            re.S,
        ).group(0)
        self.assertIn("w:embedRegular", dm)

    def test_report_compound_arial_preset_uses_arial(self):
        """Preset axia-arial: estilos base e runs da capa caem para Arial."""
        context = build_compound_context()
        context["renderModel"]["compound"]["sharedTextsJson"] = {
            "introducao": "Intro",
            "reportStyle": "axia-arial",
        }
        docx = self._run_compound(context, job_id="JOB-ARIAL")
        self.assertIn('w:ascii="Arial"', self._normal_rfonts(docx))
        # Nenhum run direto sobra em DM Sans (capa incluida).
        self.assertNotIn('w:ascii="DM Sans"', read_docx_entry(docx, "word/document.xml"))

    def test_project_dossier_header_has_custom_document_code(self):
        """T3 — document_code flows from context, hard-coded placeholder gone."""
        context = build_dossier_context()
        context["renderModel"]["dossier"]["codigo_documento"] = "TESTE.RT.999.2030"
        context["renderModel"]["dossier"]["revisao"] = "07"
        context["job"]["updatedAt"] = "2030-06-15T12:00:00Z"
        docx = self._run_dossier(context, job_id="JOB-T3")
        header_xml = read_metadata_header(docx)
        self.assertIn("TESTE.RT.999.2030", header_xml)
        self.assertNotIn("OOSEMB.RT.023.2026", header_xml)
        self.assertIn("15/06/2030", header_xml)
        # The revision cell is a separate run; check inside <w:t> to avoid
        # matching the font size attribute "07".
        self.assertTrue(
            re.search(r"<w:t[^>]*>07</w:t>", header_xml),
            msg="expected revision '07' inside a <w:t> element of the header",
        )

    def test_format_emission_date_normalizes_to_utc(self):
        """T4 — format_emission_date must normalize timezones to UTC."""
        from worker.docx_renderer import format_emission_date
        self.assertEqual(
            format_emission_date("2026-01-01T02:30:00+00:00"),
            "01/01/2026",
        )
        # 23:30 in -05:00 == 04:30 UTC next day
        self.assertEqual(
            format_emission_date("2025-12-31T23:30:00-05:00"),
            "01/01/2026",
        )

    def test_project_dossier_renders_without_template(self):
        """T5 — fallback path must work when the template file is absent."""
        import unittest.mock as mock
        import worker.docx_renderer as renderer

        real_exists = renderer.os.path.exists

        def fake_exists(path):
            if path == renderer.TEMPLATE_PATH:
                return False
            return real_exists(path)

        with mock.patch.object(renderer.os.path, "exists", side_effect=fake_exists):
            docx = self._run_dossier(build_dossier_context(), job_id="JOB-T5")
        plain = read_docx_plain_text(docx)
        self.assertIn("Resumo do Empreendimento", plain)
        # Photos section still rendered even without the Legenda style.
        self.assertIn("Foto 1", plain)

    def test_project_dossier_header_with_non_lt_project_name(self):
        """T6 — non-LT project names are prefixed and propagated."""
        context = build_dossier_context()
        context["project"]["nome"] = "SE Substacao Norte"
        docx = self._run_dossier(context, job_id="JOB-T6")
        header_xml = read_metadata_header(docx)
        self.assertIn("LT SE Substacao Norte", header_xml)

    def test_report_compound_signature_block_no_blank_page(self):
        """T7 — the signature block adds at most one page break (no blank page).

        The institutional template carries pre-existing page breaks inside
        its quarta-capa textboxes, so we compare the delta between a run
        WITH signatures and a run WITHOUT, instead of an absolute count.
        """
        def _page_break_count(context):
            docx = self._run_compound(context, job_id="JOB-T7")
            xml = read_docx_entry(docx, "word/document.xml")
            return len(re.findall(r'<w:br[^>]*w:type="page"', xml))

        baseline_ctx = build_compound_context()
        baseline_ctx["renderModel"]["compound"]["sharedTextsJson"] = {
            "introducao": "Introducao global",
        }
        baseline = _page_break_count(baseline_ctx)

        signed_ctx = build_compound_context()
        signed_ctx["renderModel"]["compound"]["sharedTextsJson"] = {
            "introducao": "Introducao global",
            "elaboradores": [{"nome": "Fulano"}],
            "revisores": [],
        }
        with_signatures = _page_break_count(signed_ctx)

        delta = with_signatures - baseline
        self.assertLessEqual(
            delta,
            1,
            msg=(
                f"signature block added {delta} page breaks "
                f"(baseline={baseline}, with_signatures={with_signatures})"
            ),
        )

    def test_report_compound_anexo_fichas_appended_after_signatures(self):
        """Anexo de fichas de erosao simplificada aparece apos assinaturas."""
        ctx = build_compound_context()
        ctx["renderModel"]["compound"]["sharedTextsJson"] = {
            "introducao": "Introducao global",
            "elaboradores": [{"nome": "Fulano", "profissao": "Geotecnico", "registro": "123"}],
            "revisores": [],
        }
        ctx["renderModel"]["compound"]["anexoFichas"] = {
            "projectName": "Projeto Alfa",
            "erosions": [
                {
                    "id": "EROS-1",
                    "torreRef": "T-01",
                    "criticalityCode": "C2",
                    "locationCoordinates": {"utmEasting": "500000", "utmNorthing": "7000000"},
                },
            ],
        }
        docx = self._run_compound(ctx, job_id="JOB-ANEXO")
        plain = read_docx_plain_text(docx)
        idx_sig = plain.find("Fulano")
        idx_anexo = plain.find("ANEXO")
        self.assertGreaterEqual(idx_sig, 0, "assinatura deveria estar no DOCX")
        self.assertGreaterEqual(idx_anexo, 0, "heading ANEXO deveria estar no DOCX")
        self.assertGreater(idx_anexo, idx_sig, "ANEXO deve vir apos as assinaturas")
        self.assertIn("FICHAS DE EROS", plain)

    def test_report_compound_without_anexo_fichas(self):
        """Sem anexoFichas, o heading ANEXO nao deve aparecer."""
        ctx = build_compound_context()
        ctx["renderModel"]["compound"]["sharedTextsJson"] = {"introducao": "Introducao"}
        docx = self._run_compound(ctx, job_id="JOB-NO-ANEXO")
        plain = read_docx_plain_text(docx)
        self.assertNotIn("ANEXO - FICHAS DE EROS", plain)

    def test_report_compound_tower_grouping_independent_of_input_order(self):
        # Regressao BOM DESPACHO 26/05/2026: agrupamento por torre era acoplado
        # ao sort_mode upstream ("tower*" => grouping). Quando o usuario escolhia
        # ordem manual/data/legenda, os subtopicos "Torre X" sumiam. Hoje
        # group_by_tower=True e fixo em render_report_compound_docx; este teste
        # protege contra re-acoplamento.
        ctx = build_compound_context()
        ctx["renderModel"]["compound"]["sharedTextsJson"] = {"introducao": "Intro"}
        ctx["renderModel"]["workspaces"] = [
            {
                "workspace": {"id": "RW-1", "nome": "Workspace 1", "status": "draft"},
                "project": {"id": "PRJ-01", "nome": "Projeto 1"},
                "photos": [
                    {"id": "RPH-1", "caption": "X1", "towerId": "T-02",
                     "includeInReport": True, "mediaAssetId": "MED-A"},
                    {"id": "RPH-2", "caption": "X2", "towerId": "T-01",
                     "includeInReport": True, "mediaAssetId": "MED-B"},
                    {"id": "RPH-3", "caption": "X3", "towerId": "T-02",
                     "includeInReport": True, "mediaAssetId": "MED-C"},
                    {"id": "RPH-4", "caption": "X4", "towerId": "T-01",
                     "includeInReport": True, "mediaAssetId": "MED-D"},
                ],
            },
        ]
        ctx["renderModel"]["compound"]["orderJson"] = ["RW-1"]
        ctx["renderModel"]["compound"]["workspaceIds"] = ["RW-1"]
        docx = self._run_compound(ctx, job_id="JOB-TOWER-GROUP")
        plain = read_docx_plain_text(docx)
        self.assertIn("Torre T-01", plain, "subtopico 'Torre T-01' deve existir mesmo com fotos intercaladas")
        self.assertIn("Torre T-02", plain, "subtopico 'Torre T-02' deve existir mesmo com fotos intercaladas")
        # As duas torres devem aparecer na ordem de primeira ocorrencia das fotos:
        # T-02 vem antes de T-01 no input -> esperado no output tambem.
        idx_t02 = plain.find("Torre T-02")
        idx_t01 = plain.find("Torre T-01")
        self.assertGreaterEqual(idx_t02, 0)
        self.assertGreaterEqual(idx_t01, 0)
        self.assertLess(idx_t02, idx_t01, "ordem dos grupos segue primeira ocorrencia no input")

    def test_report_compound_photos_without_tower_id_fall_into_fallback_group(self):
        # Quando nenhuma foto tem towerId, todas caem em "Fotos sem agrupamento"
        # — caso de KMZ sem pastas por torre. O subtopico de torre nao deve
        # aparecer, mas o agrupamento (heading nivel 1) ainda eh aplicado.
        ctx = build_compound_context()
        ctx["renderModel"]["compound"]["sharedTextsJson"] = {"introducao": "Intro"}
        ctx["renderModel"]["workspaces"] = [
            {
                "workspace": {"id": "RW-1", "nome": "Workspace 1", "status": "draft"},
                "project": {"id": "PRJ-01", "nome": "Projeto 1"},
                "photos": [
                    {"id": "RPH-1", "caption": "Y1",
                     "includeInReport": True, "mediaAssetId": "MED-X"},
                    {"id": "RPH-2", "caption": "Y2",
                     "includeInReport": True, "mediaAssetId": "MED-Y"},
                ],
            },
        ]
        ctx["renderModel"]["compound"]["orderJson"] = ["RW-1"]
        ctx["renderModel"]["compound"]["workspaceIds"] = ["RW-1"]
        docx = self._run_compound(ctx, job_id="JOB-NO-TOWER")
        plain = read_docx_plain_text(docx)
        # O heading nivel 1 de fallback aparece...
        self.assertIn("Fotos sem agrupamento", plain)
        # ...e vem depois da secao de ilustracao, junto com as fotos.
        idx_section = plain.find("ILUSTRA")
        idx_fallback = plain.find("Fotos sem agrupamento")
        idx_foto1 = plain.find("Foto 1 - Y1")
        idx_foto2 = plain.find("Foto 2 - Y2")
        self.assertGreater(idx_fallback, idx_section)
        self.assertGreater(idx_foto1, idx_fallback)
        self.assertGreater(idx_foto2, idx_foto1)

    def test_append_fichas_helper_adds_one_table_per_erosion(self):
        """Helper gera uma tabela por erosao, respeitando o template."""
        from docx import Document
        from worker.ficha_cadastro_renderer import append_fichas_cadastro_to_document

        document = Document()
        document.add_paragraph("Assinaturas ficticias")
        tables_before = len(document.tables)
        erosions = [
            {"id": "EROS-A", "torreRef": "T-02", "criticalityCode": "C3"},
            {"id": "EROS-B", "torreRef": "T-05", "criticalityCode": "C1"},
            {"id": "EROS-C", "torreRef": "T-09", "criticalityCode": "C4"},
        ]
        append_fichas_cadastro_to_document(document, erosions, "Projeto X")
        self.assertEqual(len(document.tables) - tables_before, 3)

    def test_append_fichas_helper_handles_empty_list(self):
        """Lista vazia e no-op (nao adiciona nada ao documento)."""
        from docx import Document
        from worker.ficha_cadastro_renderer import append_fichas_cadastro_to_document

        document = Document()
        body_len_before = len(document._element.body.getchildren())
        append_fichas_cadastro_to_document(document, [], "Projeto X")
        self.assertEqual(len(document._element.body.getchildren()), body_len_before)

    def test_project_dossier_photo_numbering_sequential(self):
        """T8 — multiple photos produce sequential 'Foto N' captions."""
        context = build_dossier_context()
        context["renderModel"]["sections"]["photos"] = [
            {
                "id": "RPH-A",
                "caption": "A",
                "towerId": "T-01",
                "includeInReport": True,
                "mediaAssetId": "MED-A",
            },
            {
                "id": "RPH-B",
                "caption": "B",
                "towerId": "T-01",
                "includeInReport": True,
                "mediaAssetId": "MED-B",
            },
            {
                "id": "RPH-C",
                "caption": "C",
                "towerId": "T-02",
                "includeInReport": True,
                "mediaAssetId": "MED-C",
            },
        ]
        docx = self._run_dossier(context, job_id="JOB-T8")
        plain = read_docx_plain_text(docx)
        idx1 = plain.find("Foto 1")
        idx2 = plain.find("Foto 2")
        idx3 = plain.find("Foto 3")
        self.assertGreaterEqual(idx1, 0, "Foto 1 missing")
        self.assertGreater(idx2, idx1, "Foto 2 should follow Foto 1")
        self.assertGreater(idx3, idx2, "Foto 3 should follow Foto 2")


class WorkerClientTests(unittest.TestCase):
    def test_claim_next_job_returns_none_on_204(self):
        client = WorkerClient(
            base_url="https://geomonitor-api.example.com",
            token="worker-secret",
            urlopen=lambda req, timeout=0: FakeResponse(204),
        )

        result = client.claim_next_job()

        self.assertIsNone(result)

    def test_claim_next_job_returns_data_on_200(self):
        client = WorkerClient(
            base_url="https://geomonitor-api.example.com",
            token="worker-secret",
            urlopen=lambda req, timeout=0: FakeResponse(200, {"data": {"id": "JOB-2"}}),
        )

        result = client.claim_next_job()

        self.assertEqual(result, {"id": "JOB-2"})

    def test_get_job_context_returns_payload_data(self):
        client = WorkerClient(
            base_url="https://geomonitor-api.example.com",
            token="worker-secret",
            urlopen=lambda req, timeout=0: FakeResponse(200, {"data": {"job": {"id": "JOB-CTX"}}}),
        )

        result = client.get_job_context("JOB-CTX")

        self.assertEqual(result, {"job": {"id": "JOB-CTX"}})


class ExifGpsTests(unittest.TestCase):
    def test_extracts_signed_latlon_south_west(self):
        result = extract_gps_latlon(build_minimal_jpeg_with_gps())
        self.assertIsNotNone(result)
        lat, lon = result
        self.assertAlmostEqual(lat, -22.5, places=5)
        self.assertAlmostEqual(lon, -43.25, places=5)

    def test_north_east_stays_positive(self):
        lat, lon = extract_gps_latlon(build_minimal_jpeg_with_gps(lat_ref="N", lon_ref="E"))
        self.assertAlmostEqual(lat, 22.5, places=5)
        self.assertAlmostEqual(lon, 43.25, places=5)

    def test_non_jpeg_returns_none(self):
        self.assertIsNone(extract_gps_latlon(b"definitely not a jpeg"))

    def test_png_without_exif_returns_none(self):
        self.assertIsNone(extract_gps_latlon(SAMPLE_PNG_BYTES))

    def test_zero_denominator_returns_none(self):
        self.assertIsNone(extract_gps_latlon(build_minimal_jpeg_with_gps(break_den=True)))

    def test_out_of_bounds_offset_returns_none(self):
        self.assertIsNone(extract_gps_latlon(build_minimal_jpeg_with_gps(bad_offset=True)))


if __name__ == "__main__":
    unittest.main()
