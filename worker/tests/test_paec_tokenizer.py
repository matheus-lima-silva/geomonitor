import io

import pytest
from docx import Document
from docx.enum.text import WD_COLOR_INDEX

from worker.docx_runs import collect_all_spans
from worker.tools.paec_tokenizer import (
    apply_mapping,
    build_mapping,
    document_text,
    rebase_mapping,
    render_values,
    sample_values,
)

_COLORS = {"yellow": WD_COLOR_INDEX.YELLOW, "red": WD_COLOR_INDEX.RED}


def _marked_doc():
    """Documento sintetico reproduzindo os padroes do modelo REV 10."""
    doc = Document()

    p = doc.add_paragraph()
    p.add_run("Usina de ")
    p.add_run("Marimbondo").font.highlight_color = _COLORS["yellow"]
    p.add_run(" gera energia.")

    p = doc.add_paragraph("Proteger as instalacoes da ")
    p.add_run("MARIMBONDO").font.highlight_color = _COLORS["yellow"]

    # campo fragmentado em 3 runs, como o Word costuma salvar
    p = doc.add_paragraph()
    for chunk in ("Ger", "encia ", "Regional"):
        p.add_run(chunk).font.highlight_color = _COLORS["yellow"]

    p = doc.add_paragraph("Contato: ")
    p.add_run("(34) 99730-4626").font.highlight_color = _COLORS["yellow"]

    p = doc.add_paragraph("Derramamento de ")
    p.add_run("95.814 L").font.highlight_color = _COLORS["yellow"]
    p.add_run(" de oleo isolante.")

    p = doc.add_paragraph()
    p.add_run("12.1.2. Rede de Hidrantes").font.highlight_color = _COLORS["yellow"]

    p = doc.add_paragraph()
    p.add_run("anexo VII - ROTA DE FUGA").font.highlight_color = _COLORS["yellow"]

    p = doc.add_paragraph()
    p.add_run("N PONTO DE ENCONTRO").font.highlight_color = _COLORS["red"]

    p = doc.add_paragraph()
    p.add_run("   ").font.highlight_color = _COLORS["yellow"]
    p.add_run("fim")

    return doc


def _entry(mapping, text):
    return next(e for e in mapping["spans"] if e["text"].strip() == text)


# ---------------------------------------------------------------------------
# extract / build_mapping
# ---------------------------------------------------------------------------
def test_extract_sugere_kinds_e_chaves():
    mapping = build_mapping(_marked_doc(), "sintetico.docx")

    assert mapping["stats"] == {"spans": 9, "yellow": 8, "red": 1}
    assert len(mapping["pageAudit"]) == 1

    assert _entry(mapping, "Marimbondo")["kind"] == "field"
    assert _entry(mapping, "Gerencia Regional")["kind"] == "field"
    assert _entry(mapping, "(34) 99730-4626")["key"] == "telefone_1"
    # "95.814 L" casa com o regex de heading (numero.numero) mas nao tem
    # letras suficientes para ser um titulo de secao real -> deve ficar campo
    assert _entry(mapping, "95.814 L")["kind"] == "field"
    assert _entry(mapping, "12.1.2. Rede de Hidrantes")["kind"] == "section_title"
    assert _entry(mapping, "anexo VII - ROTA DE FUGA")["kind"] == "manual"
    assert _entry(mapping, "N PONTO DE ENCONTRO")["kind"] == "list"

    blank = next(e for e in mapping["spans"] if not e["text"].strip())
    assert blank["kind"] == "whitespace"


def test_extract_agrupa_variante_maiuscula_na_mesma_chave():
    mapping = build_mapping(_marked_doc(), "sintetico.docx")
    canonical = _entry(mapping, "Marimbondo")
    upper = _entry(mapping, "MARIMBONDO")
    assert canonical["key"] == upper["key"] == "marimbondo"
    assert canonical["transform"] == "none"
    assert upper["transform"] == "upper"


def test_extract_variante_nao_derivavel_ganha_chave_propria():
    doc = Document()
    doc.add_paragraph().add_run("Elvio Zampier").font.highlight_color = _COLORS["yellow"]
    # mesmo texto normalizado (acentos ignorados), mas nao derivavel por upper/title
    doc.add_paragraph().add_run("elvio zampier").font.highlight_color = _COLORS["yellow"]
    mapping = build_mapping(doc, "s.docx")
    keys = [e["key"] for e in mapping["spans"]]
    assert keys[0] == "elvio_zampier"
    assert keys[1] == "elvio_zampier_v2"


def test_extract_textos_distintos_que_colidem_no_slug_ganham_chaves_diferentes():
    """Regressao: slugify trunca em 5 tokens e descarta pontuacao, entao
    "Gerencia de O&M Marimbondo e Colombia G Sudeste" e "...e P. Colombia G
    Sudeste" (textos DIFERENTES, nao variantes de caixa um do outro)
    colidiam no mesmo slug e se fundiam incorretamente no mesmo campo."""
    doc = Document()
    doc.add_paragraph().add_run(
        "Gerencia de O&M Marimbondo e Colombia G Sudeste"
    ).font.highlight_color = _COLORS["yellow"]
    doc.add_paragraph().add_run(
        "Gerencia de O&M Marimbondo e P. Colombia G Sudeste"
    ).font.highlight_color = _COLORS["yellow"]
    doc.add_paragraph().add_run("Telefone / Fax").font.highlight_color = _COLORS["yellow"]
    doc.add_paragraph().add_run("Telefone/Fax").font.highlight_color = _COLORS["yellow"]

    mapping = build_mapping(doc, "s.docx")
    keys = [e["key"] for e in mapping["spans"]]
    assert len(keys) == len(set(keys)), f"chaves colidiram: {keys}"

    manifest = apply_mapping(doc, mapping, "PAEC", "REV")
    assert len(manifest["fields"]) == 4
    values = sample_values(manifest)
    unresolved = render_values(doc, values)
    assert unresolved == []
    assert "Gerencia de O&M Marimbondo e Colombia G Sudeste" in document_text(doc)
    assert "Gerencia de O&M Marimbondo e P. Colombia G Sudeste" in document_text(doc)
    assert "Telefone / Fax" in document_text(doc)
    assert "Telefone/Fax" in document_text(doc)


def test_extract_canonico_nunca_e_a_variante_toda_maiuscula():
    doc = Document()
    # a variante MAIUSCULA aparece primeiro no documento
    doc.add_paragraph().add_run("MARIMBONDO").font.highlight_color = _COLORS["yellow"]
    doc.add_paragraph().add_run("Marimbondo").font.highlight_color = _COLORS["yellow"]
    mapping = build_mapping(doc, "s.docx")
    assert mapping["spans"][0]["transform"] == "upper"
    assert mapping["spans"][1]["transform"] == "none"


# ---------------------------------------------------------------------------
# apply
# ---------------------------------------------------------------------------
def test_apply_tokeniza_campos_e_limpa_highlights():
    doc = _marked_doc()
    mapping = build_mapping(doc, "sintetico.docx")
    manifest = apply_mapping(doc, mapping, "PAEC AXIA", "REV TESTE")

    text = document_text(doc)
    assert "{{marimbondo}}" in text
    assert "{{marimbondo|upper}}" in text
    assert "{{telefone_1}}" in text
    assert "{{95_814_l}}" in text
    assert "Marimbondo" not in text.replace("{{marimbondo", "")

    # sobram apenas os spans NAO tokenizados nesta fase: titulos de secao e
    # anexos manuais mantem o realce (pendencia visual) e o bloco vermelho
    # so e tokenizado na fase 2
    remaining = collect_all_spans(doc)
    assert [(s.color, s.text) for s in remaining] == [
        ("yellow", "12.1.2. Rede de Hidrantes"),
        ("yellow", "anexo VII - ROTA DE FUGA"),
        ("red", "N PONTO DE ENCONTRO"),
    ]

    field = next(f for f in manifest["fields"] if f["key"] == "marimbondo")
    assert field["sampleValue"] == "Marimbondo"
    assert len(field["occurrences"]) == 2
    assert {o["transform"] for o in field["occurrences"]} == {"none", "upper"}

    assert [(b["key"], b["kind"]) for b in manifest["blocks"]] == [
        ("anexo_vii_rota_de_fuga", "manual"),
        ("n_ponto_de_encontro", "list"),
    ]
    assert [s["sectionKey"] for s in manifest["sections"]] == ["12_1_2_rede_de"]
    assert manifest["sections"][0]["renumberGroup"] == "12.1"
    assert manifest["revisionLabel"] == "REV TESTE"


def test_apply_repassa_columns_curadas_para_o_bloco():
    """kind=list curado com `columns` (grade de brigadistas, recursos etc.)
    repassa a coluna pro manifest; sem curadoria, cai em lista vazia (Fase 2
    ainda decide as colunas manualmente na curadoria do mapping.yaml)."""
    doc = Document()
    doc.add_paragraph().add_run("ITEM | NOME").font.highlight_color = _COLORS["red"]
    mapping = build_mapping(doc, "s.docx")
    mapping["spans"][0]["columns"] = [
        {"key": "item", "label": "Item"},
        {"key": "nome", "label": "Nome"},
    ]

    manifest = apply_mapping(doc, mapping, "PAEC", "REV")

    block = manifest["blocks"][0]
    assert block["kind"] == "list"
    assert block["columns"] == [
        {"key": "item", "label": "Item"},
        {"key": "nome", "label": "Nome"},
    ]

    other_doc = Document()
    other_doc.add_paragraph().add_run("anexo VII - ROTA DE FUGA").font.highlight_color = (
        _COLORS["yellow"]
    )
    other_mapping = build_mapping(other_doc, "s.docx")
    other_manifest = apply_mapping(other_doc, other_mapping, "PAEC", "REV")
    assert other_manifest["blocks"][0]["columns"] == []


def test_apply_kind_image_vira_image_slot_com_max_images():
    """Anexo com imagem variavel por usina (rota de fuga, unifilar) curado
    como kind=image + maxImages vira imageSlot no manifest, fora de blocks.
    Sem maxImages curado, default 1 (slot simples)."""
    doc = Document()
    doc.add_paragraph().add_run("anexo VII - ROTA DE FUGA").font.highlight_color = (
        _COLORS["yellow"]
    )
    doc.add_paragraph().add_run("anexo X - UNIFILAR").font.highlight_color = (
        _COLORS["yellow"]
    )
    mapping = build_mapping(doc, "s.docx")
    mapping["spans"][0]["kind"] = "image"
    mapping["spans"][0]["maxImages"] = 5
    mapping["spans"][1]["kind"] = "image"

    manifest = apply_mapping(doc, mapping, "PAEC", "REV")

    assert manifest["blocks"] == []
    assert len(manifest["imageSlots"]) == 2
    rota, unifilar = manifest["imageSlots"]
    assert rota["assetKey"] == "anexo_vii_rota_de_fuga"
    assert rota["maxImages"] == 5
    assert rota["anchorContext"] == "anexo VII - ROTA DE FUGA"
    assert rota["page"] == {"preset": "a4-portrait"}
    assert unifilar["maxImages"] == 1


def test_apply_generated_list_blocks_entram_no_manifest_sem_span():
    """Tabelas sem NENHUM span marcado (nem o cabecalho) -- ex. contatos
    internos/externos no REV 10 -- entram no manifest via
    mapping['generatedListBlocks'], curado a parte de `spans`. Sem
    anchorContext (nao ha span pra apontar); localizadas no render por
    headerMatch."""
    doc = Document()
    mapping = build_mapping(doc, "s.docx")
    mapping["generatedListBlocks"] = [
        {
            "key": "contatos_internos",
            "label": "Contatos internos",
            "headerMatch": ["APOIO", "TELEFONES DE CONTATO"],
            "columns": [{"key": "apoio", "label": "Apoio"}, {"key": "telefones", "label": "Telefones"}],
        },
    ]

    manifest = apply_mapping(doc, mapping, "PAEC", "REV")

    assert len(manifest["blocks"]) == 1
    block = manifest["blocks"][0]
    assert block == {
        "key": "contatos_internos",
        "kind": "list",
        "label": "Contatos internos",
        "anchorContext": None,
        "headerMatch": ["APOIO", "TELEFONES DE CONTATO"],
        "columns": [{"key": "apoio", "label": "Apoio"}, {"key": "telefones", "label": "Telefones"}],
        "spanIds": [],
    }


def test_apply_ignore_mantem_texto_mas_remove_realce():
    """kind=ignore preserva o texto original (rotulo estatico que nao varia
    por usina) mas remove o realce de curadoria; ao contrario de kind=field,
    nao vira placeholder nem entra no manifest."""
    doc = Document()
    doc.add_paragraph().add_run("Razao Social").font.highlight_color = _COLORS["yellow"]
    mapping = build_mapping(doc, "s.docx")
    mapping["spans"][0]["kind"] = "ignore"
    mapping["spans"][0]["reviewed"] = True

    manifest = apply_mapping(doc, mapping, "PAEC", "REV")

    assert document_text(doc) == "Razao Social"
    remaining = collect_all_spans(doc)
    assert remaining == []
    assert manifest["fields"] == []
    assert manifest["stats"]["unreviewedSpans"] == 0


def test_apply_preserva_indentacao_da_capa():
    doc = Document()
    doc.add_paragraph().add_run("          Plano de Emergencia ").font.highlight_color = (
        _COLORS["yellow"]
    )
    mapping = build_mapping(doc, "s.docx")
    apply_mapping(doc, mapping, "PAEC", "REV")
    assert document_text(doc) == "          {{plano_de_emergencia}} "


def test_apply_recusa_documento_divergente_do_mapping():
    doc = _marked_doc()
    mapping = build_mapping(doc, "sintetico.docx")
    mapping["spans"][0]["text"] = "outro texto"
    with pytest.raises(ValueError, match="divergiu do mapping"):
        apply_mapping(doc, mapping, "PAEC", "REV")

    other = Document()
    other.add_paragraph().add_run("x").font.highlight_color = _COLORS["yellow"]
    with pytest.raises(ValueError, match="confira se e o MESMO arquivo"):
        apply_mapping(other, build_mapping(_marked_doc(), "s.docx"), "PAEC", "REV")


# ---------------------------------------------------------------------------
# render / round-trip
# ---------------------------------------------------------------------------
def test_render_aplica_transform_e_reporta_nao_resolvidas():
    doc = _marked_doc()
    mapping = build_mapping(doc, "s.docx")
    manifest = apply_mapping(doc, mapping, "PAEC", "REV")

    values = sample_values(manifest)
    values.pop("telefone_1")
    unresolved = render_values(doc, values)

    text = document_text(doc)
    assert "Usina de Marimbondo gera energia." in text
    assert "Proteger as instalacoes da MARIMBONDO" in text
    assert "{{telefone_1}}" in text
    assert set(unresolved) == {"telefone_1"}


def test_roundtrip_texto_identico_ao_modelo_marcado():
    original = _marked_doc()
    original_text = document_text(original)

    doc = _marked_doc()
    mapping = build_mapping(doc, "s.docx")
    manifest = apply_mapping(doc, mapping, "PAEC", "REV")

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    rendered = Document(buffer)
    unresolved = render_values(rendered, sample_values(manifest))

    assert unresolved == []
    assert document_text(rendered) == original_text


def test_roundtrip_com_campo_multilinha():
    def _build():
        doc = Document()
        p = doc.add_paragraph()
        run = p.add_run("Tel.: ")
        run.font.highlight_color = _COLORS["yellow"]
        run.add_break()
        run2 = p.add_run("17 98820-6833")
        run2.font.highlight_color = _COLORS["yellow"]
        return doc

    original_text = document_text(_build())
    doc = _build()
    mapping = build_mapping(doc, "s.docx")
    manifest = apply_mapping(doc, mapping, "PAEC", "REV")

    field = manifest["fields"][0]
    assert field["type"] == "multiline"

    unresolved = render_values(doc, sample_values(manifest))
    assert unresolved == []
    assert document_text(doc) == original_text


# ---------------------------------------------------------------------------
# rebase (evolucao de revisao)
# ---------------------------------------------------------------------------
def _rev_old_doc():
    """REV antiga sintetica: campo, heading, anexo de imagem e um campo que
    sera REMOVIDO na REV nova."""
    doc = Document()
    doc.add_paragraph().add_run("Marimbondo").font.highlight_color = _COLORS["yellow"]
    doc.add_paragraph().add_run("12.1.1. Rede de Hidrantes").font.highlight_color = _COLORS["yellow"]
    doc.add_paragraph().add_run("anexo VII - ROTA DE FUGA").font.highlight_color = _COLORS["yellow"]
    doc.add_paragraph().add_run("Campo Removido Na Rev Nova").font.highlight_color = _COLORS["yellow"]
    return doc


def _rev_new_doc():
    """REV nova: mantem Marimbondo e o anexo identicos, muda a CAIXA do
    heading (casamento aproximado), remove um campo e adiciona um inedito."""
    doc = Document()
    doc.add_paragraph().add_run("Marimbondo").font.highlight_color = _COLORS["yellow"]
    doc.add_paragraph().add_run("12.1.1. REDE DE HIDRANTES").font.highlight_color = _COLORS["yellow"]
    doc.add_paragraph().add_run("anexo VII - ROTA DE FUGA").font.highlight_color = _COLORS["yellow"]
    doc.add_paragraph().add_run("Campo Novo Da Rev").font.highlight_color = _COLORS["yellow"]
    return doc


def _curated_old_mapping():
    mapping = build_mapping(_rev_old_doc(), "rev_antiga.docx")
    entries = {e["text"].strip(): e for e in mapping["spans"]}

    campo = entries["Marimbondo"]
    campo.update({"section": "Identificação", "label": "Usina", "reviewed": True})

    heading = entries["12.1.1. Rede de Hidrantes"]
    heading["reviewed"] = True

    anexo = entries["anexo VII - ROTA DE FUGA"]
    anexo.update({"kind": "image", "maxImages": 5, "label": "Rota de fuga", "reviewed": True})

    removido = entries["Campo Removido Na Rev Nova"]
    removido.update({"label": "Sai na proxima", "reviewed": True})

    mapping["generatedListBlocks"] = [
        {"key": "contatos_internos", "label": "Contatos internos",
         "headerMatch": ["APOIO"], "columns": [{"key": "apoio", "label": "Apoio"}]},
    ]
    return mapping


def test_rebase_herda_curadoria_por_texto_identico_preservando_reviewed():
    new_mapping, report = rebase_mapping(_rev_new_doc(), _curated_old_mapping(), "rev_nova.docx")
    entries = {e["text"].strip(): e for e in new_mapping["spans"]}

    campo = entries["Marimbondo"]
    assert campo["section"] == "Identificação"
    assert campo["label"] == "Usina"
    assert campo["reviewed"] is True

    anexo = entries["anexo VII - ROTA DE FUGA"]
    assert anexo["kind"] == "image"
    assert anexo["maxImages"] == 5
    assert anexo["reviewed"] is True

    assert report["matchedExact"] == 2


def test_rebase_casamento_aproximado_herda_mas_volta_pra_nao_revisado():
    new_mapping, report = rebase_mapping(_rev_new_doc(), _curated_old_mapping(), "rev_nova.docx")
    heading = next(e for e in new_mapping["spans"] if "REDE DE HIDRANTES" in e["text"])

    # herdou kind/key da curadoria antiga (mesmo texto normalizado)...
    assert heading["kind"] == "section_title"
    assert heading["key"] == "12_1_1_rede_de"
    # ...mas caixa mudou -> conferir de novo
    assert heading["reviewed"] is False
    assert report["matchedFuzzy"] == 1


def test_rebase_reporta_ineditos_e_orfaos():
    new_mapping, report = rebase_mapping(_rev_new_doc(), _curated_old_mapping(), "rev_nova.docx")

    novo = next(e for e in new_mapping["spans"] if e["text"].strip() == "Campo Novo Da Rev")
    assert novo["reviewed"] is False
    assert novo["kind"] == "field"  # sugestao heuristica intacta
    assert [s["text"] for s in report["newSpans"]] == ["Campo Novo Da Rev"]

    assert [s["text"] for s in report["orphanedSpans"]] == ["Campo Removido Na Rev Nova"]
    assert report["orphanedSpans"][0]["key"] == "campo_removido_na_rev_nova"


def test_rebase_carrega_generated_list_blocks_e_o_apply_funciona():
    doc = _rev_new_doc()
    new_mapping, _report = rebase_mapping(doc, _curated_old_mapping(), "rev_nova.docx")
    assert new_mapping["generatedListBlocks"][0]["key"] == "contatos_internos"

    # o mapping rebased alinha com o documento novo — apply nao pode divergir
    manifest = apply_mapping(doc, new_mapping, "PAEC", "REV NOVA")
    assert [s["assetKey"] for s in manifest["imageSlots"]] == ["anexo_vii_rota_de_fuga"]
    assert manifest["imageSlots"][0]["maxImages"] == 5
    assert any(b["key"] == "contatos_internos" for b in manifest["blocks"])
