import io
import logging
import os
import re
import zipfile
from datetime import datetime, timezone

from copy import deepcopy

logger = logging.getLogger(__name__)

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

from worker.coordinate_format import format_tower_coordinate
from worker.kmz_renderer import build_tower_lookup, normalize_tower_id


MAX_IMAGE_WIDTH_CM = 15
HEADING_NUM_ID = "12"
SIGNATURE_LINE = "_________________________________________"
TEMPLATE_PATH = os.path.join(os.path.dirname(__file__), "assets", "template_relatorio.docx")


def normalize_text(value):
    if value is None:
        return ""
    if isinstance(value, str):
        return value.strip()
    if isinstance(value, (int, float, bool)):
        return str(value).strip()
    # Dicts/lists/other — don't emit their repr into the DOCX.
    return ""


def ensure_dict(value):
    return value if isinstance(value, dict) else {}


def safe_list(value):
    return value if isinstance(value, list) else []


def format_timestamp(value=None):
    if not value:
        return datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")

    text = normalize_text(value)
    if not text:
        return datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")

    try:
        parsed = datetime.fromisoformat(text.replace("Z", "+00:00"))
        if parsed.tzinfo is None:
            parsed = parsed.replace(tzinfo=timezone.utc)
        return parsed.astimezone(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")
    except ValueError:
        return text


def format_emission_date(value=None):
    if not value:
        return datetime.now(timezone.utc).strftime("%d/%m/%Y")

    text = normalize_text(value)
    if not text:
        return datetime.now(timezone.utc).strftime("%d/%m/%Y")

    try:
        parsed = datetime.fromisoformat(text.replace("Z", "+00:00"))
        if parsed.tzinfo is None:
            parsed = parsed.replace(tzinfo=timezone.utc)
        return parsed.astimezone(timezone.utc).strftime("%d/%m/%Y")
    except ValueError:
        parts = text.split("-")
        if len(parts) == 3:
            return f"{parts[2].zfill(2)}/{parts[1].zfill(2)}/{parts[0]}"
        return text


def normalize_lt_name(value, fallback="Relatorio"):
    text = normalize_text(value) or normalize_text(fallback) or "Relatorio"
    return text if text.upper().startswith("LT ") else f"LT {text}"


# NOTE: header metadata is injected by `rewrite_template_header_metadata`
# after the document is saved, because python-docx cannot traverse headers
# that are linked to the previous section (they return empty proxies).


def clear_template_body(document):
    body = document._element.body
    children = list(body)
    content_sectpr = None
    start_index = None

    # Save the content section properties (inline sectPr with footerReference)
    # and remove it from the original paragraph so it doesn't create an extra section break.
    # Must scan ALL paragraphs — the first <w:p> with pPr isn't necessarily the one
    # carrying the inline sectPr.
    for child in children:
        if child.tag != qn("w:p"):
            continue
        ppr = child.find(qn("w:pPr"))
        if ppr is None:
            continue
        sectpr = ppr.find(qn("w:sectPr"))
        if sectpr is None:
            continue
        if sectpr.find(qn("w:footerReference")) is None:
            continue
        content_sectpr = deepcopy(sectpr)
        ppr.remove(sectpr)
        break

    for index, child in enumerate(children):
        if child.tag not in {qn("w:p"), qn("w:tbl")}:
            continue
        text = "".join(child.itertext()).strip()
        if text.startswith("Introdução") or text.startswith("Introducao"):
            start_index = index
            break

    if start_index is None:
        return content_sectpr

    for child in children[start_index:]:
        if child.tag == qn("w:sectPr"):
            continue
        body.remove(child)

    return content_sectpr


def insert_content_section_break(document, content_sectpr):
    """Insert the content section break after the last paragraph, before the body sectPr.

    This separates the main content (with proper margins, header, and footer)
    from the quarta capa (body sectPr with zero margins and back-cover image).
    """
    if content_sectpr is None:
        return
    body = document._element.body
    p = OxmlElement("w:p")
    ppr = OxmlElement("w:pPr")
    ppr.append(deepcopy(content_sectpr))
    p.append(ppr)
    body_sectpr = body.find(qn("w:sectPr"))
    if body_sectpr is not None:
        body_sectpr.addprevious(p)
    else:
        body.append(p)


def create_document_from_template(title, document_code="", emission_date="", revision="00"):
    # title/document_code/emission_date/revision are accepted for API
    # compatibility and consumed later by `rewrite_template_header_metadata`.
    del title, document_code, emission_date, revision
    if os.path.exists(TEMPLATE_PATH):
        document = Document(TEMPLATE_PATH)
        content_sectpr = clear_template_body(document)
        return document, True, content_sectpr

    return Document(), False, None


_WT_TEXT_RE = re.compile(r"<w:t(?P<attrs>(?:\s[^>]*)?)>(?P<text>[^<]*)</w:t>")


def replace_header_xml_value_after_marker(xml, marker, value):
    """Replace the value portion of a labelled field in header XML.

    The institutional template stores each metadata field as a label run
    ("N° do Documento:", "Emissão Inicial:", "Rev.:") followed by one or
    more sibling runs inside the same `<w:p>` that carry the current value
    (sometimes split across several `<w:t>` elements for date separators,
    etc.).

    We locate the label by literal substring, then:
      1. Find every `<w:t>` that follows the label within the same paragraph.
      2. Write `value` into the FIRST such `<w:t>` that isn't whitespace-only
         (so we overwrite the first real value token rather than a space).
      3. Blank out every subsequent `<w:t>` in the same paragraph, so stale
         fragments from the template (e.g. "26/01/2026") don't leak through.
    """
    if not marker:
        return xml
    marker_index = xml.find(marker)
    if marker_index < 0:
        return xml

    # Find the `<w:t>` containing the marker to anchor our scan.
    search_from = marker_index
    for match in _WT_TEXT_RE.finditer(xml, marker_index):
        if marker in match.group("text"):
            search_from = match.end()
            break

    # Find the paragraph boundary so we never bleed into the next field.
    para_end = xml.find("</w:p>", search_from)
    if para_end < 0:
        para_end = len(xml)

    result_parts = [xml[:search_from]]
    cursor = search_from
    written = False
    for match in _WT_TEXT_RE.finditer(xml, search_from, para_end):
        result_parts.append(xml[cursor:match.start()])
        attrs = match.group("attrs") or ""
        current_text = match.group("text")
        if not written and current_text.strip():
            # First real value run: replace it with the new value.
            if "xml:space" not in attrs:
                attrs = f'{attrs} xml:space="preserve"'
            result_parts.append(f"<w:t{attrs}>{value}</w:t>")
            written = True
        elif written:
            # Blank out trailing value fragments from the template.
            result_parts.append(f"<w:t{attrs}></w:t>")
        else:
            # Whitespace-only run preceding the value — keep it intact.
            result_parts.append(match.group(0))
        cursor = match.end()
    result_parts.append(xml[cursor:])
    return "".join(result_parts)


def replace_header_lt_title(xml, title):
    paragraphs = re.findall(r"<w:p\b[^>]*>[\s\S]*?</w:p>", xml)
    for paragraph in paragraphs:
        text_runs = re.findall(r"<w:t[^>]*>([\s\S]*?)</w:t>", paragraph)
        text = "".join(text_runs).strip()
        if not text.startswith("LT "):
            continue
        run_blocks = re.findall(r"<w:r\b[^>]*>[\s\S]*?<w:t[^>]*>[\s\S]*?</w:t>[\s\S]*?</w:r>", paragraph)
        if not run_blocks:
            continue
        first_run = run_blocks[0]
        open_tag_match = re.match(r"^<w:r\b[^>]*>", first_run)
        open_tag = open_tag_match.group(0) if open_tag_match else "<w:r>"
        rpr_match = re.search(r"<w:rPr[\s\S]*?</w:rPr>", first_run)
        rpr = rpr_match.group(0) if rpr_match else ""
        replacement_run = f'{open_tag}{rpr}<w:t xml:space="preserve">{normalize_lt_name(title)}</w:t></w:r>'
        updated_paragraph = re.sub(
            r"(?:<w:r\b[^>]*>[\s\S]*?<w:t[^>]*>[\s\S]*?</w:t>[\s\S]*?</w:r>)+",
            replacement_run,
            paragraph,
            count=1,
        )
        return xml.replace(paragraph, updated_paragraph, 1)
    return xml


def apply_template_header_metadata(document, title, document_code="", emission_date="", revision="00"):
    """Injeta os metadados de cabecalho diretamente nas HeaderParts do documento.

    Substitui a antiga ``rewrite_template_header_metadata`` que abria o ZIP
    depois do ``document.save()`` e o reescrevia inteiro, gerando ~2-3s de
    I/O redundante em compostos com muitas fotos.

    Acessamos os headers via ``document.part.package.iter_parts()`` porque
    ``document.sections[i].header`` retorna um proxy vazio para headers
    ligados a secoes anteriores (``link_to_previous``). Iterando as parts
    diretamente vemos todos os ``HeaderPart`` reais do pacote.

    Deve ser chamada **antes** de ``document.save()``.
    """

    normalized_title = normalize_lt_name(title)
    normalized_code = normalize_text(document_code)
    normalized_date = format_emission_date(emission_date)
    normalized_revision = normalize_text(revision) or "00"

    package = getattr(getattr(document, "part", None), "package", None)
    if package is None or not hasattr(package, "iter_parts"):
        return

    for part in package.iter_parts():
        partname = str(getattr(part, "partname", ""))
        if not partname.startswith("/word/header") or not partname.endswith(".xml"):
            continue
        try:
            blob = part.blob
        except Exception:  # pragma: no cover - defensive
            continue
        try:
            xml = blob.decode("utf-8")
        except (AttributeError, UnicodeDecodeError):
            continue
        if "N\u00b0 do Documento:" not in xml:
            continue

        updated_xml = replace_header_lt_title(xml, normalized_title)
        updated_xml = replace_header_xml_value_after_marker(
            updated_xml, "N\u00b0 do Documento:", normalized_code
        )
        updated_xml = replace_header_xml_value_after_marker(
            updated_xml, "Emiss\u00e3o Inicial:", normalized_date
        )
        updated_xml = replace_header_xml_value_after_marker(
            updated_xml, "Rev.:", normalized_revision
        )

        if updated_xml == xml:
            continue

        try:
            new_element = parse_xml(updated_xml.encode("utf-8"))
        except Exception as exc:  # pragma: no cover - defensive
            logger.warning(
                "header_metadata_parse_failed",
                extra={
                    "partname": partname,
                    "errorType": type(exc).__name__,
                    "errorMessage": str(exc),
                },
            )
            continue

        part._element = new_element


TEMPLATE_HEADING_STYLE = "Ttulo1"
FALLBACK_HEADING_STYLE = "Heading 1"
BODY_STYLE_NAME = "NormalWeb"


def _has_style(document, style_name):
    styles = getattr(document, "styles", None)
    if styles is None:
        return False
    try:
        styles[style_name]
        return True
    except KeyError:
        return False


def _add_body_paragraph(document, text):
    """Add a body paragraph using the template's NormalWeb style when available."""
    style = BODY_STYLE_NAME if _has_style(document, BODY_STYLE_NAME) else None
    if style:
        return document.add_paragraph(text, style=style)
    return document.add_paragraph(text)


def _resolve_heading_style(document):
    """Pick the template's custom heading style if available, else a built-in.

    The institutional template ships a style named 'Ttulo1'. Documents
    created via Document() (fallback path) only have the default 'Heading 1'
    style, so attempting to use 'Ttulo1' raises KeyError.
    """
    if _has_style(document, TEMPLATE_HEADING_STYLE):
        return TEMPLATE_HEADING_STYLE
    return FALLBACK_HEADING_STYLE


# ----------------------------------------------------------------------------
# Eletrobras/Biocev formatting overrides — scope: report_compound only.
#
# The institutional template `template_relatorio.docx` ships with body/caption
# styles that drift from the Eletrobras style guide (NormalWeb is Times New
# Roman 12 with autospacing; Legenda is 9pt italic gray; section margins are
# 2.54/1.91cm instead of 4/2cm). We apply these fixes at runtime on the loaded
# Document so they don't leak into dossier/ficha renders — each render loads a
# fresh template copy, so the mutation is scoped to a single render call.
# ----------------------------------------------------------------------------
ELETROBRAS_BODY_FONT = "Arial"
ELETROBRAS_BODY_SIZE_PT = 11
ELETROBRAS_CAPTION_SIZE_PT = 10


def _clear_autospacing(style):
    """Strip w:beforeAutospacing/afterAutospacing from a style's spacing element.

    Without this, Word ignores space_before/space_after because the autospacing
    flags take precedence. NormalWeb ships with both flags enabled.
    """
    pPr = style.element.get_or_add_pPr()
    spacing = pPr.find(qn("w:spacing"))
    if spacing is None:
        return
    for attr_name in ("beforeAutospacing", "afterAutospacing"):
        key = qn(f"w:{attr_name}")
        if key in spacing.attrib:
            del spacing.attrib[key]


def _clear_run_color(style):
    """Remove any explicit w:color child on the style's rPr so color falls back to auto."""
    rPr = style.element.find(qn("w:rPr"))
    if rPr is None:
        return
    color = rPr.find(qn("w:color"))
    if color is not None:
        rPr.remove(color)


def apply_eletrobras_formatting_compound(document):
    """Mutate a template-backed Document to match the Eletrobras/Biocev guide.

    Only called from render_report_compound_docx, and only when used_template
    is True (the fallback Document() path relies on python-docx defaults that
    are already closer to the spec). Mutates margins and the custom styles
    NormalWeb, caption and Normal in place.
    """
    # --- 1. Content section margins (leave back-cover section untouched) ---
    content_sections = list(document.sections)
    if len(content_sections) > 1:
        content_sections = content_sections[:-1]
    for section in content_sections:
        section.top_margin = Cm(4)
        section.right_margin = Cm(2)
        section.left_margin = Cm(2)
        section.bottom_margin = Cm(2)
        # Header distance is already 1.25cm in the template.
        # Footer distance deliberately left at template default (user asked
        # us not to touch the footer area).

    styles = document.styles

    # --- 2. Normal: Arial 11 (was 11.5) ---
    try:
        normal = styles["Normal"]
        normal.font.name = ELETROBRAS_BODY_FONT
        normal.font.size = Pt(ELETROBRAS_BODY_SIZE_PT)
    except KeyError:
        pass

    # --- 3. NormalWeb (body paragraphs via _add_body_paragraph) ---
    # Arial 11, justified, 12pt before/after, single line spacing.
    try:
        body = styles["Normal (Web)"]
        body.font.name = ELETROBRAS_BODY_FONT
        body.font.size = Pt(ELETROBRAS_BODY_SIZE_PT)
        body.font.italic = False
        fmt = body.paragraph_format
        fmt.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        fmt.space_before = Pt(12)
        fmt.space_after = Pt(12)
        fmt.line_spacing_rule = WD_LINE_SPACING.SINGLE
        _clear_autospacing(body)
    except KeyError:
        pass

    # --- 4. Legenda (photo caption) ---
    # Arial 10, bold, not italic, no gray color.
    try:
        caption = styles["caption"]
        caption.font.name = ELETROBRAS_BODY_FONT
        caption.font.size = Pt(ELETROBRAS_CAPTION_SIZE_PT)
        caption.font.bold = True
        caption.font.italic = False
        _clear_run_color(caption)
    except KeyError:
        pass


def add_heading_paragraph(document, text, ilvl=0):
    style_name = _resolve_heading_style(document)
    p = document.add_paragraph(text, style=style_name)
    # The numPr numbering definition only exists in the template. Skip it
    # for the fallback path so the build doesn't reference an unknown numId.
    if style_name == TEMPLATE_HEADING_STYLE:
        pPr = p._p.get_or_add_pPr()
        numPr = OxmlElement('w:numPr')
        ilvl_el = OxmlElement('w:ilvl')
        ilvl_el.set(qn('w:val'), str(ilvl))
        numId_el = OxmlElement('w:numId')
        numId_el.set(qn('w:val'), HEADING_NUM_ID)
        numPr.append(ilvl_el)
        numPr.append(numId_el)
        pPr.append(numPr)
    return p


def _set_paragraph_runs_text(runs, new_text):
    """Write new_text into the first <w:t> of the first run and clear the rest.

    Defensive against runs without a <w:t> child (e.g. fldChar/br/drawing),
    which would otherwise raise AttributeError.
    """
    first_t = None
    for r in runs:
        t = r.find(qn("w:t"))
        if t is not None:
            first_t = t
            break
    if first_t is None:
        return False
    first_t.text = new_text
    # Clear every other w:t in the paragraph so stale fragments don't leak through.
    cleared_first = False
    for r in runs:
        for t in r.findall(qn("w:t")):
            if not cleared_first and t is first_t:
                cleared_first = True
                continue
            t.text = ""
    return True


def update_cover_page_body(document, lt_name, titulo_programa):
    """Replace cover textbox paragraphs via deterministic anchors.

    The institutional template has three lines per cover textbox:
      p0: "LT <...>"                                    -> lt_name
      p1: "Programa de monitoramento..."                -> titulo_programa
      p2: "Inspeção Técnica do Programa..."             -> left untouched

    We match by explicit prefixes instead of loose substring heuristics to
    avoid clobbering legitimate subtitles on customized templates.
    """
    body = document._element.body
    normalized_lt = normalize_lt_name(lt_name) if lt_name else None
    for txbx_content in body.iter(qn("w:txbxContent")):
        for p in txbx_content.findall(qn("w:p")):
            runs = p.findall(qn("w:r"))
            if not runs:
                continue
            text = "".join(
                (t.text or "") for r in runs for t in r.findall(qn("w:t"))
            ).strip()
            if not text:
                continue
            if normalized_lt and text.startswith("LT "):
                _set_paragraph_runs_text(runs, normalized_lt)
                continue
            if titulo_programa and text.startswith("Programa") and not text.startswith("Programa de"):
                # Custom program title already exists — don't overwrite.
                continue
            if titulo_programa and text.startswith("Programa de"):
                _set_paragraph_runs_text(runs, titulo_programa)


def add_cover(document, title, subtitle_lines):
    heading = document.add_heading(title or "Relatorio", level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for line in subtitle_lines:
        text = normalize_text(line)
        if not text:
            continue
        paragraph = document.add_paragraph(text)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_paragraph("")


def add_key_value_table(document, title, rows):
    if not rows:
        return

    add_heading_paragraph(document, title)
    table = document.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    for label, value in rows:
        if not normalize_text(value):
            continue
        cells = table.add_row().cells
        cells[0].text = normalize_text(label)
        cells[1].text = normalize_text(value)


def add_record_table(document, title, records, columns):
    rows = safe_list(records)
    add_heading_paragraph(document, title)
    if not rows:
        document.add_paragraph("Nenhum registro encontrado para esta secao.")
        return

    table = document.add_table(rows=1, cols=len(columns))
    table.style = "Table Grid"
    for index, column in enumerate(columns):
        table.rows[0].cells[index].text = column["label"]

    for record in rows:
        record = ensure_dict(record)
        values = []
        for column in columns:
            value = column["getter"](record)
            values.append(normalize_text(value))

        if not any(values):
            continue

        table_row = table.add_row().cells
        for index, value in enumerate(values):
            table_row[index].text = value


def add_workspace_summary(document, workspaces):
    rows = safe_list(workspaces)
    add_heading_paragraph(document, "Workspaces")
    if not rows:
        document.add_paragraph("Nenhum workspace encontrado para esta secao.")
        return

    table = document.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    headers = ["ID", "Nome", "Status", "Importado em"]
    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = header

    for workspace in rows:
        row = table.add_row().cells
        row[0].text = normalize_text(workspace.get("id"))
        row[1].text = normalize_text(workspace.get("nome"))
        row[2].text = normalize_text(workspace.get("status"))
        row[3].text = format_timestamp(workspace.get("importedAt"))


def add_photo_entry(document, photo, image_loader, photo_index):
    media_asset_id = normalize_text(photo.get("mediaAssetId"))
    if not media_asset_id:
        logger.warning(
            "photo_missing_media_id",
            extra={"photoIndex": photo_index},
        )
        document.add_paragraph("[Imagem indisponível]")
    else:
        try:
            media = image_loader(media_asset_id)
            buffer = media.get("buffer") if isinstance(media, dict) else None
            if not buffer:
                raise ValueError("conteudo vazio")
            document.add_picture(io.BytesIO(buffer), width=Cm(MAX_IMAGE_WIDTH_CM))
            document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception as exc:  # pragma: no cover - defensive fallback
            logger.warning(
                "photo_load_failed",
                extra={
                    "mediaAssetId": media_asset_id,
                    "photoIndex": photo_index,
                    "errorType": type(exc).__name__,
                    "errorMessage": str(exc),
                },
            )
            document.add_paragraph("[Imagem indisponível]")

    caption = normalize_text(photo.get("caption"))
    caption_style = "Legenda" if _has_style(document, "Legenda") else None
    paragraph = document.add_paragraph(style=caption_style) if caption_style else document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run_prefix = paragraph.add_run("Foto ")
    run_prefix.bold = False

    fld_char_begin = OxmlElement('w:fldChar')
    fld_char_begin.set(qn('w:fldCharType'), 'begin')
    run_field_begin = paragraph.add_run()
    run_field_begin._r.append(fld_char_begin)

    instr_text = OxmlElement('w:instrText')
    instr_text.set(qn('xml:space'), 'preserve')
    instr_text.text = ' SEQ Foto \\* ARABIC '
    run_instr = paragraph.add_run()
    run_instr._r.append(instr_text)

    fld_char_separate = OxmlElement('w:fldChar')
    fld_char_separate.set(qn('w:fldCharType'), 'separate')
    run_field_sep = paragraph.add_run()
    run_field_sep._r.append(fld_char_separate)

    paragraph.add_run(str(photo_index))

    fld_char_end = OxmlElement('w:fldChar')
    fld_char_end.set(qn('w:fldCharType'), 'end')
    run_field_end = paragraph.add_run()
    run_field_end._r.append(fld_char_end)

    if caption:
        paragraph.add_run(f" - {caption}")


def add_photos_section(
    document,
    photos,
    image_loader,
    section_title="ILUSTRAÇÃO FOTOGRÁFICA",
    group_by_tower=True,
    tower_lookup=None,
    coordinate_format=None,
):
    rows = [photo for photo in safe_list(photos) if photo.get("includeInReport") is True]

    add_heading_paragraph(document, section_title, ilvl=0)

    if not rows:
        logger.info(
            "photos_section_empty",
            extra={"sectionTitle": section_title},
        )
        _add_body_paragraph(document, "Nenhuma foto marcada para inclusao no relatorio.")
        return

    logger.info(
        "photos_section_start",
        extra={
            "sectionTitle": section_title,
            "photoCount": len(rows),
            "groupByTower": bool(group_by_tower),
            "coordinateFormat": coordinate_format or "",
        },
    )

    if group_by_tower:
        grouped = []
        lookup = {}
        for photo in rows:
            tower_id = normalize_text(photo.get("towerId"))
            if tower_id:
                group_key = f"Torre {tower_id}"
            else:
                group_key = "Fotos sem agrupamento"
            if group_key not in lookup:
                lookup[group_key] = []
                grouped.append((group_key, lookup[group_key]))
            lookup[group_key].append(photo)

        logger.info(
            "photos_section_grouped",
            extra={
                "sectionTitle": section_title,
                "groupCount": len(grouped),
                "groupSizes": [len(items) for _, items in grouped],
            },
        )

        photo_index = 1
        for group_label, items in grouped:
            add_heading_paragraph(document, group_label, ilvl=1)
            if tower_lookup and coordinate_format and items:
                raw_tower_id = items[0].get("towerId") if isinstance(items[0], dict) else None
                tid = normalize_tower_id(raw_tower_id)
                tower = tower_lookup.get(tid) if tid else None
                if tower:
                    coord_str = format_tower_coordinate(
                        tower.get("latitude"),
                        tower.get("longitude"),
                        coordinate_format,
                    )
                    if coord_str:
                        _add_body_paragraph(document, f"Coordenada: {coord_str}")
            for photo in items:
                add_photo_entry(document, photo, image_loader, photo_index)
                photo_index += 1
    else:
        photo_index = 1
        for photo in rows:
            add_photo_entry(document, photo, image_loader, photo_index)
            photo_index += 1


def build_project_summary_rows(project, defaults):
    project = project if isinstance(project, dict) else {}
    defaults = defaults if isinstance(defaults, dict) else {}
    return [
        ("Empreendimento", normalize_text(project.get("nome")) or normalize_text(project.get("id"))),
        ("Codigo", normalize_text(project.get("id"))),
        ("Buffer faixa (m)", str(defaults.get("faixaBufferMetersSide") or 200)),
        ("Raio sugestao torre (m)", str(defaults.get("towerSuggestionRadiusMeters") or 300)),
    ]


def first_nonempty(*values):
    for value in values:
        text = normalize_text(value)
        if text:
            return text
    return ""


def resolve_template_metadata(defaults=None, source=None):
    defaults_texts = ensure_dict(ensure_dict(defaults).get("textosBase"))
    payload = ensure_dict(source)
    return {
        "document_code": first_nonempty(
            payload.get("codigo_documento"),
            payload.get("codigoRt"),
            payload.get("codigo_rt"),
            payload.get("numeroDocumento"),
            defaults_texts.get("codigoRt"),
            defaults_texts.get("codigo_rt"),
            defaults_texts.get("numeroDocumento"),
        ),
        "revision": first_nonempty(
            payload.get("revisao"),
            payload.get("rev"),
            defaults_texts.get("revisao"),
            defaults_texts.get("rev"),
            "00",
        ),
    }


def add_numbered_text_section(document, title, text):
    add_heading_paragraph(document, title, ilvl=0)
    for paragraph_text in text.split("\n\n"):
        stripped = normalize_text(paragraph_text)
        if stripped:
            _add_body_paragraph(document, stripped)


def format_registro(person):
    conselho = person.get("registro_conselho", "")
    estado = person.get("registro_estado", "")
    numero = person.get("registro_numero", "")
    sufixo = person.get("registro_sufixo", "")
    parts = []
    if conselho and estado:
        parts.append(f"{conselho}-{estado}")
    elif conselho:
        parts.append(conselho)
    if numero:
        parts.append(f"{numero}/{sufixo}" if sufixo else numero)
    return " ".join(parts)


def add_signature_block(document, elaboradores, revisores):
    # Only start a new page if there is signature content AND there is
    # already content on the current page; otherwise we risk emitting a
    # blank page (e.g. when the previous section ended with its own break).
    if safe_list(elaboradores) or safe_list(revisores):
        has_preceding_content = any(
            (p.text or "").strip() for p in document.paragraphs
        )
        if has_preceding_content:
            document.add_page_break()

    def _add_centered(text, bold=False, font_size=None):
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        if bold:
            run.bold = True
        if font_size:
            run.font.size = font_size
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        return p

    def _render_group(label, people):
        if not people:
            return
        label_p = document.add_paragraph()
        label_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        label_run = label_p.add_run(label)
        label_run.bold = True
        label_p.paragraph_format.space_before = Pt(12)
        label_p.paragraph_format.space_after = Pt(6)
        for person in people:
            nome = normalize_text(person.get("nome", "")) if isinstance(person, dict) else normalize_text(person)
            profissao = normalize_text(person.get("profissao", "")) if isinstance(person, dict) else ""
            registro = normalize_text(person.get("registro", "")) if isinstance(person, dict) else ""
            spacer = document.add_paragraph()
            spacer.paragraph_format.space_before = Pt(24)
            spacer.paragraph_format.space_after = Pt(0)
            _add_centered(SIGNATURE_LINE)
            _add_centered(nome, bold=True)
            if profissao:
                _add_centered(profissao)
            if registro:
                _add_centered(registro)

    _render_group("Elaborado por:", safe_list(elaboradores))
    if elaboradores and revisores:
        spacer = document.add_paragraph()
        spacer.paragraph_format.space_before = Pt(12)
        spacer.paragraph_format.space_after = Pt(0)
    _render_group("Revisto por:", safe_list(revisores))


def render_project_dossier_docx(context, output_path, image_loader):
    render_model = ensure_dict(context.get("renderModel"))
    dossier = ensure_dict(render_model.get("dossier"))
    sections = ensure_dict(render_model.get("sections"))
    project = ensure_dict(context.get("project"))
    defaults = ensure_dict(context.get("defaults"))
    job = ensure_dict(context.get("job"))
    metadata = resolve_template_metadata(defaults=defaults, source=dossier)

    document, used_template, content_sectpr = create_document_from_template(
        normalize_text(project.get("nome")) or normalize_text(project.get("id")) or normalize_text(dossier.get("nome")),
        metadata["document_code"],
        job.get("updatedAt") or job.get("createdAt"),
        metadata["revision"],
    )
    if not used_template:
        add_cover(
            document,
            normalize_text(dossier.get("nome")) or f"Dossie {normalize_text(project.get('id'))}",
            [
                normalize_text(project.get("nome")) or normalize_text(project.get("id")),
                f"Gerado em {format_timestamp()}",
            ],
        )

    add_key_value_table(document, "Resumo do Empreendimento", build_project_summary_rows(project, defaults))

    observacoes = normalize_text(dossier.get("observacoes"))
    if observacoes:
        add_numbered_text_section(document, "OBSERVACOES", observacoes)

    scope = ensure_dict(dossier.get("scopeJson"))

    if scope.get("includeLicencas", True):
        add_record_table(
            document,
            "Licencas",
            sections.get("licencas"),
            [
                {"label": "ID", "getter": lambda item: item.get("id")},
                {"label": "Orgao", "getter": lambda item: item.get("orgao") or item.get("agencia")},
                {"label": "Numero", "getter": lambda item: item.get("numero") or item.get("numeroLicenca")},
                {"label": "Status", "getter": lambda item: item.get("status")},
            ],
        )

    if scope.get("includeInspecoes", True):
        add_record_table(
            document,
            "Inspecoes",
            sections.get("inspecoes"),
            [
                {"label": "ID", "getter": lambda item: item.get("id")},
                {"label": "Inicio", "getter": lambda item: item.get("dataInicio")},
                {"label": "Fim", "getter": lambda item: item.get("dataFim")},
                {"label": "Status", "getter": lambda item: item.get("status")},
            ],
        )

    if scope.get("includeErosoes", True):
        add_record_table(
            document,
            "Erosoes",
            sections.get("erosoes"),
            [
                {"label": "ID", "getter": lambda item: item.get("id")},
                {"label": "Status", "getter": lambda item: item.get("status")},
                {"label": "Criticidade", "getter": lambda item: item.get("criticalityCode")},
                {"label": "Score", "getter": lambda item: item.get("criticalityScore")},
            ],
        )

    if scope.get("includeEntregas", True):
        add_record_table(
            document,
            "Entregas",
            sections.get("entregas"),
            [
                {"label": "ID", "getter": lambda item: item.get("id")},
                {"label": "Competencia", "getter": lambda item: item.get("monthKey")},
                {"label": "Status", "getter": lambda item: item.get("operationalStatus")},
                {"label": "Atualizado em", "getter": lambda item: item.get("updatedAt")},
            ],
        )

    if scope.get("includeWorkspaces", True):
        add_workspace_summary(document, sections.get("workspaces"))

    if scope.get("includeFotos", True):
        add_photos_section(document, sections.get("photos"), image_loader)

    if used_template:
        insert_content_section_break(document, content_sectpr)

    if used_template:
        apply_template_header_metadata(
            document,
            normalize_text(project.get("nome")) or normalize_text(project.get("id")) or normalize_text(dossier.get("nome")),
            metadata["document_code"],
            job.get("updatedAt") or job.get("createdAt"),
            metadata["revision"],
        )

    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    document.save(output_path)
    return output_path


def render_report_compound_docx(context, output_path, image_loader):
    render_model = ensure_dict(context.get("renderModel"))
    compound = ensure_dict(render_model.get("compound"))
    workspaces = safe_list(render_model.get("workspaces"))
    job = ensure_dict(context.get("job"))
    shared = ensure_dict(compound.get("sharedTextsJson"))
    metadata = resolve_template_metadata(source=shared)

    lt_name = normalize_text(shared.get("nome_lt")) or normalize_text(compound.get("nome")) or "Relatorio composto"
    doc_code = metadata["document_code"]
    revision = metadata["revision"]

    document, used_template, content_sectpr = create_document_from_template(
        lt_name,
        doc_code,
        job.get("updatedAt") or job.get("createdAt"),
        revision,
    )
    if used_template:
        apply_eletrobras_formatting_compound(document)
    titulo_programa = normalize_text(shared.get("titulo_programa"))

    if used_template:
        update_cover_page_body(document, lt_name, titulo_programa)
    else:
        subtitle_lines = []
        if titulo_programa:
            subtitle_lines.append(titulo_programa)
        if doc_code:
            subtitle_lines.append(doc_code)
        subtitle_lines.append(f"Gerado em {format_timestamp()}")
        add_cover(document, lt_name, subtitle_lines)

    intro_text = normalize_text(shared.get("introducao"))
    if intro_text:
        add_numbered_text_section(document, "INTRODUÇÃO", intro_text)

    caracterizacao_subtopicos = [
        ("geologia", "Geologia"),
        ("geotecnia", "Geotecnia"),
        ("geomorfologia", "Geomorfologia"),
    ]
    subtopico_texts = [(label, normalize_text(shared.get(key))) for key, label in caracterizacao_subtopicos if normalize_text(shared.get(key))]
    legacy_caracterizacao = normalize_text(shared.get("caracterizacao_tecnica"))
    if subtopico_texts or legacy_caracterizacao:
        add_heading_paragraph(document, "CARACTERIZAÇÃO TÉCNICA", ilvl=0)
        if subtopico_texts:
            for label, text in subtopico_texts:
                add_heading_paragraph(document, label, ilvl=1)
                for paragraph_text in text.split("\n\n"):
                    stripped = normalize_text(paragraph_text)
                    if stripped:
                        _add_body_paragraph(document, stripped)
        elif legacy_caracterizacao:
            for paragraph_text in legacy_caracterizacao.split("\n\n"):
                stripped = normalize_text(paragraph_text)
                if stripped:
                    _add_body_paragraph(document, stripped)

    descricao_text = normalize_text(shared.get("descricao_atividades"))
    if descricao_text:
        add_numbered_text_section(document, "DESCRIÇÃO DAS ATIVIDADES", descricao_text)

    include_tower_coords = bool(shared.get("includeTowerCoordinates"))
    tower_coord_format = normalize_text(shared.get("towerCoordinateFormat")) or "decimal"

    for bundle in workspaces:
        if not isinstance(bundle, dict):
            continue
        photos = safe_list(bundle.get("photos"))
        if not photos:
            continue
        sort_mode = normalize_text(bundle.get("photoSortMode")) or "tower_asc"
        use_tower_grouping = sort_mode.startswith("tower")
        ws = ensure_dict(bundle.get("workspace"))
        ws_name = normalize_text(ws.get("nome")) or normalize_text(ws.get("id"))
        section_title = f"ILUSTRAÇÃO FOTOGRÁFICA - {ws_name}" if len(workspaces) > 1 and ws_name else "ILUSTRAÇÃO FOTOGRÁFICA"

        tower_lookup = None
        if include_tower_coords and use_tower_grouping:
            project = ensure_dict(bundle.get("project"))
            tower_lookup = build_tower_lookup(project) or None

        add_photos_section(
            document,
            photos,
            image_loader,
            section_title=section_title,
            group_by_tower=use_tower_grouping,
            tower_lookup=tower_lookup,
            coordinate_format=tower_coord_format if include_tower_coords else None,
        )

    post_photo_sections = [
        ("CONCLUSÕES E RECOMENDAÇÕES", "conclusoes"),
        ("ANÁLISE DA EVOLUÇÃO DOS PROCESSOS EROSIVOS", "analise_evolucao"),
        ("CONSIDERAÇÕES FINAIS", "observacoes"),
    ]
    for title, key in post_photo_sections:
        text = normalize_text(shared.get(key))
        if text:
            add_numbered_text_section(document, title, text)

    elaboradores = safe_list(shared.get("elaboradores"))
    revisores = safe_list(shared.get("revisores"))
    if elaboradores or revisores:
        add_signature_block(document, elaboradores, revisores)

    if used_template:
        insert_content_section_break(document, content_sectpr)

    if used_template:
        apply_template_header_metadata(
            document,
            lt_name,
            doc_code,
            job.get("updatedAt") or job.get("createdAt"),
            revision,
        )

    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    document.save(output_path)
    return output_path


def render_context_to_docx(context, output_path, image_loader):
    job = context.get("job") if isinstance(context, dict) else {}
    kind = normalize_text(job.get("kind"))
    if kind == "project_dossier":
        return render_project_dossier_docx(context, output_path, image_loader)
    if kind == "report_compound":
        return render_report_compound_docx(context, output_path, image_loader)
    if kind == "ficha_cadastro":
        from worker.ficha_cadastro_renderer import render_ficha_cadastro_docx
        return render_ficha_cadastro_docx(context, output_path, image_loader)
    raise ValueError(f"Tipo de job sem renderizador DOCX: {kind or 'desconhecido'}")
